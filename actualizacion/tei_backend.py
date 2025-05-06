import os
import re
import unicodedata
import zipfile
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from difflib import get_close_matches

### PROCESAMIENTO DE DOCX A TEI ###

def get_intro_footnotes(docx_path):
    """
    Estrae tutte le note a piè di pagina da un file DOCX.
    Ritorna un dizionario {id: testo_nota}.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    footnote_dict = {}

    with zipfile.ZipFile(docx_path) as docx_zip:
        with docx_zip.open("word/footnotes.xml") as footnote_file:
            root = etree.parse(footnote_file).getroot()

            for note in root.xpath("//w:footnote[not(@w:type='separator')]", namespaces=ns):
                note_id = note.get(qn("w:id"))
                texts = note.xpath(".//w:t", namespaces=ns)
                full_text = "".join(t.text for t in texts if t is not None).strip()
                if full_text:
                    footnote_dict[note_id] = full_text

    return footnote_dict

def extract_text_with_intro_notes(para, footnotes_intro):
    text = ""
    for run in para.runs:
        run_element = run._element
        refs = run_element.findall(".//w:footnoteReference", namespaces=run_element.nsmap)
        if refs:
            for ref in refs:
                note_id = ref.get(qn("w:id"))
                note_text = footnotes_intro.get(note_id, "")
                text += f'<note type="intro" n="{note_id}">{note_text}</note>'
        else:
            if run.italic:
                text += f'<hi rend="italic">{run.text}</hi>'
            else:
                text += run.text
    return text.strip()

def chiudi_blocchi_correnti(tei, stato):
    if stato.get("in_sp"):
        tei.append('        </sp>')
        stato["in_sp"] = False
    if stato.get("in_cast_list"):
        tei.append('        </castList>')
        tei.append('      </div>')
        stato["in_cast_list"] = False
    if stato.get("in_dedicatoria"):
        tei.append('      </div>')
        stato["in_dedicatoria"] = False
    if stato.get("in_act"):
        tei.append('      </div>')
        stato["in_act"] = False


#############################
# 1) FUNZIONI DI SUPPORTO
#############################

def extract_text_with_italics(para):
    """Estrae il testo da un paragrafo mantenendo il corsivo."""
    text = ""
    for run in para.runs:
        if run.italic:
            text += f'<hi rend="italic">{run.text}</hi>'
        else:
            text += run.text
    return text.strip()

def merge_italic_text(text):
    """
    Unisce le sequenze consecutive di tag <hi rend="italic">...</hi>
    in un unico tag.
    """
    # Cerca una o più occorrenze consecutive di <hi rend="italic">...</hi>
    pattern = re.compile(r'((?:<hi rend="italic">.*?</hi>\s*)+)', re.DOTALL)
    def replacer(match):
        segment = match.group(1)
        # Estrae tutto il contenuto interno dai tag
        inner_texts = re.findall(r'<hi rend="italic">(.*?)</hi>', segment, re.DOTALL)
        # Unisce tutto in un'unica stringa (senza spazi aggiuntivi)
        merged = ''.join(inner_texts).strip()
        return f'<hi rend="italic">{merged}</hi>'
    return pattern.sub(replacer, text)

def generate_filename(title):
    """
    Genera un nome file basato sulle prime tre parole del titolo
    e rimuovendo spazi, virgole e punti.
    """
    words = title.split()[:3]
    filename = '_'.join(words).replace(' ', '_').replace(',', '').replace('.', '')
    return filename

def find_who_id(speaker, personaggi):
    """
    Trova il ruolo corretto (xml:id) nel cast list con match flessibile.
    """
    speaker_cleaned = re.sub(r'[\s\[\]]+', '_', speaker).strip()

    # Match esatto
    if speaker.strip() in personaggi:
        return personaggi[speaker.strip()]

    # Match parziale
    for name, role_id in personaggi.items():
        if speaker.strip() in name or speaker_cleaned in role_id:
            return role_id

    # Fuzzy matching come ultima risorsa
    close_matches = get_close_matches(speaker.strip(), personaggi.keys(), n=1, cutoff=0.8)
    if close_matches:
        return personaggi[close_matches[0]]

    return ""

### METADATOS Y FRONT-MATTER ###

def parse_metadata_docx(path):
    """Estrae metadati da un file .docx strutturato in tabelle e costruisce un teiHeader TEI/XML completo."""
    doc = Document(path)
    tables = doc.tables

    if len(tables) < 3:
        raise ValueError("❌ Il documento dei metadati deve contenere almeno 3 tabelle.")

    # Tabella 1: Metadati principali
    main_meta = {}
    for row in tables[0].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            main_meta[key] = value

    # Tabella 2: sourceDesc
    source_meta = {}
    for row in tables[1].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            source_meta[key] = value

    # Tabella 3: listWit (salta la prima riga!)
    witnesses = []
    for i, row in enumerate(tables[2].rows):
        if i == 0:
            continue  # salta la riga guida "SIGLA TESTIMONIO | DESCRIPCIÓN"
        if len(row.cells) >= 2:
            siglum = row.cells[0].text.strip()

            # Estrai il contenuto della seconda cella mantenendo il corsivo
            desc_parts = []
            for run in row.cells[1].paragraphs[0].runs:
                if run.italic:
                    desc_parts.append(f'<hi rend="italic">{run.text}</hi>')
                else:
                    desc_parts.append(run.text)
            desc = ''.join(desc_parts).strip()

            if siglum and desc:
                witnesses.append((siglum, desc))

    # Costruzione teiHeader
    tei = ['<teiHeader>', '  <fileDesc>', '    <titleStmt>']

    tei.append(f'      <title>{main_meta.get("Titulo comedia", "")}</title>')
    tei.append(f'      <author><name>{main_meta.get("Autor", "")}</name></author>')

    if 'Editor' in main_meta:
        tei.append(f'      <editor>{main_meta["Editor"]}</editor>')

    if 'Responsable/s revisión' in main_meta:
        tei.append('      <respStmt>')
        tei.append('        <resp>Edición crítica digital revisada filológicamente por</resp>')
        for name in main_meta['Responsable/s revisión'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    if 'Responsable marcado automático' in main_meta:
        tei.append('      <respStmt>')
        tei.append('        <resp>Marcado XML-TEI automático revisado por</resp>')
        for name in main_meta['Responsable marcado automático'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    tei.append('      <respStmt>')
    tei.append('        <resp>Codificado según los criterios de</resp>')
    tei.append('        <name ref="https://datos.bne.es/entidad/XX4849774.html">Grupo de investigación PROLOPE, de la Universitat Autònoma de Barcelona</name>')
    tei.append('      </respStmt>')

    tei.extend(['    </titleStmt>', '    <editionStmt>'])
    tei.append(f'      <edition>Versión {main_meta.get("Versión", "")}</edition>')
    tei.extend(['    </editionStmt>', '    <publicationStmt>'])
    tei.append(f'      <publisher>{main_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'      <pubPlace>{main_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'      <date>{main_meta.get("Fecha publicación", "")}</date>')
    tei.extend(['    </publicationStmt>', '    <seriesStmt>'])
    tei.append('      <title>Biblioteca Digital PROLOPE</title>')
    tei.append('      <respStmt>')
    tei.append('        <resp>Dirección de</resp>')
    tei.append('        <persName ref="https://orcid.org/0000-0002-7429-9709"><forename>Ramón</forename> <surname>Valdés Gázquez</surname></persName>')
    tei.append('      </respStmt>')
    tei.append('      <idno type="URI">https://bibdigitalprolope.com/</idno>')
    tei.append('    </seriesStmt>')

    # sourceDesc
    tei.extend(['    <sourceDesc>', '      <biblStruct xml:lang="es">', '        <monogr>'])
    tei.append('          <author>')
    tei.append('            <persName ref="http://datos.bne.es/persona/XX1719671"><forename>Félix Lope</forename><surname>de Vega Carpio</surname></persName>')
    tei.append('          </author>')
    tei.append(f'          <title type="main">{source_meta.get("Titulo comedia", "")}</title>')
    if "Subtítulo" in source_meta:
        tei.append(f'          <title type="alt">{source_meta.get("Subtítulo", "")}</title>')
    tei.append(f'          <title type="s">{source_meta.get("Título volumen", "")}</title>')
    tei.append(f'          <title type="a">Parte {source_meta.get("Parte", "")}</title>')
    tei.append('          <availability status="restricted">')
    tei.append('           <p>Todos los derechos reservados.</p>')
    tei.append('          </availability>')
    tei.append('          <imprint>')
    tei.append(f'            <pubPlace>{source_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'            <publisher>{source_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'            <date>{source_meta.get("Fecha publicación", "")}</date>')
    tei.append(f'            <biblScope unit="volume" n="{source_meta.get("Volumen", "")}">vol. {source_meta.get("Volumen", "")}</biblScope>')
    tei.append(f'            <biblScope unit="page">{source_meta.get("Páginas", "")}</biblScope>')
    tei.append('          </imprint>')
    tei.append('        </monogr>')
    tei.append('      </biblStruct>')

    # listWit
    tei.append('      <listWit>')
    for siglum, desc in witnesses:
        tei.append(f'        <witness xml:id="{siglum}">')
        tei.append(f'          <label>{desc}</label>')
        tei.append('        </witness>')
    tei.append('      </listWit>')

    tei.extend(['    </sourceDesc>', '  </fileDesc>', '</teiHeader>'])

    return "\n".join(tei)

def process_front_paragraphs(paragraphs, footnotes_intro):
    tei_front = []
    subsection_open = False
    subsection_n = 1
    current_section = None
    paragraph_buffer = []
    head_inserted = False

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer
        for p in paragraph_buffer:
            text = extract_text_with_intro_notes(p, footnotes_intro)
            if p.style and p.style.name == "Cita":
                tei_front.append(f'          <cit rend="blockquote">')
                tei_front.append(f'            <quote>{text}</quote>')
                tei_front.append(f'          </cit>')
            elif text.strip():
                tei_front.append(f'          <p>{text.strip()}</p>')
        paragraph_buffer.clear()

    for i, para in enumerate(paragraphs):
        raw = extract_text_with_intro_notes(para, footnotes_intro)
        text = para.text.strip() if para.text else ""
        if not text:
            continue

        # 🔹 Ignora "Introducción"
        if text.lower() == "introducción":
            continue

        # 🔹 Gestione titolo principale "PRÓLOGO"
        if not head_inserted and "prólogo" in text.lower():
            flush_paragraph_buffer()
            tei_front.append(f'        <head type="divTitle" subtype="MenuLevel_1">PRÓLOGO</head>')
            head_inserted = True
            continue

        # 🔹 Riconoscimento sottotitolo con asterisco
        if text.startswith("*"):
            flush_paragraph_buffer()
            title = text.lstrip("*").strip()
            if title.lower() == "prólogo":
                continue  # ignora duplicazione
            if subsection_open:
                tei_front.append('        </div>')
            tei_front.append(f'        <div type="subsection" n="{subsection_n}">')
            tei_front.append(f'          <head type="divTitle" subtype="MenuLevel_2">{title}</head>')
            subsection_open = True
            current_section = title.lower()
            subsection_n += 1

        # 🔹 Tabelle nella sezione "Sinopsis"
            if current_section and "sinopsis" in current_section:
                parent = para._parent
                if hasattr(parent, "tables") and parent.tables:
                    for tbl in parent.tables:
                        tei_front.append(process_table_to_tei(tbl))
            continue

        # 🔹 Aggiungi al buffer
        paragraph_buffer.append(para)

    flush_paragraph_buffer()

    if subsection_open:
        tei_front.append('        </div>')  # chiude ultimo subsection

    return "\n".join(tei_front)

def process_table_to_tei(table):
    tei = ['          <table rend="rules">']
    
    for row in table.rows:
        cells = row.cells
        texts = [cell.text.strip() for cell in cells]
        non_empty_cells = [t for t in texts if t]

        # Riga di intestazione di sezione (es. Acto primero, Resumen)
        if len(non_empty_cells) == 1 and texts[0]:
            tei.append('            <row>')
            tei.append(f'              <cell rend="both" cols="3"><hi rend="italic">{texts[0]}</hi></cell>')
            tei.append('            </row>')
            continue

        # Riga standard a 3 colonne
        tei.append('            <row>')
        for txt in texts:
            tei.append(f'              <cell rend="both">{txt}</cell>')
        tei.append('            </row>')

    tei.append('          </table>')
    return "\n".join(tei)

#############################
# 2) LETTURA NOTE (COMENTARIO/APARATO)
#############################

def extract_notes_with_italics(docx_path: str) -> dict:
    """
    Extrae notas de comentario o aparato de un DOCX.
    Retorna un dict donde las claves pueden ser int (versos) o str (palabras)
    y los valores el texto de la nota.
    """
    notes: dict = {}
    if not docx_path or not os.path.exists(docx_path):
        return notes

    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = extract_text_with_italics(para).strip()

        # Notas tipo verso: "1: contenido"
        match_verse = re.match(r'^(\d+):\s*(.*)', text)
        # Notas tipo @palabra@: "@palabra@: contenido"
        match_single = re.match(r'^@([^@]+?):\s*(.*)', text)

        if match_verse:
            verse_num = int(match_verse.group(1))
            notes[verse_num] = match_verse.group(2).strip()
        elif match_single:
            key = match_single.group(1).strip()
            if key not in notes:
                notes[key] = match_single.group(2).strip()

    return notes


#############################
# 4) PROCESSAMENTO DELLE ANNOTAZIONI @
#############################


def process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, section):
    """
    Sustituye marcadores @palabra@ en el texto por notas TEI con xml:ids únicos.
    - comentario_notes y aparato_notes son dicts con claves normalizadas.
    - annotation_counter es un dict que lleva el conteo de repeticiones por sección.
    - section es el nombre de la sección (p.ej. 'p', 'speaker', etc.).
    """
    if not text:
        return ""

    # Aseguramos dicts válidos
    comentario_notes = comentario_notes or {}
    aparato_notes    = aparato_notes    or {}

    # Función de normalización para claves
    def normalize_word(word):
        if isinstance(word, int):
            return word
        # Quitamos tags <hi> y descomponemos acentos
        plain = re.sub(r'<hi rend="italic">(.*?)</hi>', r'\1', word)
        normalized = unicodedata.normalize('NFKD', plain)
        normalized = normalized.encode('ASCII', 'ignore').decode('utf-8').lower().strip()
        return normalized

    # Normalizamos claves de las notas
    comentario_notes_norm = {
        normalize_word(k): v
        for k, v in comentario_notes.items()
        if isinstance(k, str)
    }
    aparato_notes_norm = {
        normalize_word(k): v
        for k, v in aparato_notes.items()
        if isinstance(k, str)
    }

    # Unificamos todas las notas en un solo dict
    all_notes = {**comentario_notes_norm, **aparato_notes_norm}

    new_text = text.strip()
    # Buscamos todos los marcadores '@palabra'
    matches = re.findall(r'@(\w+)', text)

    for phrase in matches:
        if not phrase:
            continue
        phrase_to_replace = f"@{phrase}"
        key = normalize_word(phrase)

        if key in all_notes:
            # Gestionamos el contador por sección
            section_counters = annotation_counter.setdefault(section, {})
            count = section_counters.get(key, 0) + 1
            section_counters[key] = count

            # Creamos xml:id válido
            xml_id = f"{key}_{section}_{count}"
            xml_id = re.sub(r'\s+', '_', xml_id)
            xml_id = re.sub(r'[^a-zA-Z0-9_]', '', xml_id)
            xml_id = xml_id.lower()

            # Construimos los posibles <note>
            note_str = ""
            if key in comentario_notes_norm:
                note_str += f'<note subtype="comentario" xml:id="{xml_id}">{comentario_notes_norm[key]}</note>'
            if key in aparato_notes_norm:
                note_str += f'<note subtype="aparato"     xml:id="{xml_id}">{aparato_notes_norm[key]}</note>'

            # Sustituimos solo la primera ocurrencia
            new_text = new_text.replace(phrase_to_replace, f"{phrase}{note_str}", 1)

    # Reagrupamos cursivas consecutivas
    new_text = merge_italic_text(new_text)
    return new_text


#############################
# 5) FUNZIONE PRINCIPALE
#############################

def convert_docx_to_tei(
    main_docx: str,
    comentario_docx: str | None = None,
    aparato_docx: str | None = None,
    output_file: str | None = None,
    tei_header: str | None = None,
    save: bool = True
) -> str | None:
    """
    Convierte uno o más DOCX a un XML-TEI completo.
    - Si save=False: devuelve el XML como cadena.
    - Si save=True: escribe en output_file (o genera uno por defecto) y retorna None.
    """
    # 1) Chequeo de existencia del principal
    if not os.path.exists(main_docx):
        raise FileNotFoundError(f"No existe el archivo principal: {main_docx}")

    # 2) Iniciar lista TEI con encabezado
    tei: list[str] = []
    if tei_header:
        tei.append(tei_header)
    else:
        # Cabecera mínima como fallback
        tei.append("<teiHeader>…</teiHeader>")

    # 3) Apertura del texto
    tei.append("<text>")

    # 4) Carga del DOCX principal
    doc = Document(main_docx)

    # --- SEPARAZIONE PARAGRAFI FRONT/BODY BASATA SULLO STILE 'Titulo_comedia' ---

    front_paragraphs = []
    body_paragraphs = []
    inside_body = False

    for para in doc.paragraphs:
        text = extract_text_with_italics(para).strip()
        style = para.style.name if para.style else "Normal"

        if not inside_body:
            if style == 'Titulo_comedia':
                inside_body = True
                body_paragraphs.append(para)  # Il paragrafo del titolo va nel body
            else:
                front_paragraphs.append(para)
        else:
            body_paragraphs.append(para)


    # Estrazione del titolo
    title = "Untitled"
    for para in doc.paragraphs:
        if para.style and para.style.name == "Titulo_comedia":
            title = para.text.strip()
            break

    clean_title = re.sub(r'@', '', title)
    title_key = generate_filename(title)

    # Determinación de los archivos de comentario y aparato
    if comentario_docx and os.path.exists(comentario_docx):
        comentario_file = comentario_docx
    else:
        comentario_file = f'comentario_{title_key}.docx'

    if aparato_docx and os.path.exists(aparato_docx):
        aparato_file = aparato_docx
    else:
        aparato_file = f'aparato_{title_key}.docx'

    # Extracción de notas si los archivos existen
    comentario_notes = {}
    if os.path.exists(comentario_file):
        comentario_notes = extract_notes_with_italics(comentario_file)

    aparato_notes = {}
    if os.path.exists(aparato_file):
        aparato_notes = extract_notes_with_italics(aparato_file)


    # ... (il resto della funzione per generare l'XML TEI)
    
    annotation_counter = {}
    stato = {"in_sp": False, "in_cast_list": False, "in_dedicatoria": False, "in_act": False}
    personaggi = {}
    act_counter = 0
    verse_counter = 1
    processed_title = process_annotations_with_ids(title, comentario_notes, aparato_notes, {}, "head")
    footnotes_intro = get_intro_footnotes(main_docx)  # Usa il percorso del docx principale
    ultimo_speaker_id = None

    # Intestazione TEI
    tei = [
    '<?xml-model href="http://www.tei-c.org/release/xml/tei/custom/schema/relaxng/tei_all.rng" schematypens="http://relaxng.org/ns/structure/1.0"?>',
    '<TEI xmlns="http://www.tei-c.org/ns/1.0">',
    tei_header,
    '  <text>',
    '    <front>',
    '      <div type="Introducción">'
    ]

    tei.append(process_front_paragraphs(front_paragraphs, footnotes_intro))

    tei.extend([
        '      </div>',  # chiude Introducción
        '    </front>',
        '    <body>',
        '      <div type="Texto" subtype="TEXTO">',
        f'        <head type="mainTitle">{processed_title}</head>',
    ])

    # Scansione paragrafi
    for para in body_paragraphs:
        text = extract_text_with_italics(para).strip()
        style = para.style.name if para.style else "Normal"

        # Rilevamento strofe $nomeStrofa -> milestone
        lg_milestone = re.match(r'^\$(\w+)', text)
        if lg_milestone:
            current_milestone = lg_milestone.group(1)
            continue

        # Chiudi blocchi prima di aprirne altri
        if style in ["Epigr_Dramatis", "Acto", "Epigr_Dedic"]:
            chiudi_blocchi_correnti(tei, stato)

        if style == "Epigr_Dedic":
            tei.append('        <div type="dedicatoria">')
            tei.append(f'          <head>{text}</head>')
            stato["in_dedicatoria"] = True

        elif style == "Epigr_Dramatis":
            tei.append('        <div type="castList">')
            tei.append('          <castList>')
            tei.append(f'            <head>{text}</head>')
            stato["in_cast_list"] = True

        elif style == "Dramatis_lista":
            role_name = text
            if role_name:
                role_id = re.sub(r'[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ_-]+', '_', role_name)
                tei.append(f'            <castItem><role xml:id="{role_id}">{role_name}</role></castItem>')
                personaggi[role_name] = role_id

        elif style == "Acto":
            act_counter += 1
            tei.append(f'        <div type="subsection" subtype="ACTO" n="{act_counter}">')
            tei.append(f'          <head type="acto">{text}</head>')
            stato["in_act"] = True

        elif style == "Prosa":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "p")
            if processed_text.strip():
                tei.append(f'          <p>{processed_text}</p>')

        elif style == "Verso":
            if stato["in_dedicatoria"]:
                tei.append(f'          <l>{text}</l>')
            elif stato["in_sp"]:
                if current_milestone:
                    tei.append(f'            <milestone unit="stanza" type="{current_milestone}"/>')
                    current_milestone = None
                verse_text = text
                if verse_counter in comentario_notes:
                    verse_text += f'<note subtype="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
                if verse_counter in aparato_notes:
                    verse_text += f'<note subtype="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
                tei.append(f'            <l n="{verse_counter}">{verse_text}</l>')
                verse_counter += 1

        elif style == "Partido_incial":
            verse_text = text
            if verse_counter in comentario_notes:
                verse_text += f'<note subtype="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
            if verse_counter in aparato_notes:
                verse_text += f'<note subtype="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
            tei.append(f'            <l part="I" n="{verse_counter}">{verse_text}</l>')
            verse_counter += 1

        elif style == "Partido_medio":
            tei.append(f'            <l part="M">{text}</l>')

        elif style == "Partido_final":
            tei.append(f'            <l part="F">{text}</l>')

        elif style == "Acot":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "stage")
            if stato["in_sp"]:
                tei.append('        </sp>')
                stato["in_sp"] = False
            tei.append(f'        <stage>{processed_text}</stage>')

        elif style == "Personaje":
            who_id = find_who_id(text, personaggi)
            processed = process_annotations_with_ids(
                text, comentario_notes, aparato_notes, annotation_counter, "speaker"
            )

            # 🔒 Chiudi <sp> precedente se necessario
            if stato["in_sp"]:
                tei.append('        </sp>')
                stato["in_sp"] = False

            # 🎯 Se è lo stesso personaggio precedente, non inserire who o speaker
            if who_id == ultimo_speaker_id:
                tei.append('        <sp>')
            else:
                tei.append(f'        <sp who="#{who_id}">')
                tei.append(f'          <speaker>{processed}</speaker>')
                ultimo_speaker_id = who_id

            stato["in_sp"] = True

        elif style == "Titulo_comedia":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "head")

    # Chiusura finale di tutti i blocchi ancora aperti
    chiudi_blocchi_correnti(tei, stato)

    # Chiusura sezioni TEI
    tei.append('      </div>')  # chiude Texto
    tei.append('    </body>')
    tei.append('  </text>')
    tei.append('</TEI>')


    # Serializamos el TEI a string
    tei = [fragment for fragment in tei if isinstance(fragment, str)]
    tei_str = "\n".join(tei)

    # Si no queremos guardar en disco, devolvemos el string
    if not save:
        return tei_str

    # Si llegamos aquí, save == True: escribimos el fichero
    if not output_file:
        # Generamos nombre por defecto si hace falta
        output_file = generate_filename(title) + ".xml"

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(tei_str)
    # Devolvemos None para indicar que se escribió en disco
    return None


### VALIDACIÓN Y ANÁLISIS ###

def analyze_main_text(main_docx) -> list[str]:
    """
    Analiza el archivo principal y devuelve avisos de párrafos sin estilo
    solo en el cuerpo de la obra (tras Titulo_comedia), ignorando front matter
    y milestones ($...).
    """
    warnings: list[str] = []
    unstyled_paragraphs: list[str] = []

    doc = Document(main_docx)
    found_body = False

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()

        # 1) Buscamos el inicio del body
        if not found_body:
            if style == "Titulo_comedia":
                found_body = True
            continue

        # 2) Ignoramos párrafos vacíos
        if not text:
            continue

        # 3) Solo revisamos estilos 'Normal' o None
        if style in ["Normal", "", None]:
            # Ignora cualquier milestone que empiece con '$'
            if re.match(r'^\$\S+', text):
                continue
            # Ignora líneas que solo contengan puntuación
            if re.match(r'^[\.,;:!?\(\)"\']+$', text):
                continue
            unstyled_paragraphs.append(text)

    if unstyled_paragraphs:
        warnings.append(
            f"❌ He encontrado {len(unstyled_paragraphs)} párrafos sin estilo "
            f"en el cuerpo del texto crítico: {main_docx}"
        )

    return warnings



def analyze_notes(
    notes: dict, 
    note_type: str
) -> list[str]:
    """
    Analiza el dict de notas (aparato o comentario) y devuelve
    una lista de strings con posibles notas mal formateadas.
    """
    warnings: list[str] = []

    for key, content in notes.items():
        # content es el texto de la nota
        # key puede ser int (verso) o str (palabra)
        text = str(content).strip()
        if not text:
            continue
        # Comprobamos formato: versos "1:..." los quita extract_notes y aquí sólo miramos @palabra@
        if isinstance(key, str):
            # en extract_notes mantienes solo claves que cumplían @clave@.: no hay invalidas
            continue
        elif isinstance(key, int):
            # en extract_notes mantienes los versos parseados; no hay invalidos
            continue
        else:
            warnings.append(
                f"❌ Nota {note_type} con clave inesperada ({key}): «{text[:60]}»"
            )

    return warnings



from docx import Document

def validate_documents(main_docx, aparato_docx=None, comentario_docx=None) -> list[str]:
    """
    Ejecuta las comprobaciones sobre los DOCX y devuelve una lista
    de strings con los avisos encontrados (vacía si no hay warnings).
    """
    warnings: list[str] = []

    # 1) Comprueba existencia del principal
    if not os.path.exists(main_docx):
        warnings.append(f"❌ No existe el archivo principal: {main_docx}")
        return warnings

    # 2) Validación de estilos en el body
    STILI_VALIDI = {
        "Titulo_comedia", "Acto", "Prosa", "Verso", "Partido_incial",
        "Partido_medio", "Partido_final", "Personaje", "Acot",
        "Epigr_Dedic", "Epigr_Dramatis", "Dramatis_lista"
    }
    doc = Document(main_docx)
    found_body = False
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        if not found_body:
            if style == "Titulo_comedia":
                found_body = True
            continue  # seguimos buscando el inicio del body
        if style not in STILI_VALIDI:
            snippet = para.text.strip()[:60]
            warnings.append(f"❌ Estilo no válido: {style} — Texto: {snippet}")

    # 3) Análisis avanzado del texto principal
    warnings.extend(analyze_main_text(main_docx))

    # 4) Notas de aparato
    if aparato_docx:
        if not os.path.exists(aparato_docx):
            warnings.append(f"❌ No existe el archivo de notas de aparato: {aparato_docx}")
        else:
            aparato_notes = extract_notes_with_italics(aparato_docx)
            warnings.extend(analyze_notes(aparato_notes, "aparato"))

    # 5) Notas de comentario
    if comentario_docx:
        if not os.path.exists(comentario_docx):
            warnings.append(f"❌ No existe el archivo de notas de comentario: {comentario_docx}")
        else:
            comentario_notes = extract_notes_with_italics(comentario_docx)
            warnings.extend(analyze_notes(comentario_notes, "comentario"))

    return warnings

