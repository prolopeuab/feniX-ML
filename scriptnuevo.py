import os
import re
import unicodedata
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import filedialog, messagebox, font
from ttkthemes import ThemedTk
from docx import Document
from difflib import get_close_matches
import zipfile
from lxml import etree
from docx.oxml.ns import qn

def get_intro_footnotes(docx_path):
    """
    Estrae tutte le note a piè di pagina da un file DOCX.
    Ritorna un dizionario {id: testo_nota}.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    footnote_dict = {}

    with zipfile.ZipFile(docx_path) as docx_zip:
        with docx_zip.open("word/footnotes.xml") as footnote_file:
            root = etree.parse(footnote_file).getroot()

            for note in root.xpath("//w:footnote[not(@w:type='separator')]", namespaces=ns):
                note_id = note.get(qn("w:id"))
                texts = note.xpath(".//w:t", namespaces=ns)
                full_text = "".join(t.text for t in texts if t is not None).strip()
                if full_text:
                    footnote_dict[note_id] = full_text

    return footnote_dict

def extract_text_with_intro_notes(para, footnotes_intro):
    text = ""
    for run in para.runs:
        run_element = run._element
        refs = run_element.findall(".//w:footnoteReference", namespaces=run_element.nsmap)
        if refs:
            for ref in refs:
                note_id = ref.get(qn("w:id"))
                note_text = footnotes_intro.get(note_id, "")
                text += f'<note type="intro" n="{note_id}">{note_text}</note>'
        else:
            if run.italic:
                text += f'<hi rend="italic">{run.text}</hi>'
            else:
                text += run.text
    return text.strip()

def chiudi_blocchi_correnti(tei, stato):
    if stato.get("in_sp"):
        tei.append('        </sp>')
        stato["in_sp"] = False
    if stato.get("in_cast_list"):
        tei.append('        </castList>')
        tei.append('      </div>')
        stato["in_cast_list"] = False
    if stato.get("in_dedicatoria"):
        tei.append('      </div>')
        stato["in_dedicatoria"] = False
    if stato.get("in_act"):
        tei.append('      </div>')
        stato["in_act"] = False
comentario_notes = {}
aparato_notes = {}

#############################
# 1) FUNZIONI DI SUPPORTO
#############################

def extract_text_with_italics(para):
    """Estrae il testo da un paragrafo mantenendo il corsivo."""
    text = ""
    for run in para.runs:
        if run.italic:
            text += f'<hi rend="italic">{run.text}</hi>'
        else:
            text += run.text
    return text.strip()

def merge_italic_text(text):
    """
    Unisce le sequenze consecutive di tag <hi rend="italic">...</hi>
    in un unico tag.
    """
    # Cerca una o più occorrenze consecutive di <hi rend="italic">...</hi>
    pattern = re.compile(r'((?:<hi rend="italic">.*?</hi>\s*)+)', re.DOTALL)
    def replacer(match):
        segment = match.group(1)
        # Estrae tutto il contenuto interno dai tag
        inner_texts = re.findall(r'<hi rend="italic">(.*?)</hi>', segment, re.DOTALL)
        # Unisce tutto in un'unica stringa (senza spazi aggiuntivi)
        merged = ''.join(inner_texts).strip()
        return f'<hi rend="italic">{merged}</hi>'
    return pattern.sub(replacer, text)

def generate_filename(title):
    """
    Genera un nome file basato sulle prime tre parole del titolo
    e rimuovendo spazi, virgole e punti.
    """
    words = title.split()[:3]
    filename = '_'.join(words).replace(' ', '_').replace(',', '').replace('.', '')
    return filename

def find_who_id(speaker, personaggi):
    """
    Trova il ruolo corretto (xml:id) nel cast list con match flessibile.
    """
    speaker_cleaned = re.sub(r'[\s\[\]]+', '_', speaker).strip()

    # Match esatto
    if speaker.strip() in personaggi:
        return personaggi[speaker.strip()]

    # Match parziale
    for name, role_id in personaggi.items():
        if speaker.strip() in name or speaker_cleaned in role_id:
            return role_id

    # Fuzzy matching come ultima risorsa
    close_matches = get_close_matches(speaker.strip(), personaggi.keys(), n=1, cutoff=0.8)
    if close_matches:
        return personaggi[close_matches[0]]

    return ""

def parse_metadata_docx(path):
    """Estrae metadati da un file .docx strutturato in tabelle e costruisce un teiHeader TEI/XML completo."""
    doc = Document(path)
    tables = doc.tables

    if len(tables) < 3:
        raise ValueError("❌ Il documento dei metadati deve contenere almeno 3 tabelle.")

    # Tabella 1: Metadati principali
    main_meta = {}
    for row in tables[0].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            main_meta[key] = value

    # Tabella 2: sourceDesc
    source_meta = {}
    for row in tables[1].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            source_meta[key] = value

    # Tabella 3: listWit (salta la prima riga!)
    witnesses = []
    for i, row in enumerate(tables[2].rows):
        if i == 0:
            continue  # salta la riga guida "SIGLA TESTIMONIO | DESCRIPCIÓN"
        if len(row.cells) >= 2:
            siglum = row.cells[0].text.strip()

            # Estrai il contenuto della seconda cella mantenendo il corsivo
            desc_parts = []
            for run in row.cells[1].paragraphs[0].runs:
                if run.italic:
                    desc_parts.append(f'<hi rend="italic">{run.text}</hi>')
                else:
                    desc_parts.append(run.text)
            desc = ''.join(desc_parts).strip()

            if siglum and desc:
                witnesses.append((siglum, desc))

    # Costruzione teiHeader
    tei = ['<teiHeader>', '  <fileDesc>', '    <titleStmt>']

    tei.append(f'      <title>{main_meta.get("Titulo comedia", "")}</title>')
    tei.append(f'      <author><name>{main_meta.get("Autor", "")}</name></author>')

    if 'Editor' in main_meta:
        tei.append(f'      <editor>{main_meta["Editor"]}</editor>')

    if 'Responsable/s revisión' in main_meta:
        tei.append('      <respStmt>')
        tei.append('        <resp>Edición crítica digital revisada filológicamente por</resp>')
        for name in main_meta['Responsable/s revisión'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    if 'Responsable marcado automático' in main_meta:
        tei.append('      <respStmt>')
        tei.append('        <resp>Marcado XML-TEI automático revisado por</resp>')
        for name in main_meta['Responsable marcado automático'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    tei.append('      <respStmt>')
    tei.append('        <resp>Codificado según los criterios de</resp>')
    tei.append('        <name ref="https://datos.bne.es/entidad/XX4849774.html">Grupo de investigación PROLOPE, de la Universitat Autònoma de Barcelona</name>')
    tei.append('      </respStmt>')

    tei.extend(['    </titleStmt>', '    <editionStmt>'])
    tei.append(f'      <edition>Versión {main_meta.get("Versión", "")}</edition>')
    tei.extend(['    </editionStmt>', '    <publicationStmt>'])
    tei.append(f'      <publisher>{main_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'      <pubPlace>{main_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'      <date>{main_meta.get("Fecha publicación", "")}</date>')
    tei.extend(['    </publicationStmt>', '    <seriesStmt>'])
    tei.append('      <title>Biblioteca Digital PROLOPE</title>')
    tei.append('      <respStmt>')
    tei.append('        <resp>Dirección de</resp>')
    tei.append('        <persName ref="https://orcid.org/0000-0002-7429-9709"><forename>Ramón</forename> <surname>Valdés Gázquez</surname></persName>')
    tei.append('      </respStmt>')
    tei.append('      <idno type="URI">https://bibdigitalprolope.com/</idno>')
    tei.append('    </seriesStmt>')

    # sourceDesc
    tei.extend(['    <sourceDesc>', '      <biblStruct xml:lang="es">', '        <monogr>'])
    tei.append('          <author>')
    tei.append('            <persName ref="http://datos.bne.es/persona/XX1719671"><forename>Félix Lope</forename><surname>de Vega Carpio</surname></persName>')
    tei.append('          </author>')
    tei.append(f'          <title type="main">{source_meta.get("Titulo comedia", "")}</title>')
    if "Subtítulo" in source_meta:
        tei.append(f'          <title type="alt">{source_meta.get("Subtítulo", "")}</title>')
    tei.append(f'          <title type="s">{source_meta.get("Título volumen", "")}</title>')
    tei.append(f'          <title type="a">Parte {source_meta.get("Parte", "")}</title>')
    tei.append('          <availability status="restricted">')
    tei.append('           <p>Todos los derechos reservados.</p>')
    tei.append('          </availability>')
    tei.append('          <imprint>')
    tei.append(f'            <pubPlace>{source_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'            <publisher>{source_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'            <date>{source_meta.get("Fecha publicación", "")}</date>')
    tei.append(f'            <biblScope unit="volume" n="{source_meta.get("Volumen", "")}">vol. {source_meta.get("Volumen", "")}</biblScope>')
    tei.append(f'            <biblScope unit="page">{source_meta.get("Páginas", "")}</biblScope>')
    tei.append('          </imprint>')
    tei.append('        </monogr>')
    tei.append('      </biblStruct>')

    # listWit
    tei.append('      <listWit>')
    for siglum, desc in witnesses:
        tei.append(f'        <witness xml:id="{siglum}">')
        tei.append(f'          <label>{desc}</label>')
        tei.append('        </witness>')
    tei.append('      </listWit>')

    tei.extend(['    </sourceDesc>', '  </fileDesc>', '</teiHeader>'])

    return "\n".join(tei)

def process_front_paragraphs(paragraphs, footnotes_intro):
    tei_front = []
    subsection_open = False
    subsection_n = 1
    current_section = None
    paragraph_buffer = []
    head_inserted = False

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer
        for p in paragraph_buffer:
            text = extract_text_with_intro_notes(p, footnotes_intro)
            if p.style and p.style.name == "Cita":
                tei_front.append(f'          <cit rend="blockquote">')
                tei_front.append(f'            <quote>{text}</quote>')
                tei_front.append(f'          </cit>')
            elif text.strip():
                tei_front.append(f'          <p>{text.strip()}</p>')
        paragraph_buffer.clear()

    for i, para in enumerate(paragraphs):
        raw = extract_text_with_intro_notes(para, footnotes_intro)
        text = para.text.strip() if para.text else ""
        if not text:
            continue

        # 🔹 Ignora "Introducción"
        if text.lower() == "introducción":
            continue

        # 🔹 Gestione titolo principale "PRÓLOGO"
        if not head_inserted and "prólogo" in text.lower():
            flush_paragraph_buffer()
            tei_front.append(f'        <head type="divTitle" subtype="MenuLevel_1">PRÓLOGO</head>')
            head_inserted = True
            continue

        # 🔹 Riconoscimento sottotitolo con asterisco
        if text.startswith("*"):
            flush_paragraph_buffer()
            title = text.lstrip("*").strip()
            if title.lower() == "prólogo":
                continue  # ignora duplicazione
            if subsection_open:
                tei_front.append('        </div>')
            tei_front.append(f'        <div type="subsection" n="{subsection_n}">')
            tei_front.append(f'          <head type="divTitle" subtype="MenuLevel_2">{title}</head>')
            subsection_open = True
            current_section = title.lower()
            subsection_n += 1

        # 🔹 Tabelle nella sezione "Sinopsis"
            if current_section and "sinopsis" in current_section:
                parent = para._parent
                if hasattr(parent, "tables") and parent.tables:
                    for tbl in parent.tables:
                        tei_front.append(process_table_to_tei(tbl))
            continue

        # 🔹 Aggiungi al buffer
        paragraph_buffer.append(para)

    flush_paragraph_buffer()

    if subsection_open:
        tei_front.append('        </div>')  # chiude ultimo subsection

    return "\n".join(tei_front)

def process_table_to_tei(table):
    tei = ['          <table rend="rules">']
    
    for row in table.rows:
        cells = row.cells
        texts = [cell.text.strip() for cell in cells]
        non_empty_cells = [t for t in texts if t]

        # Riga di intestazione di sezione (es. Acto primero, Resumen)
        if len(non_empty_cells) == 1 and texts[0]:
            tei.append('            <row>')
            tei.append(f'              <cell rend="both" cols="3"><hi rend="italic">{texts[0]}</hi></cell>')
            tei.append('            </row>')
            continue

        # Riga standard a 3 colonne
        tei.append('            <row>')
        for txt in texts:
            tei.append(f'              <cell rend="both">{txt}</cell>')
        tei.append('            </row>')

    tei.append('          </table>')
    return "\n".join(tei)

#############################
# 2) LETTURA NOTE (COMENTARIO/APARATO)
#############################

def extract_notes_with_italics(docx_path):
    notes = {}
    if not docx_path or not os.path.exists(docx_path):
        print(f"❌ Il file note '{docx_path}' non esiste o non è stato selezionato.")
        return notes
    print(f"📂 Estrazione note da: {docx_path}")
    doc = Document(docx_path)
        
    for para in doc.paragraphs:
        text = extract_text_with_italics(para).strip()

        print(f"   ➜ Paragrafo note: '{text[:50]}'")

        match_verse = re.match(r'^(\d+):\s*(.*)', text)
        match_single_word = re.match(r'^@([^@]+?):\s*(.*)', text)

        if match_verse:
            verse_number = int(match_verse.group(1))
            note_content = match_verse.group(2).strip()
            notes[verse_number] = note_content
            print(f"✔️ Nota VERSO {verse_number}: {note_content}")

        elif match_single_word:
            word = match_single_word.group(1).strip()
            note_content = match_single_word.group(2).strip()
            if word not in notes:
                notes[word] = note_content
                print(f"✔️ Nota PAROLA '{word}': {note_content}")

    return notes


#############################
# 4) PROCESSAMENTO DELLE ANNOTAZIONI @
#############################

def process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, section):
    if not text:
        print(f"❌ [DEBUG] Il testo è None nella sezione '{section}'")
        return ""

    print(f"🔍 [DEBUG] Analizzando testo in sezione '{section}': {text[:100]}...")

    matches = re.findall(r'@(\w+)', text)  # Cerca solo @parola

    if matches is None:
        print("⚠️ [DEBUG] Attenzione: `matches` è None")
        return text

    print(f"🟢 [DEBUG] Annotazioni trovate in '{section}': {matches}")

    comentario_notes = comentario_notes or {}
    aparato_notes = aparato_notes or {}

    def normalize_word(word):
        if isinstance(word, int):
            return word
        word = re.sub(r'<hi rend="italic">(.*?)</hi>', r'\1', word)
        word = unicodedata.normalize('NFKD', word).encode('ASCII', 'ignore').decode('utf-8').lower()
        return word.strip()

    new_text = text.strip()

    comentario_notes_normalized = {normalize_word(k): v for k, v in comentario_notes.items() if isinstance(k, str)}
    aparato_notes_normalized = {normalize_word(k): v for k, v in aparato_notes.items() if isinstance(k, str)}

    all_notes = {**comentario_notes_normalized, **aparato_notes_normalized}

    for phrase in matches:
        if not phrase:
            continue

        phrase_to_replace = f"@{phrase}"
        normalized_key = normalize_word(phrase.strip())

        print(f"🔍 [DEBUG] Controllo annotazione: '{phrase_to_replace}', Normalized: '{normalized_key}'")

        note = ""

        if normalized_key in all_notes:
            annotation_counter_section = annotation_counter.get(section, {})
            annotation_counter_section[normalized_key] = annotation_counter_section.get(normalized_key, 0) + 1

            xml_id = f"{normalized_key}_{section}_{annotation_counter_section[normalized_key]}"
            xml_id = re.sub(r'\s+', '_', xml_id)
            xml_id = re.sub(r'[^a-zA-Z0-9_]', '', xml_id)
            xml_id = xml_id.lower()

            note_comentario = comentario_notes_normalized.get(normalized_key, "")
            note_aparato = aparato_notes_normalized.get(normalized_key, "")

            if note_comentario:
                note += f'<note subtype="comentario" xml:id="{xml_id}">{note_comentario}</note>'
            if note_aparato:
                note += f'<note subtype="aparato" xml:id="{xml_id}">{note_aparato}</note>'

            if new_text and phrase_to_replace in new_text:
                print(f"✅ [DEBUG] Sostituzione in corso per '{phrase_to_replace}'...")
                new_text = new_text.replace(phrase_to_replace, f"{phrase}{note}", 1)
                print(f"✔️ [DEBUG] Sostituzione completata per '{phrase_to_replace}' → '{phrase}{note}'")
            else:
                print(f"⚠️ [DEBUG] '{phrase_to_replace}' NON trovato nel testo!")

        else:
            print(f"❌ [DEBUG] Nessuna nota trovata per '{normalized_key}'!")

        print(f"🔄 [DEBUG] Tentativo di sostituzione: '{phrase_to_replace}' con '<note>{note}</note>'")

    new_text = merge_italic_text(new_text)

    return new_text if new_text else ""

#############################
# 5) FUNZIONE PRINCIPALE
#############################

def convert_docx_to_tei(main_docx, comentario_docx=None, aparato_docx=None, output_file=None, tei_header=None):
    if not os.path.exists(main_docx):
        raise FileNotFoundError(f"Il file principale '{main_docx}' non esiste!")
    
    tei = []
    
    if tei_header:
        tei.append(tei_header)
    else:
        tei.append("<teiHeader>...</teiHeader>")  # fallback
    
    tei.append("<text>")
    
    doc = Document(main_docx)
    print(f"📂 Leggo il documento principale: {main_docx}")

    # --- SEPARAZIONE PARAGRAFI FRONT/BODY BASATA SULLO STILE 'Titulo_comedia' ---

    front_paragraphs = []
    body_paragraphs = []
    inside_body = False

    for para in doc.paragraphs:
        text = extract_text_with_italics(para).strip()
        style = para.style.name if para.style else "Normal"

        if not inside_body:
            if style == 'Titulo_comedia':
                inside_body = True
                body_paragraphs.append(para)  # Il paragrafo del titolo va nel body
            else:
                front_paragraphs.append(para)
        else:
            body_paragraphs.append(para)


    # Estrazione del titolo
    title = "Untitled"
    for para in doc.paragraphs:
        if para.style and para.style.name == "Titulo_comedia":
            title = para.text.strip()
            break

    clean_title = re.sub(r'@', '', title)
    title_key = generate_filename(title)

    # Se l'utente ha fornito i file di commento e apparato, usali direttamente;
    # altrimenti, usa quelli ricostruiti.
    if comentario_docx and os.path.exists(comentario_docx):
        comentario_file = comentario_docx
    else:
        comentario_file = f'comentario_{title_key}.docx'

    if aparato_docx and os.path.exists(aparato_docx):
        aparato_file = aparato_docx
    else:
        aparato_file = f'aparato_{title_key}.docx'

    print(f"📝 File comentario: {comentario_file} → {os.path.exists(comentario_file)}")
    print(f"📝 File aparato: {aparato_file} → {os.path.exists(aparato_file)}")

    comentario_notes = extract_notes_with_italics(comentario_file) if os.path.exists(comentario_file) else {}
    aparato_notes = extract_notes_with_italics(aparato_file) if os.path.exists(aparato_file) else {}

    print("NOTES COMMENTARIO:", comentario_notes)
    print("NOTES APARATO:", aparato_notes)

    # ... (il resto della funzione per generare l'XML TEI)
    
    annotation_counter = {}
    stato = {"in_sp": False, "in_cast_list": False, "in_dedicatoria": False, "in_act": False}
    personaggi = {}
    act_counter = 0
    verse_counter = 1
    processed_title = process_annotations_with_ids(title, comentario_notes, aparato_notes, {}, "head")
    footnotes_intro = get_intro_footnotes(main_docx)  # Usa il percorso del docx principale
    ultimo_speaker_id = None

    # Intestazione TEI
    tei = [
    '<?xml-model href="http://www.tei-c.org/release/xml/tei/custom/schema/relaxng/tei_all.rng" schematypens="http://relaxng.org/ns/structure/1.0"?>',
    '<TEI xmlns="http://www.tei-c.org/ns/1.0">',
    tei_header,
    '  <text>',
    '    <front>',
    '      <div type="Introducción">'
    ]

    tei.append(process_front_paragraphs(front_paragraphs, footnotes_intro))

    tei.extend([
        '      </div>',  # chiude Introducción
        '    </front>',
        '    <body>',
        '      <div type="Texto" subtype="TEXTO">',
        f'        <head type="mainTitle">{processed_title}</head>',
    ])

    # Scansione paragrafi
    for para in body_paragraphs:
        text = extract_text_with_italics(para).strip()
        style = para.style.name if para.style else "Normal"

        # Rilevamento strofe $nomeStrofa -> milestone
        lg_milestone = re.match(r'^\$(\w+)', text)
        if lg_milestone:
            current_milestone = lg_milestone.group(1)
            continue

        # Chiudi blocchi prima di aprirne altri
        if style in ["Epigr_Dramatis", "Acto", "Epigr_Dedic"]:
            chiudi_blocchi_correnti(tei, stato)

        if style == "Epigr_Dedic":
            tei.append('        <div type="dedicatoria">')
            tei.append(f'          <head>{text}</head>')
            stato["in_dedicatoria"] = True

        elif style == "Epigr_Dramatis":
            tei.append('        <div type="castList">')
            tei.append('          <castList>')
            tei.append(f'            <head>{text}</head>')
            stato["in_cast_list"] = True

        elif style == "Dramatis_lista":
            role_name = text
            if role_name:
                role_id = re.sub(r'[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ_-]+', '_', role_name)
                tei.append(f'            <castItem><role xml:id="{role_id}">{role_name}</role></castItem>')
                personaggi[role_name] = role_id

        elif style == "Acto":
            act_counter += 1
            tei.append(f'        <div type="subsection" subtype="ACTO" n="{act_counter}">')
            tei.append(f'          <head type="acto">{text}</head>')
            stato["in_act"] = True

        elif style == "Prosa":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "p")
            if processed_text.strip():
                tei.append(f'          <p>{processed_text}</p>')

        elif style == "Verso":
            if stato["in_dedicatoria"]:
                tei.append(f'          <l>{text}</l>')
            elif stato["in_sp"]:
                if current_milestone:
                    tei.append(f'            <milestone unit="stanza" type="{current_milestone}"/>')
                    current_milestone = None
                verse_text = text
                if verse_counter in comentario_notes:
                    verse_text += f'<note subtype="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
                if verse_counter in aparato_notes:
                    verse_text += f'<note subtype="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
                tei.append(f'            <l n="{verse_counter}">{verse_text}</l>')
                verse_counter += 1

        elif style == "Partido_incial":
            verse_text = text
            if verse_counter in comentario_notes:
                verse_text += f'<note subtype="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
            if verse_counter in aparato_notes:
                verse_text += f'<note subtype="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
            tei.append(f'            <l part="I" n="{verse_counter}">{verse_text}</l>')
            verse_counter += 1

        elif style == "Partido_medio":
            tei.append(f'            <l part="M">{text}</l>')

        elif style == "Partido_final":
            tei.append(f'            <l part="F">{text}</l>')

        elif style == "Acot":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "stage")
            if stato["in_sp"]:
                tei.append('        </sp>')
                stato["in_sp"] = False
            tei.append(f'        <stage>{processed_text}</stage>')

        elif style == "Personaje":
            who_id = find_who_id(text, personaggi)
            processed = process_annotations_with_ids(
                text, comentario_notes, aparato_notes, annotation_counter, "speaker"
            )

            # 🔒 Chiudi <sp> precedente se necessario
            if stato["in_sp"]:
                tei.append('        </sp>')
                stato["in_sp"] = False

            # 🎯 Se è lo stesso personaggio precedente, non inserire who o speaker
            if who_id == ultimo_speaker_id:
                tei.append('        <sp>')
            else:
                tei.append(f'        <sp who="#{who_id}">')
                tei.append(f'          <speaker>{processed}</speaker>')
                ultimo_speaker_id = who_id

            stato["in_sp"] = True

        elif style == "Titulo_comedia":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "head")

    # Chiusura finale di tutti i blocchi ancora aperti
    chiudi_blocchi_correnti(tei, stato)

    # Chiusura sezioni TEI
    tei.append('      </div>')  # chiude Texto
    tei.append('    </body>')
    tei.append('  </text>')
    tei.append('</TEI>')


    # Se l'utente non passa output_file, generiamo un nome
    if not output_file:
        output_file = generate_filename(title) + ".xml"

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("\n".join(tei))
    print(f"✅ File XML salvato come: {output_file}")

def analyze_main_text(main_docx):
    """Analizza il file principale e cerca paragrafi senza stile"""
    print(f"📂 Analizzando file: {main_docx}")  # DEBUG: Assicuriamoci che la funzione venga chiamata
    
    main_doc = Document(main_docx)
    unstyled_paragraphs = []

    for para in main_doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else "None"

        print(f'[DEBUG] Testo: "{text[:50]}" | Stile: {style}')  # Mostra i primi 50 caratteri e lo stile

        # Ignora paragrafi vuoti
        if not text:
            continue  

        # Se il testo non ha stile e non è un milestone, lo segnaliamo
        if style in ["Normal", "None"]:
            if re.match(r'^\$\w+', text):  # Ignora milestone ($strofa)
                continue  
            if re.match(r'^[.,;:!?()"\']+$', text):  # Ignora solo punteggiatura
                continue  

            unstyled_paragraphs.append(text)

    if unstyled_paragraphs:
        messagebox.showwarning(
            "¡Cuidado! (Texto Crítico)",
            f"¡He encontrado {len(unstyled_paragraphs)} párrafos sin estilos en el texto crítico!\n\nRevisa el DOCX."
        )

def analyze_notes(doc_path, note_type):
    """Analizza il file delle note (di apparato o di commento) e verifica problemi di formattazione"""
    doc = Document(doc_path)
    incorrect_notes = []

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Ignora paragrafi vuoti
        if not text:
            continue

        # Controlliamo se il paragrafo inizia con un numero (es. 1:) o con @parola:
        if not re.match(r'^(\d+:\s+.+|@[^@]+@:\s+.+|@\S+:\s+.+)', text):
            incorrect_notes.append(text)

    if incorrect_notes:
        messagebox.showwarning(
            f"¡Cuidado! (Notas de {note_type})",
            f"¡Se han encontrado {len(incorrect_notes)} posibles notas sin formato correcto en las notas de {note_type}!\n\nRevisa el DOCX."
        )


from docx import Document

def validate_documents(main_docx, aparato_docx=None, comentario_docx=None):
    """Esegue le analisi su tutti i documenti DOCX e avvisa l'utente di eventuali errori."""
    if not os.path.exists(main_docx):
        messagebox.showerror("Error", f"El archivo DOCX principal no existe:\n{main_docx}")
        return

    doc = Document(main_docx)
    segnalazioni = []
    STILI_VALIDI = {
        "Titulo_comedia", "Acto", "Prosa", "Verso", "Partido_incial", "Partido_medio",
        "Partido_final", "Personaje", "Acot", "Epigr_Dedic", "Epigr_Dramatis", "Dramatis_lista"
    }

    found_body = False
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""

        if not found_body:
            if style == "Titulo_comedia":
                found_body = True
            continue  # Ignora l'introduzione

        if style not in STILI_VALIDI:
            segnalazioni.append(f"❌ Stile non valido: {style} — Testo: {para.text.strip()[:60]}")

    # Analisi avanzata
    analyze_main_text(main_docx)

    if aparato_docx:
        if not os.path.exists(aparato_docx):
            messagebox.showerror("Error", f"El archivo de notas de aparato no existe:\n{aparato_docx}")
        else:
            analyze_notes(aparato_docx, "Aparato")

    if comentario_docx:
        if not os.path.exists(comentario_docx):
            messagebox.showerror("Error", f"El archivo de notas de comentario no existe:\n{comentario_docx}")
        else:
            analyze_notes(comentario_docx, "Comentario")

    if segnalazioni:
        raise Exception("\n".join(segnalazioni))

def avvia_validazione():
    """Esegue la validazione dei documenti senza avviare la conversione."""
    main_docx = entry_main.get().strip()

    aparato_docx = entry_apa.get().strip() if entry_apa.get().strip() else None
    comentario_docx = entry_com.get().strip() if entry_com.get().strip() else None

    if not main_docx:
        messagebox.showerror("Error", "Seleccionar al menos el archivo DOCX principal.")
        return

    try:
        validate_documents(main_docx, aparato_docx, comentario_docx)
        messagebox.showinfo("Validación completada", "La validación de los archivos se ha completado correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"Se ha verificado un error en la validación:\n{e}")
    
def avvia_conversione():
    main_file = entry_main.get().strip()
    com_file = entry_com.get().strip()
    apa_file = entry_apa.get().strip()
    out_file = entry_out.get().strip()
    metadata_path = entry_meta.get().strip()

    if metadata_path:
        tei_header = parse_metadata_docx(metadata_path)
    else:
        tei_header = "<teiHeader>...</teiHeader>"  # fallback

    if not main_file:
        messagebox.showerror("Error", "Seleccione al menos el DOCX Principal!")
        return

    try:
        # Se non è stato fornito un nome, usiamo le prime tre parole del titolo
        if not out_file:
            doc = Document(main_file)
            title = "Untitled"
            for para in doc.paragraphs:
                if para.style and para.style.name == "Titulo_comedia":
                    title = para.text.strip()
                    break
            out_file = generate_filename(title) + ".xml"

        convert_docx_to_tei(
            main_docx=main_file,
            comentario_docx=com_file if com_file else None,
            aparato_docx=apa_file if apa_file else None,
            output_file=out_file,
            tei_header=tei_header
        )
        messagebox.showinfo(
            "Operación completada",
            f"Archivo XML guardado con éxito:\n{out_file}"
        )
    except Exception as e:
        messagebox.showerror("Error", f"Se ha producido un error:\n{e}")

def main_gui():
    global entry_main, entry_com, entry_apa, entry_out, entry_meta

    root = tk.Tk()
    root.title("DOCXtoTEI")
    root.configure(bg="#ACBDE9")
    root.geometry("700x450")

    style = ttk.Style()
    style.theme_use("vista")
    style.configure(".", font=("Open Sans MS", 10))
    style.configure("TFrame", background="#ACBDE9")
    style.configure("TLabel", background="#ACBDE9")
    style.configure("TButton", font=("Open Sans MS", 10))

    main_frame = ttk.Frame(root, padding="10 10 10 10", style="TFrame")
    main_frame.pack(fill="both", expand=True)

    ttk.Label(main_frame, text="Archivo DOCX Principal:").grid(row=0, column=0, sticky="e", pady=5)
    entry_main = ttk.Entry(main_frame, width=50)
    entry_main.grid(row=0, column=1, padx=5, sticky="ew")
    btn_main = ttk.Button(main_frame, text="Seleccionar", command=lambda: entry_main.insert(0, filedialog.askopenfilename()))
    btn_main.grid(row=0, column=2, padx=5)

    ttk.Label(main_frame, text="Notas de Comentario:").grid(row=1, column=0, sticky="e", pady=5)
    entry_com = ttk.Entry(main_frame, width=50)
    entry_com.grid(row=1, column=1, padx=5, sticky="ew")
    btn_com = ttk.Button(main_frame, text="Seleccionar", command=lambda: entry_com.insert(0, filedialog.askopenfilename()))
    btn_com.grid(row=1, column=2, padx=5)

    ttk.Label(main_frame, text="Notas de Aparato:").grid(row=2, column=0, sticky="e", pady=5)
    entry_apa = ttk.Entry(main_frame, width=50)
    entry_apa.grid(row=2, column=1, padx=5, sticky="ew")
    btn_apa = ttk.Button(main_frame, text="Seleccionar", command=lambda: entry_apa.insert(0, filedialog.askopenfilename()))
    btn_apa.grid(row=2, column=2, padx=5)

    ttk.Label(main_frame, text="Archivo de Metadatos:").grid(row=3, column=0, sticky="e", pady=5)
    entry_meta = ttk.Entry(main_frame, width=50)
    entry_meta.grid(row=3, column=1, padx=5, sticky="ew")
    btn_meta = ttk.Button(main_frame, text="Seleccionar", command=lambda: entry_meta.insert(0, filedialog.askopenfilename()))
    btn_meta.grid(row=3, column=2, padx=5)

    ttk.Label(main_frame, text="Nombre de archivo XML (opcional):").grid(row=4, column=0, sticky="e", pady=5)
    entry_out = ttk.Entry(main_frame, width=50)
    entry_out.grid(row=4, column=1, padx=5, sticky="ew")

    btn_validar = ttk.Button(main_frame, text="Validar Documentos", command=avvia_validazione)
    btn_validar.grid(row=5, column=1, pady=5)

    btn_convertir = ttk.Button(main_frame, text="Convertir a TEI", command=avvia_conversione)
    btn_convertir.grid(row=6, column=1, pady=10)

    main_frame.columnconfigure(1, weight=1)

    root.mainloop()

if __name__ == "__main__":
    main_gui()
