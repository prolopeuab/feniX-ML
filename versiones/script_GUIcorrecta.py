#pyinstaller --onefile --windowed --add-data "CETEIcean.js;." --add-data "estilos.css;." DOCXtoTEI.py
import json
import sys
import os
import re
import unicodedata
import tkinter as tk
import tkinter.ttk as ttk
import webbrowser
import tempfile
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from difflib import get_close_matches

#############################
# 1) FUNZIONI DI SUPPORTO
#############################
def resource_path(relative_path):
    """Obtiene la ruta absoluta del recurso, compatible con PyInstaller."""
    try:
        # Cuando se ejecuta el exe, PyInstaller define sys._MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def load_resource(filename):
    path = resource_path(filename)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

# Carga automática de los archivos externos
CETEI_JS = load_resource("CETEIcean.js")
ESTILOS_CSS = load_resource("estilos.css")

def extract_text_with_italics(para):
    """Estrae il testo da un paragrafo mantenendo il corsivo."""
    text = ""
    for run in para.runs:
        if run.italic:
            text += f'<hi rend="italic">{run.text}</hi>'
        else:
            text += run.text
    return text.strip()

def merge_italic_text(text):
    """
    Unisce le sequenze consecutive di tag <hi rend="italic">...</hi>
    in un unico tag.
    """
    # Cerca una o più occorrenze consecutive di <hi rend="italic">...</hi>
    pattern = re.compile(r'((?:<hi rend="italic">.*?</hi>\s*)+)', re.DOTALL)
    def replacer(match):
        segment = match.group(1)
        # Estrae tutto il contenuto interno dai tag
        inner_texts = re.findall(r'<hi rend="italic">(.*?)</hi>', segment, re.DOTALL)
        # Unisce tutto in un'unica stringa (senza spazi aggiuntivi)
        merged = ''.join(inner_texts).strip()
        return f'<hi rend="italic">{merged}</hi>'
    return pattern.sub(replacer, text)

def generate_filename(title):
    """
    Genera un nome file basato sulle prime tre parole del titolo
    e rimuovendo spazi, virgole e punti.
    """
    words = title.split()[:3]
    filename = '_'.join(words).replace(' ', '_').replace(',', '').replace('.', '')
    return filename

def find_who_id(speaker, personaggi):
    """
    Trova il ruolo corretto (xml:id) nel cast list con match flessibile.
    """
    speaker_cleaned = re.sub(r'[\s\[\]]+', '_', speaker).strip()

    # Match esatto
    if speaker.strip() in personaggi:
        return personaggi[speaker.strip()]

    # Match parziale
    for name, role_id in personaggi.items():
        if speaker.strip() in name or speaker_cleaned in role_id:
            return role_id

    # Fuzzy matching come ultima risorsa
    close_matches = get_close_matches(speaker.strip(), personaggi.keys(), n=1, cutoff=0.8)
    if close_matches:
        return personaggi[close_matches[0]]

    return ""

#############################
# 2) LETTURA NOTE (COMENTARIO/APARATO)
#############################

def extract_notes_with_italics(docx_path):
    notes = {}
    if not docx_path or not os.path.exists(docx_path):
        print(f"❌ Il file note '{docx_path}' non esiste o non è stato selezionato.")
        return notes
    print(f"📂 Estrazione note da: {docx_path}")
    doc = Document(docx_path)
        
    for para in doc.paragraphs:
        # Estrae testo con corsivo
        text = extract_text_with_italics(para)
        text = text.strip()

        # Debug paragrafi
        print(f"   ➜ Paragrafo note: '{text[:50]}'")

        # Regex
        match_verse = re.match(r'^(\d+):\s*(.*)', text)
        match_multi_word = re.match(r'^@(.+?)@:\s*(.*)', text)
        match_single_word = re.match(r'^@([^@]+?):\s*(.*)', text)

        if match_verse:
            verse_number = int(match_verse.group(1))
            note_content = match_verse.group(2).strip()
            notes[verse_number] = note_content
            print(f"✔️ Nota VERSO {verse_number}: {note_content}")

        elif match_multi_word:
            phrase = match_multi_word.group(1).strip()
            note_content = match_multi_word.group(2).strip()
            notes[phrase] = note_content
            print(f"✔️ Nota FRASE '{phrase}': {note_content}")

        elif match_single_word:
            word = match_single_word.group(1).strip()
            note_content = match_single_word.group(2).strip()
            if word not in notes:
                notes[word] = note_content
                print(f"✔️ Nota PAROLA '{word}': {note_content}")

    return notes

#############################
# 3) PROCESSAMENTO DELLE ANNOTAZIONI @
#############################

def process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, section):
    if not text:
        print(f"❌ [DEBUG] Il testo è None nella sezione '{section}'")
        return ""

    print(f"🔍 [DEBUG] Analizzando testo in sezione '{section}': {text[:100]}...")  

    matches = re.findall(r'@(\w+(?:\s\w+)*)@|@(\w+)(?=[^\w@]|$)', text)

    
    if matches is None:
        print("⚠️ [DEBUG] Attenzione: `matches` è None")
        return text  # Evita errori

    print(f"🟢 [DEBUG] Annotazioni trovate in '{section}': {matches}")  

    # 🔹 Controllo che comentario_notes e aparato_notes non siano None
    comentario_notes = comentario_notes or {}
    aparato_notes = aparato_notes or {}

    def normalize_word(word):
        if isinstance(word, int):
            return word
        word = re.sub(r'<hi rend="italic">(.*?)</hi>', r'\1', word)  
        word = unicodedata.normalize('NFKD', word).encode('ASCII', 'ignore').decode('utf-8').lower()
        return word.strip()

    new_text = text.strip()  # Evita problemi con None

    comentario_notes_normalized = {normalize_word(k): v for k, v in comentario_notes.items() if isinstance(k, str)}
    aparato_notes_normalized = {normalize_word(k): v for k, v in aparato_notes.items() if isinstance(k, str)}

    all_notes = {**comentario_notes_normalized, **aparato_notes_normalized}

    for match in matches:
        if not match or not isinstance(match, tuple) or len(match) < 2:
            print(f"⚠️ [DEBUG] Match non valido: {match}")
            continue  # Evita errore "NoneType is not subscriptable"

        phrase = match[0] if match[0] else match[1]
        if not phrase:
            continue  

        phrase_to_replace = f"@{phrase}@" if match[0] else f"@{phrase}"
        normalized_key = normalize_word(phrase.strip())

        print(f"🔍 [DEBUG] Controllo annotazione: '{phrase_to_replace}', Normalized: '{normalized_key}'")

        # 🔹 Inizializzazione di `note`
        note = ""

        if normalized_key in all_notes:
            annotation_counter_section = annotation_counter.get(section, {})
            annotation_counter_section[normalized_key] = annotation_counter_section.get(normalized_key, 0) + 1

            xml_id = f"{normalized_key}_{section}_{annotation_counter_section[normalized_key]}"
    
        # 🔹 Normalizzazione dell'xml:id per evitare errori
            xml_id = re.sub(r'\s+', '_', xml_id)  # Sostituisce spazi con underscore
            xml_id = re.sub(r'[^a-zA-Z0-9_]', '', xml_id)  # Rimuove caratteri speciali non permessi
            xml_id = xml_id.lower()  # Converte tutto in minuscolo

            note_comentario = comentario_notes_normalized.get(normalized_key, "")
            note_aparato = aparato_notes_normalized.get(normalized_key, "")

            if note_comentario:
                note += f'<note type="comentario" xml:id="{xml_id}">{note_comentario}</note>'
            if note_aparato:
                note += f'<note type="aparato" xml:id="{xml_id}">{note_aparato}</note>'

            # 🔹 Controlliamo se la sostituzione è possibile
            if new_text and phrase_to_replace in new_text:
                print(f"✅ [DEBUG] Sostituzione in corso per '{phrase_to_replace}'...")
                new_text = new_text.replace(phrase_to_replace, f"{phrase}{note}", 1)
                print(f"✔️ [DEBUG] Sostituzione completata per '{phrase_to_replace}' → '{phrase}{note}'")
            else:
                print(f"⚠️ [DEBUG] '{phrase_to_replace}' NON trovato nel testo!")

        else:
            print(f"❌ [DEBUG] Nessuna nota trovata per '{normalized_key}'!")

        print(f"🔄 [DEBUG] Tentativo di sostituzione: '{phrase_to_replace}' con '<note>{note}</note>'")

        # 🔹 **Unire i segmenti in corsivo spezzati prima del ritorno finale**
    new_text = merge_italic_text(new_text)

    return new_text if new_text else ""

#############################
# 4) FUNZIONE PRINCIPALE
#############################

def convert_docx_to_tei(main_docx, comentario_docx=None, aparato_docx=None, output_file=None, save=True):
    if not os.path.exists(main_docx):
        raise FileNotFoundError(f"Il file principale '{main_docx}' non esiste!")
    
    doc = Document(main_docx)
    print(f"📂 Leggo il documento principale: {main_docx}")

    # Estrazione del titolo
    title = "Untitled"
    for para in doc.paragraphs:
        if para.style and para.style.name == "Titulo_comedia":
            title = para.text.strip()
            break

    clean_title = re.sub(r'@', '', title)
    title_key = generate_filename(title)

    # Se l'utente ha fornito i file di commento e apparato, usali direttamente;
    # altrimenti, usa quelli ricostruiti.
    if comentario_docx and os.path.exists(comentario_docx):
        comentario_file = comentario_docx
    else:
        comentario_file = f'comentario_{title_key}.docx'

    if aparato_docx and os.path.exists(aparato_docx):
        aparato_file = aparato_docx
    else:
        aparato_file = f'aparato_{title_key}.docx'

    print(f"📝 File comentario: {comentario_file} → {os.path.exists(comentario_file)}")
    print(f"📝 File aparato: {aparato_file} → {os.path.exists(aparato_file)}")

    comentario_notes = extract_notes_with_italics(comentario_file) if os.path.exists(comentario_file) else {}
    aparato_notes = extract_notes_with_italics(aparato_file) if os.path.exists(aparato_file) else {}

    print("NOTES COMMENTARIO:", comentario_notes)
    print("NOTES APARATO:", aparato_notes)

    # ... (il resto della funzione per generare l'XML TEI)
    
    annotation_counter = {}
    in_cast_list = False
    in_sp = False
    in_dedicatoria = False
    in_act = False
    personaggi = {}
    act_counter = 0
    verse_counter = 1

    # Intestazione TEI
    tei = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<TEI xmlns:math="http://www.w3.org/1998/Math/MathML" xmlns:xi="http://www.w3.org/2001/XInclude" xmlns="http://www.tei-c.org/ns/1.0" xmlns:svg="http://www.w3.org/2000/svg">',
        '  <teiHeader>',
        '    <fileDesc>',
         '      <titleStmt>',
        f'        <title>{clean_title}</title>',
        '       </titleStmt>',
        '       <publicationStmt>',
        '          <p/>',
        '        </publicationStmt>',
        '       <sourceDesc>',
        '       <bibl/>',
        '     </sourceDesc>',
        '    </fileDesc>',
        '  </teiHeader>',
        '  <text>',
        '    <body>',
        f'      <head>{process_annotations_with_ids(title, comentario_notes, aparato_notes, {}, "head")}</head>'
    ]

    current_milestone = None

    # Scansione paragrafi
    for para in doc.paragraphs:
        text = extract_text_with_italics(para)
        text = text.strip()
        style = para.style.name if para.style else "Normal"

        # Rilevamento strofe $nomeStrofa -> milestone
        lg_milestone = re.match(r'^\$(\w+)', text)
        if lg_milestone:
            current_milestone = lg_milestone.group(1)
            print(f"📍 Milestone rilevata: {current_milestone}")
            continue

        if style == "Epigr_Dramatis":
            if in_dedicatoria:
                tei.append('      </div>')
                in_dedicatoria = False
            if not in_cast_list:
                tei.append('      <div type="castList">')
                tei.append('        <castList>')
                in_cast_list = True  # Segnala che siamo dentro la lista

            # Aggiungere il titolo dentro <head> con il contenuto di Epigr_Dramatis
            if text.strip():
                tei.append(f'        <head>{text.strip()}</head>')

        elif style == "Dramatis_lista":
            role_id = re.sub(r'[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ_-]+', '_', text.strip())

            personaggi[text.strip()] = role_id
            tei.append(f'  <castItem><role xml:id="{role_id}">{text}</role></castItem>')
           

        elif style == "Epigr_Dedic":
            if in_dedicatoria:
                tei.append('      </div>')
            tei.append('      <div type="dedicatoria">')
            tei.append(f'        <head>{text}</head>')
            in_dedicatoria = True

        elif style == "Prosa":
            print(f"🔎 [DEBUG] Testo finale per TEI: {processed_text[:300]}...")

            processed_text = ""  # Inizializzazione sicura
            if text:  # Controlla che il testo non sia None o vuoto
                processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "p")
    
            if processed_text.strip():
                tei.append(f'          <p>{processed_text}</p>')

        elif style == "Verso":
            if in_dedicatoria:
                tei.append(f'        <l>{text}</l>')
            elif in_sp:
                if current_milestone:
                    tei.append(f'            <milestone unit="stanza" type="{current_milestone}"/>')
                    current_milestone = None
                verse_text = text
                if verse_counter in comentario_notes:
                    verse_text += f'<note type="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
                if verse_counter in aparato_notes:
                    verse_text += f'<note type="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
                tei.append(f'            <l n="{verse_counter}">{verse_text}</l>')
                verse_counter += 1

        elif style == "Partido_incial":
            verse_text = text
            if verse_counter in comentario_notes:
                verse_text += f'<note type="comentario" n="{verse_counter}">{comentario_notes[verse_counter]}</note>'
            if verse_counter in aparato_notes:
                verse_text += f'<note type="aparato" n="{verse_counter}">{aparato_notes[verse_counter]}</note>'
            tei.append(f'            <l part="I" n="{verse_counter}">{verse_text}</l>')
            verse_counter += 1

        elif style == "Partido_medio":
            tei.append(f'            <l part="M">{text}</l>')

        elif style == "Partido_final":
            tei.append(f'            <l part="F">{text}</l>')

        elif style == "Acto":
            if in_sp:
                tei.append('        </sp>')
                in_sp = False
            if in_cast_list:
                tei.append('        </castList>')
                tei.append('      </div>')
                in_cast_list = False
            if in_dedicatoria:
                tei.append('      </div>')
                in_dedicatoria = False
            if in_act:
                tei.append('      </div>')
                in_act = False
            act_counter += 1
            tei.append(f'      <div type="act" n="{act_counter}">')
            tei.append(f'        <head>{text}</head>')
            in_act = True

        elif style == "Acot":
            processed_text = text if text else ""  # Assicura che processed_text sia sempre inizializzato
            processed_text = process_annotations_with_ids(
                text, comentario_notes, aparato_notes, annotation_counter, "stage"
            )
            if in_sp:
                tei.append('        </sp>')
                in_sp = False
            tei.append(f'        <stage>{processed_text}</stage>')

        elif style == "Personaje":
            processed_text = text if text else ""  # Inizializza processed_text con un valore di default
            processed_text = process_annotations_with_ids(
                text, comentario_notes, aparato_notes, annotation_counter, "speaker"
            )
            if in_sp:
                tei.append('        </sp>')
            who_id = find_who_id(text, personaggi)
            tei.append(f'        <sp who="#{who_id}">')
            tei.append(f'        <speaker>{processed_text}</speaker>')
            in_sp = True

        elif style == "Titulo_comedia":
            processed_text = process_annotations_with_ids(text, comentario_notes, aparato_notes, annotation_counter, "head")
        

    if in_cast_list:
        tei.append('        </castList>')
        tei.append('      </div>')
    if in_sp:
        tei.append('        </sp>')
    if in_dedicatoria:
        tei.append('      </div>')
    if in_act:
        tei.append('      </div>')

    tei.extend(['    </body>', '    <back>', '    </back>', '  </text>', '</TEI>'])
    tei_content = "\n".join(tei)

    if save:
        # Se l'utente non pasa output_file, generiamo un nome
        if not output_file:
            output_file = generate_filename(title) + ".xml"
    
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(tei_content)
        print(f"✅ File XML salvado como: {output_file}")
    
    return tei_content

####################################
# 5) GUI MIGLIORATA
####################################

def show_info(message):
    """Muestra una ventana emergente con información"""
    messagebox.showinfo("Información", message)

def main_gui():
    root = tk.Tk()
    root.title("DOCXtoTEI")
    root.geometry("600x500")
    root.configure(bg="#ACBDE9")

    # Stile ttk
    style = ttk.Style(root)
    style.theme_use("vista")
    # Imposta font e background su vari widget.
    # NB: alcuni controlli non ereditano bg facilmente, vedrai più effetto su TLabel/TFrame
    style.configure(".", font=("Open Sans", 10))  # font di default
    style.configure("TFrame", background="#ACBDE9")
    style.configure("TLabel", background="#ACBDE9")
    style.configure("TButton", font=("Open Sans", 10))

    main_frame = ttk.Frame(root, padding="10 10 10 10", style="TFrame")
    main_frame.pack(fill="both", expand=True)

    # Sección 1: Selección de archivos DOCX
    frame_seleccion = ttk.Frame(main_frame, padding="10", style="TFrame", borderwidth=2, relief="ridge")
    frame_seleccion.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

    # Botón info para selección de archivos
    btn_info_seleccion = ttk.Button(frame_seleccion, text="Ayuda", command=lambda: show_info("Seleccione los archivos DOCX que desea convertir a TEI."))
    btn_info_seleccion.grid(row=0, column=2, sticky="e", padx=5)


    # Título de la sección
    ttk.Label(frame_seleccion, text="Selección de archivos", font=("Open Sans", 12, "bold")).grid(row=0, column=0, columnspan=3, pady=5, sticky="w")

    # ---------- RIGA 0: DOCX Principal
    label_main = ttk.Label(frame_seleccion, text="DOCX Principal:")
    label_main.grid(row=1, column=0, sticky="e", pady=5)

    entry_main = ttk.Entry(frame_seleccion, width=60)
    entry_main.grid(row=1, column=1, padx=5, sticky="ew")

    def seleziona_main():
        path = filedialog.askopenfilename(
            title="Seleccione el DOCX Principal",
            filetypes=[("Archivo DOCX", "*.docx"), ("Todos los archivos", "*.*")]
        )
        if path:
            entry_main.delete(0, tk.END)
            entry_main.insert(0, path)

    btn_main = ttk.Button(frame_seleccion, text="Explora...", command=seleziona_main)
    btn_main.grid(row=1, column=2, padx=5)

    # ---------- RIGA 1: DOCX Comentario
    label_com = ttk.Label(frame_seleccion, text="DOCX Comentario:")
    label_com.grid(row=2, column=0, sticky="e", pady=5)

    entry_com = ttk.Entry(frame_seleccion, width=60)
    entry_com.grid(row=2, column=1, padx=5, sticky="ew")

    def seleziona_com():
        path = filedialog.askopenfilename(
            title="Seleccione DOCX Comentario",
            filetypes=[("Archivo DOCX", "*.docx"), ("Todos los archivos", "*.*")]
        )
        if path:
            entry_com.delete(0, tk.END)
            entry_com.insert(0, path)

    btn_com = ttk.Button(frame_seleccion, text="Explora...", command=seleziona_com)
    btn_com.grid(row=2, column=2, padx=5)

    # ---------- RIGA 2: DOCX Aparato
    label_apa = ttk.Label(frame_seleccion, text="DOCX Aparato:")
    label_apa.grid(row=3, column=0, sticky="e", pady=5)

    entry_apa = ttk.Entry(frame_seleccion, width=60)
    entry_apa.grid(row=3, column=1, padx=5, sticky="ew")

    def seleziona_apa():
        path = filedialog.askopenfilename(
            title="Seleccione DOCX Aparato",
            filetypes=[("Archivo DOCX", "*.docx"), ("Todos los archivos", "*.*")]
        )
        if path:
            entry_apa.delete(0, tk.END)
            entry_apa.insert(0, path)

    btn_apa = ttk.Button(frame_seleccion, text="Explora...", command=seleziona_apa)
    btn_apa.grid(row=3, column=2, padx=5)

    # Hace que la columna central (entries) sea expandible
    frame_seleccion.columnconfigure(1, weight=1)


    # Sección 2: Validación y vista previa
    frame_output = ttk.Frame(main_frame, padding="10", style="TFrame", borderwidth=2, relief="ridge")
    frame_output.grid(row=1, column=0, sticky="ew", padx=5, pady=5)

    # Título de la sección
    ttk.Label(frame_output, text="Validación y vista previa", font=("Open Sans", 12, "bold")).grid(row=0, column=0, columnspan=3, pady=5, sticky="w")


    # VISTA PREVIA XML
    def vista_previa():
        main_file = entry_main.get()
        com_file = entry_com.get()
        apa_file = entry_apa.get()

        if not main_file:
            messagebox.showerror("Error", "Seleccione al menos el DOCX Principal!")
            return

        try:
            # Generar el contenido del archivo XML sin guardarlo
            tei_content = convert_docx_to_tei(
                main_docx=main_file,
                comentario_docx=com_file if com_file else None,
                aparato_docx=apa_file if apa_file else None,
                output_file=None,
                save=False
            )

            # Crear una nueva ventana para mostrar la vista previa
            preview_window = tk.Toplevel(root)
            preview_window.title("Vista previa del XML")
            preview_window.geometry("800x600")

            # Crear un widget de texto desplazable para mostrar el contenido del XML
            text_area = scrolledtext.ScrolledText(preview_window, wrap=tk.WORD)
            text_area.pack(fill=tk.BOTH, expand=True)
            text_area.insert(tk.END, tei_content)
            text_area.configure(state='disabled')

        except Exception as e:
            messagebox.showerror("Error", f"Se ha producido un error:\n{e}")
    
    # VISTA PREVIA EDICIÓN DIGITAL
    def vista_previa_edicion_digital():
        main_file = entry_main.get()
        com_file = entry_com.get()
        apa_file = entry_apa.get()

        if not main_file:
            messagebox.showerror("Error", "Seleccione al menos el DOCX Principal!")
            return

        try:
            # Genera el XML sin guardarlo en disco
            tei_content = convert_docx_to_tei(
                main_docx=main_file,
                comentario_docx=com_file if com_file else None,
                aparato_docx=apa_file if apa_file else None,
                output_file=None,
                save=False
            )

            # Construir la plantilla HTML con el contenido inline, usando backticks para evitar problemas de escape
            html_template = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Edición Digital</title>
    <style>
    {ESTILOS_CSS}
    </style>
    <script>
    {CETEI_JS}
    </script>
</head>
<body>
    <div id="tei"></div>
    <script>
        document.addEventListener("DOMContentLoaded", function() {{
            var teiContent = `{tei_content}`;
            var ceteiInstance = new CETEI();
            var htmlNode = ceteiInstance.makeHTML5(teiContent);
            var container = document.getElementById("tei");
            container.innerHTML = "";
            container.appendChild(htmlNode);
        }});
    </script>
</body>
</html>
"""
            with tempfile.NamedTemporaryFile("w", delete=False, suffix=".html", encoding="utf-8") as tmp_file:
                tmp_file.write(html_template)
                tmp_filename = tmp_file.name

            webbrowser.open(f"file://{tmp_filename}")

        except Exception as e:
            messagebox.showerror("Error", f"Se ha producido un error:\n{e}")

    # Botón para validar el marcado
    btn_validar = ttk.Button(frame_output, text="Validar marcado", command=lambda: messagebox.showinfo("Validación", "Función de validación aún no implementada"))
    btn_validar.grid(row=1, column=0, rowspan=2, padx=5, pady=10, sticky="nsew")  # Ocupa dos filas para alinearse con los otros botones

    # Botón de vista previa XML
    btn_vista_previa = ttk.Button(frame_output, text="Vista previa (XML)", command=vista_previa)
    btn_vista_previa.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

    # Botón de vista previa Edición Digital
    btn_vista_previa_edicion = ttk.Button(frame_output, text="Vista previa (edición digital)", command=vista_previa_edicion_digital)
    btn_vista_previa_edicion.grid(row=2, column=1, padx=5, pady=5, sticky="ew")

    # Ajustar las columnas y filas para que los botones se distribuyan correctamente
    frame_output.columnconfigure(0, weight=1) 
    frame_output.columnconfigure(1, weight=2) 
    frame_output.rowconfigure(1, weight=1) 
    frame_output.rowconfigure(2, weight=1)

    # Sección 3: Configuración del output y guardar
    frame_conversion = ttk.Frame(main_frame, padding="10", style="TFrame", borderwidth=2, relief="ridge")
    frame_conversion.grid(row=2, column=0, sticky="ew", padx=5, pady=5)

    # Título de la sección
    ttk.Label(frame_conversion, text="Configuración del output", font=("Open Sans", 12, "bold")).grid(row=0, column=0, columnspan=3, pady=5, sticky="w")

    # Campo para seleccionar el archivo de salida
    label_out = ttk.Label(frame_conversion, text="Output XML:")
    label_out.grid(row=1, column=0, sticky="e", pady=5)

    entry_out = ttk.Entry(frame_conversion, width=60)
    entry_out.grid(row=1, column=1, padx=5, sticky="ew")

    def seleziona_out():
        path = filedialog.asksaveasfilename(
            defaultextension=".xml",
            filetypes=[("Archivo XML", "*.xml"), ("Todos los archivos", "*.*")]
        )
        if path:
            entry_out.delete(0, tk.END)
            entry_out.insert(0, path)

    btn_out = ttk.Button(frame_conversion, text="Guardar como...", command=seleziona_out)
    btn_out.grid(row=1, column=2, padx=5)

    # Botón Convertir
    def avvia_conversione():
        main_file = entry_main.get()
        com_file = entry_com.get()
        apa_file = entry_apa.get()
        out_file = entry_out.get()

        if not main_file:
            messagebox.showerror("Error", "Seleccione al menos el DOCX Principal!")
            return

        try:
            convert_docx_to_tei(
                main_docx=main_file,
                comentario_docx=com_file if com_file else None,
                aparato_docx=apa_file if apa_file else None,
                output_file=out_file if out_file else None,
                save=True
            )
            messagebox.showinfo(
                "Operación completada",
                f"Archivo XML guardado con éxito:\n{out_file if out_file else 'default name'}"
            )
        except Exception as e:
            messagebox.showerror("Error", f"Se ha producido un error:\n{e}")

    btn_converti = ttk.Button(frame_conversion, text="Convertir a TEI y guardar", command=avvia_conversione)
    btn_converti.grid(row=2, column=0, columnspan=3, padx=5, pady=20, sticky="ew")

    # Ajuste para expandir el campo de texto correctamente
    frame_conversion.columnconfigure(1, weight=1)


    # Hace que la columna 0 se expanda
    main_frame.columnconfigure(0, weight=1)

    root.mainloop()


if __name__ == "__main__":
    main_gui()
