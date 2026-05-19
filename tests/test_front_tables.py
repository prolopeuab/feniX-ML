import sys
import re
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "app"))

from tei_backend import convert_docx_to_tei, extract_notes_with_italics  # noqa: E402


class FrontTablesTest(unittest.TestCase):
    @staticmethod
    def _ensure_paragraph_style(doc: Document, style_name: str) -> None:
        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    @classmethod
    def _ensure_required_styles(cls, doc: Document) -> None:
        for style_name in ["Quote", "Acot", "Titulo_comedia", "Acto", "Personaje", "Verso"]:
            cls._ensure_paragraph_style(doc, style_name)

    @staticmethod
    def _append_footnote_reference(paragraph, note_id: str) -> None:
        run = paragraph.add_run()
        footnote_reference = OxmlElement("w:footnoteReference")
        footnote_reference.set(qn("w:id"), note_id)
        run._r.append(footnote_reference)

    @classmethod
    def _add_minimal_body(cls, doc: Document) -> None:
        para = doc.add_paragraph("TITULO DE LA COMEDIA")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("ACTO PRIMERO")
        para.style = "Acto"

        para = doc.add_paragraph("PERSONAJE")
        para.style = "Personaje"

        para = doc.add_paragraph("verso del cuerpo")
        para.style = "Verso"

    @classmethod
    def _build_front_table_docx(cls, output_path: Path) -> None:
        doc = Document()
        cls._ensure_required_styles(doc)

        doc.add_paragraph("Prólogo")

        marked_table = doc.add_table(rows=2, cols=2)
        marked_table.cell(0, 0).text = "^Columna 1"
        marked_table.cell(0, 1).text = "^Columna 2"
        marked_table.cell(1, 0).text = "Dato 1"
        marked_table.cell(1, 1).text = "Dato 2"

        unmarked_table = doc.add_table(rows=1, cols=2)
        unmarked_table.cell(0, 0).text = "Sin marca 1"
        unmarked_table.cell(0, 1).text = "Sin marca 2"

        cls._add_minimal_body(doc)
        doc.save(output_path)

        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    @classmethod
    def _build_footnote_table_docx(cls, output_path: Path) -> None:
        doc = Document()
        cls._ensure_required_styles(doc)

        doc.add_paragraph("Prólogo")
        para = doc.add_paragraph("Texto con nota")
        cls._append_footnote_reference(para, "1")
        para.add_run(".")

        cls._add_minimal_body(doc)
        doc.save(output_path)

        footnotes_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1"/>
  <w:footnote w:id="1">
    <w:p><w:r><w:t>Antes </w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t>cursiva</w:t></w:r></w:p>
    <w:tbl>
      <w:tr>
        <w:tc><w:p><w:r><w:t>^Enc &amp; 1</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>^Enc 2</w:t></w:r></w:p></w:tc>
      </w:tr>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Dato &lt;uno&gt;</w:t></w:r></w:p></w:tc>
        <w:tc>
          <w:p><w:r><w:t>Linea 1</w:t></w:r></w:p>
          <w:p><w:r><w:rPr><w:i/></w:rPr><w:t>Linea 2</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
    </w:tbl>
  </w:footnote>
</w:footnotes>"""
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", footnotes_xml)

    def test_tables_keep_document_order_and_minimal_markup(self):
        xml = convert_docx_to_tei(
            main_docx=str(REPO_ROOT / "test" / "test_prologoycomedia.docx"),
            save=False
        )

        estudio_head = '<head type="divTitle" subtype="MenuLevel_2">Estudio'
        sinopsis_head = '<head type="divTitle" subtype="MenuLevel_2">Sinopsis de la versificación</head>'
        generic_table = '          <table>'
        versification_table = '          <table type="sinopsisversificacion">'

        estudio_pos = xml.index(estudio_head)
        generic_table_pos = xml.index(generic_table)
        sinopsis_pos = xml.index(sinopsis_head)
        versification_pos = xml.index(versification_table)

        self.assertLess(estudio_pos, generic_table_pos)
        self.assertLess(generic_table_pos, sinopsis_pos)
        self.assertLess(sinopsis_pos, versification_pos)

        self.assertEqual(xml.count('type="sinopsisversificacion"'), 1)
        self.assertNotIn('style="', xml)

        generic_table_xml = xml[generic_table_pos:xml.index('          </table>', generic_table_pos) + len('          </table>')]
        self.assertNotIn('type="sinopsisversificacion"', generic_table_xml)
        self.assertNotIn('cols="3"', generic_table_xml)
        self.assertNotIn('rend="summary"', generic_table_xml)
        self.assertIn('<row role="data">', generic_table_xml)
        self.assertIn('<cell role="data">', generic_table_xml)

        versification_table_xml = xml[versification_pos:xml.index('          </table>', versification_pos) + len('          </table>')]
        self.assertIn('<row role="label">', versification_table_xml)
        self.assertIn('<row role="data">', versification_table_xml)
        self.assertIn('cols="3"', versification_table_xml)
        self.assertIn('<cell role="label" cols="3">', versification_table_xml)
        self.assertIn('rend="summary"', versification_table_xml)
        self.assertRegex(
            versification_table_xml,
            r'<row role="label">\s*<cell role="label">\s*<hi rend="italic">Resumen</hi>'
        )
        self.assertNotRegex(
            versification_table_xml,
            r'<row[^>]*rend="summary"[^>]*>\s*<cell[^>]*>\s*<hi rend="italic">Resumen</hi>'
        )
        self.assertRegex(
            versification_table_xml,
            r'<row role="data" rend="summary">\s*<cell role="label">\s*<hi rend="italic">Total</hi>'
        )

    def test_front_table_headers_require_caret_marker(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "front_table_headers.docx"
            self._build_front_table_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        table_positions = [match.start() for match in re.finditer(r'          <table>', xml)]
        self.assertEqual(len(table_positions), 2)

        first_table_xml = xml[table_positions[0]:xml.index('          </table>', table_positions[0]) + len('          </table>')]
        second_table_xml = xml[table_positions[1]:xml.index('          </table>', table_positions[1]) + len('          </table>')]

        self.assertRegex(
            first_table_xml,
            r'<row role="label">\s*<cell role="label">\s*Columna 1\s*</cell>\s*<cell role="label">\s*Columna 2'
        )
        self.assertNotIn("^Columna", first_table_xml)
        self.assertIn('<row role="data">', first_table_xml)

        self.assertNotIn('<row role="label">', second_table_xml)
        self.assertIn('<row role="data">', second_table_xml)
        self.assertIn('Sin marca 1', second_table_xml)

    def test_intro_footnote_tables_become_simple_tei_tables_with_marked_headers(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "footnote_table.docx"
            self._build_footnote_table_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        note_start = xml.index('<note type="intro" n="1">')
        note_end = xml.index('</note>', note_start) + len('</note>')
        note_xml = xml[note_start:note_end]

        self.assertIn('Antes <hi rend="italic">cursiva</hi>', note_xml)
        self.assertIn('<table>', note_xml)
        self.assertIn('<row role="label"><cell role="label">Enc &amp; 1</cell><cell role="label">Enc 2</cell></row>', note_xml)
        self.assertIn('<row><cell>Dato &lt;uno&gt;</cell><cell>Linea 1<lb/><hi rend="italic">Linea 2</hi></cell></row>', note_xml)
        self.assertNotIn("^Enc", note_xml)

    def test_notes_docx_tables_accept_caret_header_marker(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "notes_table.docx"
            doc = Document()
            doc.add_paragraph("1: Nota con tabla")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "^Enc 1"
            table.cell(0, 1).text = "^Enc 2"
            table.cell(1, 0).text = "Dato 1"
            table.cell(1, 1).text = "Dato 2"
            doc.save(docx_path)

            notes = extract_notes_with_italics(str(docx_path))

        self.assertIn(1, notes)
        self.assertEqual(
            notes[1][0],
            'Nota con tabla<table><row role="label"><cell role="label">Enc 1</cell><cell role="label">Enc 2</cell></row><row><cell>Dato 1</cell><cell>Dato 2</cell></row></table>'
        )


if __name__ == "__main__":
    unittest.main()
