import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "app"))

from tei_backend import convert_docx_to_tei  # noqa: E402


class FrontSplitVersesTest(unittest.TestCase):
    @staticmethod
    def _ensure_paragraph_style(doc: Document, style_name: str) -> None:
        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    def _build_test_docx(self, output_path: Path) -> None:
        doc = Document()

        for style_name in [
            "Titulo_comedia",
            "Acto",
            "Personaje",
            "Verso",
            "Partido_inicial",
            "Partido_medio",
            "Partido_final",
        ]:
            self._ensure_paragraph_style(doc, style_name)

        doc.add_paragraph("Prólogo")

        para = doc.add_paragraph("verso front normal")
        para.style = "Verso"

        para = doc.add_paragraph("primera parte del prólogo")
        para.style = "Partido_inicial"

        para = doc.add_paragraph("segunda parte del prólogo")
        para.style = "Partido_medio"

        para = doc.add_paragraph("tercera parte del prólogo")
        para.style = "Partido_final"

        para = doc.add_paragraph("ALGUIEN")
        para.style = "Personaje"

        para = doc.add_paragraph("habla en")
        para.style = "Partido_inicial"

        para = doc.add_paragraph("dos tiempos")
        para.style = "Partido_final"

        para = doc.add_paragraph("TÍTULO DE LA COMEDIA")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("ACTO PRIMERO")
        para.style = "Acto"

        para = doc.add_paragraph("PERSONAJE")
        para.style = "Personaje"

        para = doc.add_paragraph("verso del cuerpo")
        para.style = "Verso"

        para = doc.add_paragraph("cuerpo partido")
        para.style = "Partido_inicial"

        para = doc.add_paragraph("cierre del cuerpo")
        para.style = "Partido_final"

        doc.save(output_path)

        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    def test_front_split_verses_use_part_attributes_without_numbering(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "test_front_split_verses.docx"
            self._build_test_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        front_start = xml.index('<front xml:id="front">')
        front_end = xml.index('    </front>')
        front_xml = xml[front_start:front_end]

        self.assertIn('<l>verso front normal</l>', front_xml)
        self.assertIn('<l part="I">primera parte del prólogo</l>', front_xml)
        self.assertIn('<l part="M">segunda parte del prólogo</l>', front_xml)
        self.assertIn('<l part="F">tercera parte del prólogo</l>', front_xml)
        self.assertNotIn('<l part="I" n="', front_xml)
        self.assertNotIn('<l part="M" n="', front_xml)
        self.assertNotIn('<l part="F" n="', front_xml)

        sp_start = front_xml.index('<sp>')
        sp_end = front_xml.index('          </sp>', sp_start)
        sp_xml = front_xml[sp_start:sp_end]
        self.assertIn('<speaker>ALGUIEN</speaker>', sp_xml)
        self.assertIn('<l part="I">habla en</l>', sp_xml)
        self.assertIn('<l part="F">dos tiempos</l>', sp_xml)

        body_start = xml.index('<body xml:id="body">')
        body_xml = xml[body_start:]
        self.assertIn('<l n="1">verso del cuerpo</l>', body_xml)
        self.assertIn('<l part="I" n="2a">cuerpo partido</l>', body_xml)
        self.assertIn('<l part="F" n="2b">cierre del cuerpo</l>', body_xml)


if __name__ == "__main__":
    unittest.main()
