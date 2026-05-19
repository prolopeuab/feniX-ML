import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "app"))

from tei_backend import convert_docx_to_tei  # noqa: E402


class FrontLinksTest(unittest.TestCase):
    @staticmethod
    def _ensure_paragraph_style(doc: Document, style_name: str) -> None:
        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    @staticmethod
    def _append_hyperlink(paragraph, url: str, runs: list[tuple[str, bool]]) -> None:
        relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        for text, italic in runs:
            run = OxmlElement("w:r")
            if italic:
                run_properties = OxmlElement("w:rPr")
                italic_node = OxmlElement("w:i")
                run_properties.append(italic_node)
                run.append(run_properties)

            text_node = OxmlElement("w:t")
            if text != text.strip():
                text_node.set(qn("xml:space"), "preserve")
            text_node.text = text
            run.append(text_node)
            hyperlink.append(run)

        paragraph._p.append(hyperlink)

    def _build_test_docx(self, output_path: Path) -> None:
        doc = Document()

        for style_name in ["Quote", "Acot", "Titulo_comedia", "Acto", "Personaje", "Verso"]:
            self._ensure_paragraph_style(doc, style_name)

        doc.add_paragraph("Prólogo")

        para = doc.add_paragraph()
        para.add_run("Consulta ")
        self._append_hyperlink(para, "https://ejemplo.org", [("la web", False)])
        para.add_run(" hoy.")

        para = doc.add_paragraph()
        para.style = "Quote"
        para.add_run("Según ")
        self._append_hyperlink(
            para,
            "https://cita.example",
            [("esta", False), (" cita", True)],
        )

        para = doc.add_paragraph()
        para.style = "Acot"
        para.add_run("Sale con ")
        self._append_hyperlink(para, "https://acot.example", [("sigilo", False)])

        table = doc.add_table(rows=1, cols=1)
        cell_para = table.cell(0, 0).paragraphs[0]
        cell_para.add_run("Tabla con ")
        self._append_hyperlink(
            cell_para,
            "https://tabla.example",
            [("enlace", False), (" en cursiva", True)],
        )

        para = doc.add_paragraph("TÍTULO DE LA COMEDIA")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("ACTO PRIMERO")
        para.style = "Acto"

        para = doc.add_paragraph("PERSONAJE")
        para.style = "Personaje"

        para = doc.add_paragraph("verso del cuerpo")
        para.style = "Verso"

        doc.save(output_path)

        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    def _build_footnote_link_docx(self, output_path: Path) -> None:
        doc = Document()

        for style_name in ["Titulo_comedia", "Acto", "Personaje", "Verso"]:
            self._ensure_paragraph_style(doc, style_name)

        doc.add_paragraph("Prólogo")

        para = doc.add_paragraph()
        para.add_run("Nota con enlace")
        footnote_reference = OxmlElement("w:footnoteReference")
        footnote_reference.set(qn("w:id"), "1")
        para.add_run()._r.append(footnote_reference)

        para = doc.add_paragraph("TÍTULO DE LA COMEDIA")
        para.style = "Titulo_comedia"
        para = doc.add_paragraph("ACTO PRIMERO")
        para.style = "Acto"
        para = doc.add_paragraph("PERSONAJE")
        para.style = "Personaje"
        para = doc.add_paragraph("verso del cuerpo")
        para.style = "Verso"

        doc.save(output_path)

        footnotes_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:footnote w:type="separator" w:id="-1"/>
  <w:footnote w:id="1">
    <w:p>
      <w:r><w:t>Consulta </w:t></w:r>
      <w:hyperlink r:id="rIdFootnoteLink">
        <w:r><w:t>este enlace</w:t></w:r>
        <w:r><w:rPr><w:i/></w:rPr><w:t> útil</w:t></w:r>
      </w:hyperlink>
      <w:r><w:t>.</w:t></w:r>
    </w:p>
  </w:footnote>
</w:footnotes>"""
        footnotes_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rIdFootnoteLink"
                Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
                Target="https://nota.example"
                TargetMode="External"/>
</Relationships>"""
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", footnotes_xml)
            docx_zip.writestr("word/_rels/footnotes.xml.rels", footnotes_rels)

    def test_front_hyperlinks_become_tei_refs(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "test_front_links.docx"
            self._build_test_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        front_start = xml.index('<front xml:id="front">')
        front_end = xml.index('    </front>')
        front_xml = xml[front_start:front_end]

        self.assertIn(
            '<p>Consulta <ref target="https://ejemplo.org">la web</ref> hoy.</p>',
            front_xml,
        )
        self.assertIn(
            '<quote>Según <ref target="https://cita.example">esta<hi rend="italic"> cita</hi></ref></quote>',
            front_xml,
        )
        self.assertIn(
            '<stage>Sale con <ref target="https://acot.example">sigilo</ref></stage>',
            front_xml,
        )
        self.assertIn(
            '<ref target="https://tabla.example">enlace<hi rend="italic"> en cursiva</hi></ref>',
            front_xml,
        )

        body_start = xml.index('<body xml:id="body">')
        body_xml = xml[body_start:]
        self.assertIn('<l n="1">verso del cuerpo</l>', body_xml)

    def test_intro_footnote_hyperlinks_become_tei_refs(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "test_footnote_links.docx"
            self._build_footnote_link_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        self.assertIn(
            '<note type="intro" n="1">Consulta <ref target="https://nota.example">este enlace<hi rend="italic"> útil</hi></ref>.</note>',
            xml,
        )


if __name__ == "__main__":
    unittest.main()
