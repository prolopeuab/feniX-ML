import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "app"))

from tei_backend import convert_docx_to_tei  # noqa: E402


class BodyRepeatedTitlesAndDramatisTest(unittest.TestCase):
    @staticmethod
    def _ensure_paragraph_style(doc: Document, style_name: str) -> None:
        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    def _build_test_docx(self, output_path: Path) -> None:
        doc = Document()

        for style_name in [
            "Titulo_comedia",
            "Acto",
            "Epigr_Dramatis",
            "Dramatis_lista",
            "Personaje",
            "Verso",
        ]:
            self._ensure_paragraph_style(doc, style_name)

        para = doc.add_paragraph("OBRA BASE")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("PERSONAJES")
        para.style = "Epigr_Dramatis"

        para = doc.add_paragraph("AURORA, dama")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("CARLOS")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("OBRA REITERADA ANTES")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("SUBTITULO ANTES")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("Acto primero")
        para.style = "Acto"

        para = doc.add_paragraph("PERSONAS DEL ACTO PRIMERO")
        para.style = "Epigr_Dramatis"

        para = doc.add_paragraph("AURORA, criada")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("AURORA")
        para.style = "Personaje"

        para = doc.add_paragraph("Habla Aurora en acto uno")
        para.style = "Verso"

        para = doc.add_paragraph("CARLOS")
        para.style = "Personaje"

        para = doc.add_paragraph("Habla Carlos en acto uno")
        para.style = "Verso"

        para = doc.add_paragraph("Acto segundo")
        para.style = "Acto"

        para = doc.add_paragraph("OBRA REITERADA DESPUES")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("SUBTITULO DESPUES")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("PERSONAS DEL ACTO SEGUNDO")
        para.style = "Epigr_Dramatis"

        para = doc.add_paragraph("AURORA, reina")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("AURORA")
        para.style = "Personaje"

        para = doc.add_paragraph("Habla Aurora en acto dos")
        para.style = "Verso"

        doc.save(output_path)

        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    def _build_pre_act_dramatis_docx(self, output_path: Path) -> None:
        doc = Document()

        for style_name in [
            "Titulo_comedia",
            "Acto",
            "Epigr_Dramatis",
            "Dramatis_lista",
            "Personaje",
            "Verso",
        ]:
            self._ensure_paragraph_style(doc, style_name)

        para = doc.add_paragraph("OBRA BASE")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("TITULO ACTO UNO")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("PERSONAS DEL ACTO UNO")
        para.style = "Epigr_Dramatis"

        para = doc.add_paragraph("AURORA, criada")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("Acto primero")
        para.style = "Acto"

        para = doc.add_paragraph("AURORA")
        para.style = "Personaje"

        para = doc.add_paragraph("Habla Aurora en acto uno")
        para.style = "Verso"

        para = doc.add_paragraph("TITULO ACTO DOS")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("PERSONAS DEL ACTO DOS")
        para.style = "Epigr_Dramatis"

        para = doc.add_paragraph("AURORA, reina")
        para.style = "Dramatis_lista"

        para = doc.add_paragraph("Acto segundo")
        para.style = "Acto"

        para = doc.add_paragraph("AURORA")
        para.style = "Personaje"

        para = doc.add_paragraph("Habla Aurora en acto dos")
        para.style = "Verso"

        doc.save(output_path)

        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(output_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    def test_repeated_titles_and_dramatis_are_scoped_to_each_act(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "test_body_repeated_titles_and_dramatis.docx"
            self._build_test_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        self.assertIn('<head type="mainTitle" xml:id="titulo">OBRA BASE</head>', xml)
        self.assertIn('<head type="mainTitle" subtype="repeated">OBRA REITERADA ANTES</head>', xml)
        self.assertIn('<head type="subTitle" subtype="repeated">SUBTITULO ANTES</head>', xml)
        self.assertIn('<head type="mainTitle" subtype="repeated">OBRA REITERADA DESPUES</head>', xml)
        self.assertIn('<head type="subTitle" subtype="repeated">SUBTITULO DESPUES</head>', xml)

        acto1_pos = xml.index('xml:id="acto1"')
        acto2_pos = xml.index('xml:id="acto2"')
        acto1_head_pos = xml.index('<head type="acto">ACTO PRIMERO</head>')
        acto2_head_pos = xml.index('<head type="acto">ACTO SEGUNDO</head>')
        titulo_acto1_pos = xml.index('<head type="mainTitle" subtype="repeated">OBRA REITERADA ANTES</head>')
        subtitulo_acto1_pos = xml.index('<head type="subTitle" subtype="repeated">SUBTITULO ANTES</head>')
        titulo_acto2_pos = xml.index('<head type="mainTitle" subtype="repeated">OBRA REITERADA DESPUES</head>')
        subtitulo_acto2_pos = xml.index('<head type="subTitle" subtype="repeated">SUBTITULO DESPUES</head>')
        cast_act1_pos = xml.index('xml:id="personajes_acto1"')
        cast_act2_pos = xml.index('xml:id="personajes_acto2"')
        self.assertLess(acto1_pos, titulo_acto1_pos)
        self.assertLess(titulo_acto1_pos, subtitulo_acto1_pos)
        self.assertLess(subtitulo_acto1_pos, acto1_head_pos)
        self.assertLess(acto1_head_pos, cast_act1_pos)
        self.assertLess(cast_act1_pos, acto2_pos)
        self.assertLess(acto2_pos, acto2_head_pos)
        self.assertLess(acto2_head_pos, titulo_acto2_pos)
        self.assertLess(titulo_acto2_pos, subtitulo_acto2_pos)
        self.assertLess(subtitulo_acto2_pos, cast_act2_pos)

        self.assertIn('xml:id="aurora_dama"', xml)
        self.assertIn('xml:id="carlos"', xml)
        self.assertIn('xml:id="acto1_aurora_criada"', xml)
        self.assertIn('xml:id="acto2_aurora_reina"', xml)

        self.assertIn('<sp who="#acto1_aurora_criada">', xml)
        self.assertIn('<sp who="#carlos">', xml)
        self.assertIn('<sp who="#acto2_aurora_reina">', xml)

    def test_pre_act_title_and_dramatis_attach_to_following_act(self):
        with TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "test_pre_act_dramatis.docx"
            self._build_pre_act_dramatis_docx(docx_path)
            xml = convert_docx_to_tei(main_docx=str(docx_path), save=False)

        acto1_pos = xml.index('xml:id="acto1"')
        cast_act1_pos = xml.index('xml:id="personajes_acto1"')
        acto2_pos = xml.index('xml:id="acto2"')
        cast_act2_pos = xml.index('xml:id="personajes_acto2"')
        acto1_head_pos = xml.index('<head type="acto">ACTO PRIMERO</head>')
        acto2_head_pos = xml.index('<head type="acto">ACTO SEGUNDO</head>')
        titulo_acto1_pos = xml.index('<head type="mainTitle" subtype="repeated">TITULO ACTO UNO</head>')
        titulo_acto2_pos = xml.index('<head type="mainTitle" subtype="repeated">TITULO ACTO DOS</head>')

        self.assertLess(acto1_pos, titulo_acto1_pos)
        self.assertLess(titulo_acto1_pos, cast_act1_pos)
        self.assertLess(cast_act1_pos, acto1_head_pos)
        self.assertLess(acto1_head_pos, acto2_pos)
        self.assertLess(acto2_pos, titulo_acto2_pos)
        self.assertLess(titulo_acto2_pos, cast_act2_pos)
        self.assertLess(cast_act2_pos, acto2_head_pos)
        self.assertIn('<head type="mainTitle" subtype="repeated">TITULO ACTO UNO</head>', xml)
        self.assertIn('<head type="mainTitle" subtype="repeated">TITULO ACTO DOS</head>', xml)
        self.assertIn('<sp who="#acto1_aurora_criada">', xml)
        self.assertIn('<sp who="#acto2_aurora_reina">', xml)


if __name__ == "__main__":
    unittest.main()
