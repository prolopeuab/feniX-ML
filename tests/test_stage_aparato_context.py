import re
import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "app"))

from tei_backend import convert_docx_to_tei  # noqa: E402


class StageAparatoContextTest(unittest.TestCase):
    @staticmethod
    def _ensure_paragraph_style(doc: Document, style_name: str) -> None:
        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    @staticmethod
    def _add_empty_footnotes(docx_path: Path) -> None:
        empty_footnotes = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        with zipfile.ZipFile(docx_path, "a") as docx_zip:
            docx_zip.writestr("word/footnotes.xml", empty_footnotes)

    def _build_main_docx(self, output_path: Path) -> None:
        doc = Document()
        for style_name in ["Titulo_comedia", "Acto", "Personaje", "Verso", "Acot"]:
            self._ensure_paragraph_style(doc, style_name)

        para = doc.add_paragraph("COMEDIA")
        para.style = "Titulo_comedia"

        para = doc.add_paragraph("Acto primero")
        para.style = "Acto"

        para = doc.add_paragraph("UNO")
        para.style = "Personaje"

        for verse_number in range(1, 166):
            para = doc.add_paragraph(f"verso {verse_number}")
            para.style = "Verso"

        para = doc.add_paragraph("%Patean")
        para.style = "Acot"

        for verse_number in range(166, 168):
            para = doc.add_paragraph(f"verso {verse_number}")
            para.style = "Verso"

        para = doc.add_paragraph("%Patean")
        para.style = "Acot"

        doc.save(output_path)
        self._add_empty_footnotes(output_path)

    def _build_aparato_docx(self, output_path: Path, entries: list[tuple[str, str | None, str]]) -> None:
        doc = Document()
        for prefix, acot_ref, suffix in entries:
            para = doc.add_paragraph()
            para.add_run(prefix)
            if acot_ref is not None:
                number = re.match(r'^(\d+)', acot_ref).group(1)
                para.add_run(number)
                run = para.add_run("Acot")
                run.italic = True
            para.add_run(suffix)
        doc.save(output_path)

    @staticmethod
    def _stage_texts(xml: str) -> list[str]:
        return re.findall(r'<stage>(.*?)</stage>', xml, flags=re.DOTALL)

    def test_stage_aparato_uses_acot_reference_before_consuming_next_note(self):
        with TemporaryDirectory() as tmp_dir:
            main_docx = Path(tmp_dir) / "main.docx"
            aparato_docx = Path(tmp_dir) / "aparato.docx"
            self._build_main_docx(main_docx)
            self._build_aparato_docx(
                aparato_docx,
                [
                    ("%Patean: ", "165Acot", " Patean A Par Men : om Ma"),
                    ("%Patean: ", "165Acot", " En Par la acotacion figura a la derecha del verso"),
                    ("%Patean: ", "167Acot", " Ma y Par situan la acotacion a la derecha del verso"),
                ],
            )

            xml = convert_docx_to_tei(
                main_docx=str(main_docx),
                aparato_docx=str(aparato_docx),
                save=False,
            )

        stages = self._stage_texts(xml)
        self.assertEqual(2, len(stages))
        first_stage, second_stage = stages

        self.assertIn('xml:id="a_patean_stage_1"', first_stage)
        self.assertIn('xml:id="a_patean_stage_2"', first_stage)
        self.assertIn('165<hi rend="italic">Acot</hi> Patean', first_stage)
        self.assertIn('165<hi rend="italic">Acot</hi> En Par', first_stage)
        self.assertNotIn('167<hi rend="italic">Acot</hi>', first_stage)

        self.assertIn('xml:id="a_patean_stage_3"', second_stage)
        self.assertIn('167<hi rend="italic">Acot</hi> Ma y Par', second_stage)
        self.assertNotIn('xml:id="a_patean_stage_2"', second_stage)

    def test_stage_aparato_without_acot_reference_stays_sequential(self):
        with TemporaryDirectory() as tmp_dir:
            main_docx = Path(tmp_dir) / "main.docx"
            aparato_docx = Path(tmp_dir) / "aparato.docx"
            self._build_main_docx(main_docx)
            self._build_aparato_docx(
                aparato_docx,
                [
                    ("%Patean: ", None, "primera nota simple"),
                    ("%Patean: ", None, "segunda nota simple"),
                ],
            )

            xml = convert_docx_to_tei(
                main_docx=str(main_docx),
                aparato_docx=str(aparato_docx),
                save=False,
            )

        first_stage, second_stage = self._stage_texts(xml)
        self.assertIn("primera nota simple", first_stage)
        self.assertNotIn("segunda nota simple", first_stage)
        self.assertIn("segunda nota simple", second_stage)


if __name__ == "__main__":
    unittest.main()
