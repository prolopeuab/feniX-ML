# ==========================================
# feniX-ML: Marcado automático de DOCX a TEI/XML
# Desarrollado por Anna Abate, Emanuele Leboffe y David Merino Recalde.
# Grupo de investigación PROLOPE, Universitat Autònoma de Barcelona
# Descripción: Funciones para convertir textos teatrales en formato DOCX a TEI/XML, incluyendo manejo de notas, metadatos y validaciones.
# Este script debe utilizarse junto a visualizacion.py, gui.py y main.py.
# ==========================================

# --- Importaciones
import os
import re
import unicodedata
import zipfile
from lxml import etree
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from difflib import get_close_matches
from typing import Any, Optional, TypedDict, cast

APP_VERSION = "1.1.0"


# --- Funciones de escape XML
def escape_xml(text):
    """
    Escapa caracteres especiales de XML en el texto.
    """
    if not text:
        return text
    # Orden importante: & primero para no escapar los escapes
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    text = text.replace("'", "&apos;")
    return text


# --- Extracción y procesamiento de notas en el prólogo
# Funciones para extraer y procesar notas a pie de página del prólogo o introducción.

def extract_intro_footnotes(docx_path):
    """
    Extrae todas las notas a pie de página de un archivo DOCX, preservando cursivas.
    Devuelve un diccionario {id: note_text} con formato TEI (incluyendo <hi rend="italic">).
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    footnote_dict = {}

    with zipfile.ZipFile(docx_path) as docx_zip:
        with docx_zip.open("word/footnotes.xml") as footnote_file:
            root = etree.parse(footnote_file).getroot()

            for note in root.xpath("//w:footnote[not(@w:type='separator')]", namespaces=ns):
                note_id = note.get(qn("w:id"))
                
                # Procesar todos los runs dentro de la nota para detectar cursivas
                full_text = ""
                for run in note.xpath(".//w:r", namespaces=ns):
                    # Verificar si el run tiene formato cursiva
                    is_italic = run.xpath(".//w:i", namespaces=ns) or run.xpath(".//w:iCs", namespaces=ns)
                    
                    # Extraer el texto del run
                    text_elements = run.xpath(".//w:t", namespaces=ns)
                    run_text = "".join(t.text for t in text_elements if t.text is not None)
                    
                    if run_text:
                        if is_italic:
                            # Envolver en cursiva y escapar el contenido
                            full_text += f'<hi rend="italic">{escape_xml(run_text)}</hi>'
                        else:
                            # Escapar el texto normal
                            full_text += escape_xml(run_text)
                
                if full_text.strip():
                    footnote_dict[note_id] = full_text.strip()

    return footnote_dict

def extract_text_with_intro_notes(para, footnotes_intro):
    """
    Extrae el texto de un párrafo, insertando las notas en el lugar correspondiente.
    """
    def flush_intro_text_buffer(parts: list[str], buffer: list[str], italic: bool) -> None:
        if not buffer:
            return

        text_chunk = escape_xml("".join(buffer))
        if italic:
            parts.append(f'<hi rend="italic">{text_chunk}</hi>')
        else:
            parts.append(text_chunk)
        buffer.clear()

    def extract_intro_run(run: Run) -> str:
        parts: list[str] = []
        text_buffer: list[str] = []

        for child in run._element.iterchildren():
            if child.tag == qn("w:t"):
                if child.text:
                    text_buffer.append(child.text)
            elif child.tag == qn("w:tab"):
                text_buffer.append("\t")
            elif child.tag in (qn("w:br"), qn("w:cr")):
                text_buffer.append("\n")
            elif child.tag == qn("w:footnoteReference"):
                flush_intro_text_buffer(parts, text_buffer, bool(run.italic))
                note_id = child.get(qn("w:id"))
                if note_id:
                    note_text = footnotes_intro.get(note_id, "")
                    parts.append(f'<note type="intro" n="{note_id}">{note_text}</note>')

        flush_intro_text_buffer(parts, text_buffer, bool(run.italic))
        return "".join(parts)

    def extract_intro_hyperlink(hyperlink: Hyperlink) -> str:
        content = "".join(extract_intro_run(run) for run in hyperlink.runs)
        target = hyperlink.url or hyperlink.address
        if target:
            return f'<ref target="{escape_xml(target)}">{content}</ref>'
        return content

    parts: list[str] = []
    for item in para.iter_inner_content():
        if isinstance(item, Hyperlink):
            parts.append(extract_intro_hyperlink(item))
        else:
            parts.append(extract_intro_run(item))

    return "".join(parts).strip()

# --- Manejo de bloques estructurales TEI
def close_cast_list(tei, state):
    """
    Cierra un bloque <div type="castList"> abierto, respetando su indentación.
    """
    if not state.get("in_cast_list"):
        return

    cast_list_indent = state.get("cast_list_indent", "          ")
    cast_list_div_indent = state.get("cast_list_div_indent", "        ")
    tei.append(f'{cast_list_indent}</castList>')
    tei.append(f'{cast_list_div_indent}</div>')
    state["in_cast_list"] = False
    state["cast_list_div_indent"] = None
    state["cast_list_indent"] = None
    state["cast_item_indent"] = None
    state["cast_list_scope"] = None


def close_current_blocks(tei, state, current_act_characters=None):
    """
    Cierra todos los bloques TEI abiertos según el estado actual.
    
    Valida y cierra: <sp> (diálogos), <castList> (reparto), <div> (dedicatorias y actos).
    Requiere que el diccionario `state` tenga claves: in_sp, in_cast_list, in_dedicatoria, in_act.
    
    Args:
        tei: Lista de líneas TEI donde se añaden los cierres.
        state: Dict con flags de bloques abiertos {in_sp: bool, in_cast_list: bool, ...}.
    """
    # Cierra los bloques abiertos en el estado actual
    if state.get("in_sp"):
        tei.append('        </sp>')
        state["in_sp"] = False
    close_cast_list(tei, state)
    if state.get("in_dedicatoria"):
        tei.append('        </div>')
        state["in_dedicatoria"] = False
    if state.get("in_act"):
        tei.append('        </div>')
        state["in_act"] = False
        if current_act_characters is not None:
            current_act_characters.clear()

# --- Funciones de soporte para texto
def extract_text_with_italics(para):
    """
    Extrae el texto de un párrafo, preservando las cursivas.
    Mueve los espacios que están dentro de runs cursivos hacia fuera para preservar el formato.
    """
    # Recorre los runs del párrafo y envuelve en <hi rend="italic"> si es cursiva
    text = ""
    for run in para.runs:
        if run.italic:
            # Extrae espacios del principio y final del texto cursivo
            content = run.text
            leading_spaces = ""
            trailing_spaces = ""
            
            # Extrae espacios del principio
            while content.startswith(" ") or content.startswith("\t"):
                leading_spaces += content[0]
                content = content[1:]
            
            # Extrae espacios del final
            while content.endswith(" ") or content.endswith("\t"):
                trailing_spaces = content[-1] + trailing_spaces
                content = content[:-1]
            
            # Solo envuelve en cursiva el contenido sin espacios, escapando caracteres XML
            if content:  # solo si queda contenido después de quitar espacios
                text += leading_spaces + f'<hi rend="italic">{escape_xml(content)}</hi>' + trailing_spaces
            else:  # si solo había espacios, los añade sin cursiva
                text += run.text
        else:
            # Escapar caracteres XML en texto normal
            text += escape_xml(run.text)
    return text.strip()

def normalize_id(text):
    """
    Normaliza texto para crear IDs XML seguros.
    Elimina tildes, convierte a minúsculas, reemplaza espacios y caracteres especiales por guiones bajos.
    """
    # Descomponer caracteres acentuados
    normalized = unicodedata.normalize('NFKD', text)
    # Eliminar marcas diacríticas (tildes, acentos, etc.)
    normalized = normalized.encode('ascii', 'ignore').decode('utf-8')
    # Reemplazar caracteres especiales por guiones bajos
    normalized = re.sub(r'[^A-Za-z0-9_-]+', '_', normalized)
    # Convertir a minúsculas
    normalized = normalized.lower()
    return normalized

def normalize_milestone_type(raw: str) -> str:
    """
    Normaliza el valor de milestone/@type a un slug XML-safe:
    - sin tildes
    - en minúsculas
    - espacios y separadores -> guion
    - sin caracteres fuera de [a-z0-9-]
    """
    # Limpieza inicial de extremos
    normalized = raw.strip()
    # Quitar diacríticos y limitar a ASCII
    normalized = unicodedata.normalize("NFKD", normalized)
    normalized = normalized.encode("ascii", "ignore").decode("utf-8")
    # Forzar minúsculas
    normalized = normalized.lower()
    # Espacios consecutivos -> guion
    normalized = re.sub(r"\s+", "-", normalized)
    # Cualquier carácter no permitido -> guion
    normalized = re.sub(r"[^a-z0-9-]+", "-", normalized)
    # Colapsar guiones repetidos y recortar extremos
    normalized = re.sub(r"-{2,}", "-", normalized).strip("-")
    return normalized


def normalize_text_for_matching(text: str) -> str:
    """
    Normaliza texto para comparaciones flexibles:
    - sin tildes
    - en minúsculas
    - sin puntuación
    - con espacios colapsados
    """
    normalized = unicodedata.normalize("NFKD", text)
    normalized = normalized.encode("ascii", "ignore").decode("utf-8")
    normalized = normalized.lower()
    normalized = re.sub(r"[^a-z0-9\s]+", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


class PendingSplitVerse(TypedDict):
    verse_number: int
    initial_text: str
    parts: list[str]
    has_notes: bool


def get_pending_split_verse(state: dict[str, Any]) -> Optional[PendingSplitVerse]:
    """
    Devuelve el estado tipado del verso partido pendiente, si existe.
    """
    return cast(Optional[PendingSplitVerse], state.get("pending_split_verse"))


def iter_document_blocks(doc: Any):
    """
    Itera por los bloques de primer nivel del documento en su orden real:
    párrafos y tablas.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def get_front_blocks(doc: Any, title_paragraph: Paragraph):
    """
    Devuelve los bloques del front-matter (párrafos y tablas) hasta el
    primer párrafo con estilo Titulo_comedia.
    """
    front_blocks = []
    title_element = title_paragraph._element

    for block in iter_document_blocks(doc):
        if block._element is title_element:
            break
        front_blocks.append(block)

    return front_blocks


def extract_table_cell_contents(cell, footnotes_intro):
    """
    Extrae el contenido no vacío de todos los párrafos de una celda.
    """
    contents = []
    for para in cell.paragraphs:
        raw = extract_text_with_intro_notes(para, footnotes_intro).strip()
        if raw:
            contents.append(raw)
    return contents


def append_tei_cell(tei_lines, cell_contents, indent, attrs=""):
    """
    Añade una celda TEI con su contenido ya procesado.
    """
    if not cell_contents:
        tei_lines.append(f"{indent}<cell{attrs}></cell>")
        return

    tei_lines.append(f"{indent}<cell{attrs}>")
    if len(cell_contents) == 1:
        tei_lines.append(f"{indent}  {cell_contents[0]}")
    else:
        for content in cell_contents:
            tei_lines.append(f"{indent}  <p>{content}</p>")
    tei_lines.append(f"{indent}</cell>")


def is_versification_section(current_section: Optional[str]) -> bool:
    """
    Detecta si la subsección actual corresponde a la sinopsis de versificación.
    """
    return current_section == "sinopsis de la versificacion"


def is_versification_act_heading(texts) -> bool:
    """
    Detecta filas-título del tipo "Acto X" o "X acto".
    """
    if not texts or not texts[0].strip() or any(text.strip() for text in texts[1:]):
        return False

    first = normalize_text_for_matching(texts[0])
    return bool(re.match(r"^acto\s+\w+", first) or re.match(r"^\w+(?:\s+\w+)?\s+acto$", first))


def is_versification_summary_header_row(texts) -> bool:
    """
    Detecta la fila-cabecera del bloque de resumen en la sinopsis de versificación.
    """
    if not texts or not texts[0].strip():
        return False

    first = normalize_text_for_matching(texts[0])
    return first == "resumen"


def is_versification_total_row(texts) -> bool:
    """
    Detecta filas de cierre con el literal "Total" en la sinopsis de versificación.
    """
    if not texts or not texts[0].strip():
        return False

    first = normalize_text_for_matching(texts[0])
    return first == "total"

def uppercase_preserve_tags(text):
    """
    Convierte texto a mayúsculas preservando etiquetas XML.
    Ejemplo: "<hi rend='italic'>palabra</hi>" → "<hi rend='italic'>PALABRA</hi>"
    """
    parts = re.split(r'(<[^>]+>)', text)
    return ''.join(part.upper() if not part.startswith('<') else part for part in parts)


def uppercase_preserve_tags_and_note_content(text):
    """
    Convierte texto a mayúsculas preservando etiquetas XML, pero sin alterar
    el contenido interno de bloques <note ...>...</note>.

    Útil para secciones que se muestran en mayúsculas (speaker, títulos, actos)
    donde las notas editoriales deben conservar su capitalización original.
    """
    if not text:
        return text

    # Proteger temporalmente las notas completas para no alterar su contenido.
    note_pattern = re.compile(r'(<note\b[^>]*>.*?</note>)', flags=re.IGNORECASE | re.DOTALL)
    protected_notes = []

    def _protect_note(match):
        protected_notes.append(match.group(1))
        return f"<<<NOTE_BLOCK_{len(protected_notes) - 1}>>>"

    protected_text = note_pattern.sub(_protect_note, text)
    uppercased = uppercase_preserve_tags(protected_text)

    # Restaurar las notas exactamente como estaban.
    for idx, note_block in enumerate(protected_notes):
        uppercased = uppercased.replace(f"<<<NOTE_BLOCK_{idx}>>>", note_block)

    return uppercased

def extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, section):
    """
    Extrae texto de un párrafo preservando cursivas y procesando anotaciones (@palabra, %palabra, @%palabra).
    
    Utiliza marcadores internos para cursivas antes de procesar anotaciones, evitando interferencias
    entre símbolos de anotación y marcas de formato. Genera referencias XML con IDs únicos.
    
    Args:
        para: Párrafo de python-docx con texto y formato.
        nota_notes: Dict con notas filológicas {palabra_normalizada: contenido}.
        aparato_notes: Dict con notas de aparato crítico {palabra_normalizada: contenido}.
        annotation_counter: Dict para mantener sincronización de índices de notas.
        section: Identificador de sección para generar xml:ids únicos.
    
    Returns:
        str: Texto con etiquetas XML de cursiva (<hi rend="italic">) y notas (<note>) integradas.
    """
    
    # Función interna para normalizar palabras
    def normalize_word(word):
        normalized = unicodedata.normalize('NFKD', word)
        normalized = normalized.encode('ASCII', 'ignore').decode('utf-8').lower().strip()
        return normalized
    
    # Crear conjunto de todas las palabras que tienen anotaciones
    all_keys = set()
    if nota_notes:
        all_keys.update(nota_notes.keys())
    if aparato_notes:
        all_keys.update(aparato_notes.keys())
    
    # Mantener contadores separados de notas filológicas y aparato crítico para sincronización secuencial
    nota_counters = annotation_counter.setdefault("_occurrences_nota", {})
    aparato_counters = annotation_counter.setdefault("_occurrences_aparato", {})
    
    # PASO 1: Construir texto con placeholders de cursiva antes de procesar anotaciones
    marked_text = ""
    prev_italic = False
    
    for run in para.runs:
        if not run.text:
            continue
        
        # Detectar cambios en cursiva y agregar marcadores
        if run.italic and not prev_italic:
            marked_text += '<<<ITALIC_START>>>'
        elif not run.italic and prev_italic:
            marked_text += '<<<ITALIC_END>>>'
        
        marked_text += run.text
        prev_italic = run.italic
    
    # Cerrar cursiva si quedó abierta
    if prev_italic:
        marked_text += '<<<ITALIC_END>>>'
    
    # PASO 2: Procesar anotaciones (@palabra, %palabra, @%palabra)
    # Usar placeholders Unicode para no interferir con regex mientras se preservan cursivas
    
    # Usar caracteres Unicode de control que no aparecen en texto normal
    ITALIC_START_PLACEHOLDER = '\u0001'  # SOH - no causa conflicto con regex
    ITALIC_END_PLACEHOLDER = '\u0002'    # STX - no causa conflicto con regex
    
    # Reemplazar marcadores temporales por placeholders para protegerlos del regex
    text_with_placeholders = marked_text.replace('<<<ITALIC_START>>>', ITALIC_START_PLACEHOLDER)
    text_with_placeholders = text_with_placeholders.replace('<<<ITALIC_END>>>', ITALIC_END_PLACEHOLDER)
    
    # Función para reemplazar anotaciones (@palabra, %palabra, @%palabra) con notas XML
    # Maneja sincronización de múltiples notas por palabra (lista de anotaciones)
    def replace_at_word(match):
        symbol = match.group(1)  # '@', '%' o '@%'
        placeholders_before = match.group(2) if match.group(2) else ''  # Placeholders entre símbolo y palabra
        word = match.group(3)  # palabra
        key = normalize_word(word)
        
        if key not in all_keys:
            # Sin notas, solo quitar el símbolo y mantener placeholders
            return placeholders_before + word
        
        # Construir las notas
        notes = []
        
        # Notas filológicas - solo si tiene @ (@ o @%)
        if '@' in symbol and key in nota_notes:
            nota_index = nota_counters.get(key, 0)
            # Validación defensiva: asegurar que nota_index nunca sea None
            if nota_index is None:
                nota_index = 0
            nota_list = nota_notes[key] if isinstance(nota_notes[key], list) else [nota_notes[key]]
            if nota_index < len(nota_list):
                content = nota_list[nota_index]
                xml_id = f"n_{key}_{section}_{nota_index + 1}"
                xml_id = re.sub(r'[^a-zA-Z0-9_]', '', xml_id).lower()
                notes.append(f'<<<NOTE>>><note subtype="nota" xml:id="{xml_id}">{content}</note><<<ENDNOTE>>>')
            nota_counters[key] = nota_index + 1
        
        # Aparato crítico - solo si tiene % (% o @%)
        if '%' in symbol and key in aparato_notes:
            aparato_index = aparato_counters.get(key, 0)
            # Validación defensiva: asegurar que aparato_index nunca sea None
            if aparato_index is None:
                aparato_index = 0
            aparato_list = aparato_notes[key] if isinstance(aparato_notes[key], list) else [aparato_notes[key]]
            if aparato_index < len(aparato_list):
                content = aparato_list[aparato_index]
                xml_id = f"a_{key}_{section}_{aparato_index + 1}"
                xml_id = re.sub(r'[^a-zA-Z0-9_]', '', xml_id).lower()
                notes.append(f'<<<NOTE>>><note subtype="aparato" xml:id="{xml_id}">{content}</note><<<ENDNOTE>>>')
            aparato_counters[key] = aparato_index + 1
        
        # Devolver: placeholders + palabra + notas
        return placeholders_before + word + ''.join(notes)
    
    # Regex para capturar anotaciones: símbolo(s), placeholders opcionales, palabra
    # Grupo 1: @, % o @%  (símbolos de anotación)
    # Grupo 2: \u0001|\u0002 opcionales (placeholders de cursiva entre símbolo y palabra)
    # Grupo 3: \w+ (la palabra a anotar)
    pattern = r'(@%?|%)([\u0001\u0002]*)(\w+)'
    processed_text = re.sub(pattern, replace_at_word, text_with_placeholders)
    
    # PASO 3: Restaurar marcadores de cursiva
    # Convertir placeholders de vuelta a marcadores legibles
    processed_text = processed_text.replace(ITALIC_START_PLACEHOLDER, '<<<ITALIC_START>>>')
    processed_text = processed_text.replace(ITALIC_END_PLACEHOLDER, '<<<ITALIC_END>>>')
    
    # PASO 4: Escapar XML (excepto notas y marcadores)
    # Separar notas para no escapar su contenido XML ya procesado
    parts = re.split(r'(<<<NOTE>>>.*?<<<ENDNOTE>>>)', processed_text, flags=re.DOTALL)
    escaped_parts = []
    for part in parts:
        if part.startswith('<<<NOTE>>>'):
            # Es una nota, solo quitar marcadores
            note_content = part.replace('<<<NOTE>>>', '').replace('<<<ENDNOTE>>>', '')
            escaped_parts.append(f'<<<NOTE>>>{note_content}<<<ENDNOTE>>>')
        else:
            # Es texto, escapar pero preservar marcadores de cursiva
            # Separar por marcadores de cursiva
            italic_parts = re.split(r'(<<<ITALIC_START>>>|<<<ITALIC_END>>>)', part)
            for ip in italic_parts:
                if ip in ['<<<ITALIC_START>>>', '<<<ITALIC_END>>>']:
                    escaped_parts.append(ip)
                else:
                    escaped_parts.append(escape_xml(ip))
    
    processed_text = ''.join(escaped_parts)
    
    # PASO 5: Convertir marcadores internos a etiquetas XML
    # Limpiar marcadores de delimitación de notas y convertir cursivas a <hi>
    processed_text = processed_text.replace('<<<NOTE>>>', '').replace('<<<ENDNOTE>>>', '')
    
    # Convertir todos los marcadores de cursiva a etiquetas XML válidas
    processed_text = processed_text.replace('<<<ITALIC_START>>>', '<hi rend="italic">')
    processed_text = processed_text.replace('<<<ITALIC_END>>>', '</hi>')
    
    result = processed_text
    
    return result.strip()

def merge_italic_text(text):
    """
    Fusiona SOLO etiquetas cursivas que están completamente pegadas (sin espacios ni contenido entre ellas).
    Versión conservadora para evitar fusiones incorrectas.
    """
    # Busca SOLO etiquetas que están completamente pegadas: </hi><hi rend="italic">
    pattern = re.compile(r'</hi><hi rend="italic">')
    
    # Reemplaza la secuencia </hi><hi rend="italic"> por nada (fusiona las etiquetas)
    result = pattern.sub('', text)
    
    return result

def generate_filename(title):
    """
    Genera un nombre de archivo a partir de las primeras tres palabras del título.
    """
    # Toma las primeras tres palabras del título y elimina espacios, comas y puntos
    words = title.split()[:3]
    filename = '_'.join(words).replace(' ', '_').replace(',', '').replace('.', '')
    return filename

def find_who_id(speaker, characters):
    """
    Busca el xml:id correcto de un personaje en la lista de personajes, usando coincidencia flexible.
    """
    if not characters:
        return ""
    
    # Normalizar el speaker para comparación
    speaker_normalized = speaker.strip().upper()
    
    # 1. Coincidencia exacta (case-insensitive)
    for name, role_id in characters.items():
        if speaker.strip().upper() == name.upper():
            return role_id
    
    # 2. Coincidencia parcial: el speaker está contenido al inicio del nombre completo
    for name, role_id in characters.items():
        name_upper = name.upper()
        # Verificar si el speaker es el primer término (antes de coma)
        if name_upper.startswith(speaker_normalized + ','):
            return role_id
        # O si el speaker coincide con la primera palabra
        first_word = name.split(',')[0].split()[0].upper()
        if speaker_normalized == first_word:
            return role_id
        # O si el speaker coincide con la segunda palabra (para "DON CARLOS")
        words = name.split(',')[0].split()
        if len(words) > 1 and speaker_normalized == words[1].upper():
            return role_id
    
    # 3. Coincidencia difusa (fuzzy matching) como último recurso
    close_matches = get_close_matches(speaker.strip(), characters.keys(), n=1, cutoff=0.6)
    if close_matches:
        return characters[close_matches[0]]

    return ""


def find_who_id_with_fallback(speaker, primary_characters, fallback_characters):
    """
    Busca primero en el dramatis activo del acto y, si no encuentra coincidencia,
    recurre al dramatis global.
    """
    who_id = find_who_id(speaker, primary_characters)
    if who_id:
        return who_id
    return find_who_id(speaker, fallback_characters)


def get_paragraph_style_name(para) -> str:
    """
    Devuelve el nombre de estilo del párrafo o 'Normal' si no existe.
    """
    return para.style.name if para.style else "Normal"


def collect_consecutive_title_paragraphs(paragraphs, start_idx):
    """
    Recoge un bloque consecutivo de párrafos con estilo Titulo_comedia.
    """
    title_paragraphs = []
    idx = start_idx

    while idx < len(paragraphs) and get_paragraph_style_name(paragraphs[idx]) == "Titulo_comedia":
        title_paragraphs.append(paragraphs[idx])
        idx += 1

    return title_paragraphs, idx


def collect_dramatis_block(paragraphs, start_idx):
    """
    Recoge un bloque de dramatis personae: cabecera Epigr_Dramatis y sus entradas.
    """
    if start_idx >= len(paragraphs) or get_paragraph_style_name(paragraphs[start_idx]) != "Epigr_Dramatis":
        return None, [], start_idx

    header_para = paragraphs[start_idx]
    item_paragraphs = []
    idx = start_idx + 1

    while idx < len(paragraphs) and get_paragraph_style_name(paragraphs[idx]) in ["Dramatis_lista", "Prosa"]:
        item_paragraphs.append(paragraphs[idx])
        idx += 1

    return header_para, item_paragraphs, idx


def looks_like_pre_act_sequence(paragraphs, start_idx, require_dramatis=False):
    """
    Detecta si desde una posición arranca un bloque de acto del tipo
    Titulo_comedia (+ opcional segundo Titulo_comedia) + dramatis opcional + Acto.
    """
    idx = start_idx
    saw_title = False

    while idx < len(paragraphs):
        para = paragraphs[idx]
        if is_parse_empty_paragraph(para):
            idx += 1
            continue
        if get_paragraph_style_name(para) != "Titulo_comedia":
            break
        saw_title = True
        idx += 1

    if not saw_title:
        return False

    while idx < len(paragraphs) and is_parse_empty_paragraph(paragraphs[idx]):
        idx += 1

    dramatis_head, _, after_dramatis_idx = collect_dramatis_block(paragraphs, idx)
    if require_dramatis and dramatis_head is None:
        return False
    if dramatis_head is not None:
        idx = after_dramatis_idx
        while idx < len(paragraphs) and is_parse_empty_paragraph(paragraphs[idx]):
            idx += 1

    return idx < len(paragraphs) and get_paragraph_style_name(paragraphs[idx]) == "Acto"


def append_repeated_title_heads(tei, title_paragraphs, nota_notes, aparato_notes, annotation_counter):
    """
    Inserta títulos repetidos de acto como heads anidados dentro del div del acto.
    """
    for idx, title_para in enumerate(title_paragraphs):
        processed_title = extract_text_with_italics_and_annotations(
            title_para,
            nota_notes,
            aparato_notes,
            annotation_counter,
            "head"
        )
        processed_title = uppercase_preserve_tags_and_note_content(processed_title)
        head_type = "mainTitle" if idx == 0 else "subTitle"
        tei.append(f'          <head type="{head_type}" subtype="repeated">{processed_title}</head>')


def open_act_block(tei, state, act_counter):
    """
    Abre el div del acto actual sin imponer todavía el orden interno de sus hijos.
    """
    tei.append(f'        <div type="subsection" subtype="ACTO" n="{act_counter}" xml:id="acto{act_counter}">')
    state["in_act"] = True


def append_act_head(tei, act_para, nota_notes, aparato_notes, annotation_counter):
    """
    Inserta el encabezado del acto respetando el momento en que aparece en Word.
    """
    processed_text = extract_text_with_italics_and_annotations(
        act_para, nota_notes, aparato_notes, annotation_counter, "head"
    )
    processed_text_upper = uppercase_preserve_tags_and_note_content(processed_text)
    tei.append(f'          <head type="acto">{processed_text_upper}</head>')


def append_dramatis_entries(
    tei,
    entry_paragraphs,
    state,
    nota_notes,
    aparato_notes,
    annotation_counter,
    global_characters,
    current_act_characters,
    act_counter,
):
    """
    Inserta las entradas de un dramatis abierto, reutilizando la lógica normal de castList.
    """
    for para in entry_paragraphs:
        style = get_paragraph_style_name(para)
        if style == "Dramatis_lista":
            processed_role_name = extract_text_with_italics_and_annotations(
                para, nota_notes, aparato_notes, annotation_counter, "role"
            )
            role_name = para.text.strip()
            if not role_name:
                continue
            role_name_clean = re.sub(r'@', '', role_name)
            role_slug = normalize_id(role_name_clean)
            if state.get("cast_list_scope") == "act" and state["in_act"]:
                role_id = f"acto{act_counter}_{role_slug}"
                current_act_characters[role_name_clean] = role_id
            else:
                role_id = role_slug
                global_characters[role_name_clean] = role_id
            item_indent = state.get("cast_item_indent", "            ")
            tei.append(f'{item_indent}<castItem><role xml:id="{role_id}">{processed_role_name}</role></castItem>')
        elif style == "Prosa":
            processed_text = extract_text_with_italics_and_annotations(
                para, nota_notes, aparato_notes, annotation_counter, "p"
            )
            item_indent = state.get("cast_item_indent", "            ")
            tei.append(f'{item_indent}<p>{processed_text}</p>')


def open_cast_list_block(tei, state, processed_text, xml_id, inside_act=False):
    """
    Abre un bloque de dramatis personae, ya sea global o asociado al acto actual.
    """
    div_indent = '          ' if inside_act else '        '
    content_indent = '            ' if inside_act else '          '
    item_indent = '              ' if inside_act else '            '

    tei.append(f'{div_indent}<div type="castList" xml:id="{xml_id}">')
    tei.append(f'{content_indent}<head type="castListTitle">{processed_text}</head>')
    tei.append(f'{content_indent}<castList>')
    state["in_cast_list"] = True
    state["cast_list_div_indent"] = div_indent
    state["cast_list_indent"] = content_indent
    state["cast_item_indent"] = item_indent
    state["cast_list_scope"] = "act" if inside_act else "global"


# --- Procesamiento de metadatos y front-matter
def parse_metadata_docx(path, header_mode="prolope"):
    """
    Extrae metadatos de un archivo .docx estructurado en tablas y construye un teiHeader TEI/XML.
    
    Args:
        path: Ruta al archivo DOCX de metadatos.
        header_mode: "prolope" para header completo con datos PROLOPE, 
                     "minimo" para header solo con datos del usuario y referencia a la app.
    """
    doc = Document(path)
    tables = doc.tables

    if len(tables) < 3:
        raise ValueError("❌ El documento de metadatos debe contener al menos 3 tablas.")

    # Tabla 1: Metadatos principales
    main_meta = {}
    for row in tables[0].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            main_meta[key] = value

    # Tabla 2: sourceDesc
    source_meta = {}
    for row in tables[1].rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            source_meta[key] = value

    # Tabla 3: listWit (salta la primera fila)
    witnesses = []
    for i, row in enumerate(tables[2].rows):
        if i == 0:
            continue  # salta la fila guía "SIGLA TESTIMONIO | DESCRIPCIÓN"
        if len(row.cells) >= 2:
            siglum = row.cells[0].text.strip()

            # Extrae el contenido de la segunda celda manteniendo la cursiva
            desc_parts = []
            for run in row.cells[1].paragraphs[0].runs:
                if run.italic:
                    desc_parts.append(f'<hi rend="italic">{run.text}</hi>')
                else:
                    desc_parts.append(run.text)
            desc = ''.join(desc_parts).strip()

            if siglum and desc:
                witnesses.append((siglum, desc))

    # Construcción del teiHeader
    tei = ['<teiHeader>', '  <fileDesc>', '    <titleStmt>']

    tei.append(f'      <title>{main_meta.get("Título comedia", "")}</title>')
    tei.append(f'      <author><name>{main_meta.get("Autor", "")}</name></author>')

    if 'Editor' in main_meta and main_meta['Editor']:
        tei.append(f'      <editor>{main_meta["Editor"]}</editor>')

    if 'Responsable/s revisión' in main_meta and main_meta['Responsable/s revisión']:
        tei.append('      <respStmt>')
        if header_mode == "prolope":
            tei.append('        <resp>Edición crítica digital revisada filológicamente por</resp>')
        else:
            tei.append('        <resp>Responsable/s revisión</resp>')
        for name in main_meta['Responsable/s revisión'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    if 'Responsable marcado automático' in main_meta and main_meta['Responsable marcado automático']:
        tei.append('      <respStmt>')
        if header_mode == "prolope":
            tei.append('        <resp>Marcado XML-TEI automático revisado por</resp>')
        else:
            tei.append('        <resp>Responsable marcado automático</resp>')
        for name in main_meta['Responsable marcado automático'].split(','):
            tei.append(f'        <persName>{name.strip()}</persName>')
        tei.append('      </respStmt>')

    # Solo en modo PROLOPE: añadir referencia al grupo de investigación
    if header_mode == "prolope":
        tei.append('      <respStmt>')
        tei.append('        <resp>Codificado según los criterios de</resp>')
        tei.append('        <name ref="https://datos.bne.es/entidad/XX4849774.html">Grupo de investigación PROLOPE, de la Universitat Autònoma de Barcelona</name>')
        tei.append('      </respStmt>')

    tei.extend(['    </titleStmt>', '    <editionStmt>'])
    tei.append(f'      <edition>Versión {main_meta.get("Versión", "")}</edition>')
    tei.extend(['    </editionStmt>', '    <publicationStmt>'])
    tei.append(f'      <publisher>{main_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'      <pubPlace>{main_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'      <date>{main_meta.get("Fecha publicación", "")}</date>')
    tei.append('    </publicationStmt>')
    
    # Solo en modo PROLOPE: añadir seriesStmt
    if header_mode == "prolope":
        tei.append('    <seriesStmt>')
        tei.append('      <title>Biblioteca Digital PROLOPE</title>')
        tei.append('      <respStmt>')
        tei.append('        <resp>Dirección de</resp>')
        tei.append('        <persName ref="https://orcid.org/0000-0002-7429-9709"><forename>Ramón</forename> <surname>Valdés Gázquez</surname></persName>')
        tei.append('      </respStmt>')
        tei.append('      <idno type="URI">https://bibdigitalprolope.com/</idno>')
        tei.append('    </seriesStmt>')

    # sourceDesc
    tei.extend(['    <sourceDesc>', '      <biblStruct xml:lang="es">', '        <monogr>'])
    
    # Autor
    if header_mode == "prolope":
        tei.append('          <author>')
        tei.append('            <persName ref="http://datos.bne.es/persona/XX1719671"><forename>Félix Lope</forename><surname>de Vega Carpio</surname></persName>')
        tei.append('          </author>')
    else:
        # Modo mínimo: autor desde metadatos
        tei.append(f'          <author>{main_meta.get("Autor", "")}</author>')
    
    # Títulos
    tei.append(f'          <title type="main">{source_meta.get("Título comedia", "")}</title>')
    if "Subtítulo" in source_meta and source_meta.get("Subtítulo"):
        tei.append(f'          <title type="alt">{source_meta.get("Subtítulo", "")}</title>')
    
    # Título del volumen (si existe)
    if source_meta.get("Título volumen"):
        tei.append(f'          <title level="s">{source_meta.get("Título volumen", "")}</title>')
    
    # Solo en modo PROLOPE: añadir "Parte"
    if header_mode == "prolope" and source_meta.get("Parte"):
        tei.append(f'          <title level="a">Parte {source_meta.get("Parte", "")}</title>')
    
    # Editor (si existe)
    if 'Editor' in main_meta and main_meta['Editor']:
        tei.append(f'          <editor>{main_meta["Editor"]}</editor>')
    
    # Coordinadores del volumen (si existen)
    if 'Coordinadores volumen' in source_meta and source_meta['Coordinadores volumen']:
        tei.append('          <respStmt>')
        if header_mode == "prolope":
            tei.append('            <resp>Coordinación del volumen a cargo de</resp>')
        else:
            tei.append('            <resp>Coordinadores volumen</resp>')
        for name in source_meta['Coordinadores volumen'].split(','):
            tei.append(f'            <persName>{name.strip()}</persName>')
        tei.append('          </respStmt>')
    
    # Solo en modo PROLOPE: añadir availability
    if header_mode == "prolope":
        tei.append('          <availability status="restricted">')
        tei.append('            <p>Todos los derechos reservados.</p>')
        tei.append('          </availability>')
    
    # Imprint
    tei.append('          <imprint>')
    tei.append(f'            <pubPlace>{source_meta.get("Lugar publicación", "")}</pubPlace>')
    tei.append(f'            <publisher>{source_meta.get("Publicado por", "")}</publisher>')
    tei.append(f'            <date>{source_meta.get("Fecha publicación", "")}</date>')
    
    # Volumen y páginas (si existen)
    if source_meta.get("Volumen"):
        tei.append(f'            <biblScope unit="volume" n="{source_meta.get("Volumen", "")}">vol. {source_meta.get("Volumen", "")}</biblScope>')
    if source_meta.get("Páginas"):
        tei.append(f'            <biblScope unit="page">{source_meta.get("Páginas", "")}</biblScope>')
    
    tei.append('          </imprint>')
    tei.append('        </monogr>')
    tei.append('      </biblStruct>')
    tei.append('      <listWit>')
    for siglum, desc in witnesses:
        tei.append(f'        <witness xml:id="{siglum}">')
        tei.append(f'          <label>{desc}</label>')
        tei.append('        </witness>')
    tei.append('      </listWit>')
    tei.append('    </sourceDesc>')
    tei.append('  </fileDesc>')
    tei.append('  <encodingDesc>')
    
    # Solo en modo PROLOPE: añadir editorialDecl
    if header_mode == "prolope":
        tei.append('    <editorialDecl>')
        tei.append(f'      <p>El texto se transformó desde archivos DOCX mediante un flujo semiautomático con feniX-ML (versión {APP_VERSION}).</p>')
        tei.append('    </editorialDecl>')
    
    # Siempre incluir appInfo (en ambos modos)
    tei.append('    <appInfo>')
    tei.append(f'      <application ident="feniX-ML" version="{APP_VERSION}">')
    tei.append('        <label>feniX-ML</label>')
    if header_mode == "prolope":
        tei.append('        <desc>Conversor de ediciones críticas de teatro del Siglo de Oro de DOCX a XML-TEI, desarrollado por Anna Abate, Emanuele Leboffe y David Merino Recalde (PROLOPE).</desc>')
        tei.append('        <ref target="https://github.com/prolopeuab/feniX-ML">Repositorio y documentación</ref>')
    else:
        tei.append('        <desc>Conversor de ediciones críticas de DOCX a XML-TEI.</desc>')
        tei.append('        <ref target="https://github.com/prolopeuab/feniX-ML">https://github.com/prolopeuab/feniX-ML</ref>')
    tei.append('      </application>')
    tei.append('    </appInfo>')
    tei.append('  </encodingDesc>')
    tei.append('</teiHeader>')

    return "\n".join(tei)

def process_front_paragraphs_with_tables(front_blocks, footnotes_intro):
    """
    Procesa bloques del front-matter generando secciones TEI <front> con soporte para tablas.

    Maneja diálogos (<sp>), citas (<cit>), versos (<l>), párrafos (<p>) y tablas en cualquier
    sección del prólogo. Usa marcas de estilo (# prefijo, "Quote", "Personaje", "Verso",
    "Partido_inicial", "Partido_medio", "Partido_final" y "Acot").

    Args:
        front_blocks: Lista de bloques del front-matter (párrafos y tablas) en orden real.
        footnotes_intro: Dict de notas al pie del prólogo {id: contenido_html}.

    Returns:
        str: Líneas de XML TEI para el bloque <front>, listas para ser unidas con \\n.
    """
    tei_front = []
    subsection_open = False
    subsection_n = 1
    current_section = None
    paragraph_buffer = []
    head_inserted = False
    in_sp_front = False  # Estado para controlar <sp> en el front

    def flush_paragraph_buffer():
        nonlocal paragraph_buffer, in_sp_front
        for p in paragraph_buffer:
            text = extract_text_with_intro_notes(p, footnotes_intro)
            style = ""
            if p.style:
                style = p.style.name

            if style == "Quote":
                if in_sp_front:
                    tei_front.append('          </sp>')
                    in_sp_front = False
                tei_front.append(f'          <cit rend="blockquote">')
                tei_front.append(f'            <quote>{text}</quote>')
                tei_front.append(f'          </cit>')

            elif style == "Personaje":
                if in_sp_front:
                    tei_front.append('          </sp>')
                tei_front.append('          <sp>')
                tei_front.append(f'            <speaker>{text}</speaker>')
                in_sp_front = True

            elif style in ["Verso", "Partido_inicial", "Partido_medio", "Partido_final"]:
                if text.strip():
                    part_attr = ""
                    if style == "Partido_inicial":
                        part_attr = ' part="I"'
                    elif style == "Partido_medio":
                        part_attr = ' part="M"'
                    elif style == "Partido_final":
                        part_attr = ' part="F"'

                    if in_sp_front:
                        tei_front.append(f'            <l{part_attr}>{text.strip()}</l>')
                    else:
                        tei_front.append(f'          <l{part_attr}>{text.strip()}</l>')

            elif style == "Acot":
                if text.strip():
                    if in_sp_front:
                        tei_front.append(f'            <stage>{text.strip()}</stage>')
                    else:
                        tei_front.append(f'          <stage>{text.strip()}</stage>')

            elif text.strip():
                if in_sp_front:
                    tei_front.append('          </sp>')
                    in_sp_front = False
                tei_front.append(f'          <p>{text.strip()}</p>')

        paragraph_buffer.clear()

        if in_sp_front:
            tei_front.append('          </sp>')
            in_sp_front = False


    for block in front_blocks:
        if isinstance(block, Table):
            flush_paragraph_buffer()
            tei_front.append(process_table_to_tei(block, footnotes_intro, current_section=current_section))
            continue

        para = block
        raw = extract_text_with_intro_notes(para, footnotes_intro)
        text = para.text.strip() if para.text else ""
        if not text:
            continue

        # Ignora "Introducción"
        if text.lower() == "introducción":
            continue

        # Gestión del título principal "PRÓLOGO"
        if not head_inserted and "prólogo" in text.lower():
            flush_paragraph_buffer()
            # Usar 'raw' para preservar notas al pie, quitando el # si existe
            title_prologo = raw.lstrip("#").strip()
            tei_front.append(f'        <head type="divTitle" subtype="MenuLevel_1">{title_prologo}</head>')
            head_inserted = True
            continue

        # Reconocimiento de subtítulo con almohadilla
        if text.startswith("#"):
            flush_paragraph_buffer()
            # Usar 'raw' que ya tiene las notas procesadas, y quitar el # del texto procesado
            title = raw.lstrip("#").strip()
            plain_title = text.lstrip("#").strip()
            if title.lower() == "prólogo":
                continue
            if subsection_open:
                tei_front.append('        </div>')
            tei_front.append(f'        <div type="subsection" n="{subsection_n}">')
            tei_front.append(f'          <head type="divTitle" subtype="MenuLevel_2">{title}</head>')
            subsection_open = True
            current_section = normalize_text_for_matching(plain_title)
            subsection_n += 1
            continue

        # Añade al buffer
        paragraph_buffer.append(para)

    # Procesar cualquier contenido restante
    flush_paragraph_buffer()

    if subsection_open:
        tei_front.append('        </div>')

    return "\n".join(tei_front)

def process_table_to_tei(table, footnotes_intro=None, current_section=None):
    """
    Convierte una tabla DOCX en una tabla TEI, incluyendo notas al pie.
    """
    if footnotes_intro is None:
        footnotes_intro = {}

    ncols = len(table.columns)
    versification_table = is_versification_section(current_section)
    table_attrs = ' type="sinopsisversificacion"' if versification_table else ""
    tei = [f'          <table{table_attrs}>']

    for row_idx, row in enumerate(table.rows):
        texts = [cell.text.strip() for cell in row.cells]
        non_empty = [text for text in texts if text]

        if row_idx == 0:
            tei.append('            <row role="label">')
            for cell in row.cells:
                append_tei_cell(
                    tei,
                    extract_table_cell_contents(cell, footnotes_intro),
                    '              ',
                    attrs=' role="label"'
                )
            tei.append('            </row>')
            continue

        if versification_table and is_versification_act_heading(texts) and len(non_empty) == 1:
            tei.append('            <row role="label">')
            append_tei_cell(
                tei,
                extract_table_cell_contents(row.cells[0], footnotes_intro),
                '              ',
                attrs=f' role="label" cols="{ncols}"'
            )
            tei.append('            </row>')
            continue

        if versification_table and is_versification_summary_header_row(texts):
            tei.append('            <row role="label">')
            for cell in row.cells:
                append_tei_cell(
                    tei,
                    extract_table_cell_contents(cell, footnotes_intro),
                    '              ',
                    attrs=' role="label"'
                )
            tei.append('            </row>')
            continue

        total_row = versification_table and is_versification_total_row(texts)
        row_attrs = ' role="data"'
        if total_row:
            row_attrs += ' rend="summary"'

        tei.append(f'            <row{row_attrs}>')

        for cell_idx, cell in enumerate(row.cells):
            cell_attrs = ' role="data"'
            if total_row and cell_idx == 0:
                cell_attrs = ' role="label"'
            append_tei_cell(
                tei,
                extract_table_cell_contents(cell, footnotes_intro),
                '              ',
                attrs=cell_attrs
            )
        tei.append('            </row>')

    tei.append('          </table>')
    return "\n".join(tei)

# --- Extracción de notas de notas y aparato

def process_annotations_raw(raw_text, nota_notes, aparato_notes, annotation_counter, section):
    """
    Procesa anotaciones en texto plano: @palabra, %palabra o @%palabra.
    
    Busca símbolos de anotación en el texto y genera elementos XML <note> con IDs únicos
    y sincronización de índices. Mantiene sincronización secuencial entre anotaciones
    filológicas y críticas.
    
    Args:
        raw_text: Texto plano con anotaciones sin marcas XML.
        nota_notes: Dict con notas filológicas {palabra_normalizada: contenido o [lista]}.
        aparato_notes: Dict con notas de aparato crítico {palabra_normalizada: contenido o [lista]}.
        annotation_counter: Dict para mantener contadores de cada palabra procesada.
        section: Identificador de sección para generar xml:ids únicos.
    
    Returns:
        str: Texto con anotaciones reemplazadas por palabra + elementos <note> XML integrados.
    """
    # Función para normalizar palabras (sin acentos, minúsculas)
    def normalize_word(word):
        normalized = unicodedata.normalize('NFKD', word)
        normalized = normalized.encode('ASCII', 'ignore').decode('utf-8').lower().strip()
        return normalized
    
    # Crear conjunto de todas las claves normalizadas
    all_keys = set()
    if nota_notes:
        all_keys.update(nota_notes.keys())
    if aparato_notes:
        all_keys.update(aparato_notes.keys())
    
    # Contadores de ocurrencias separados para notas y aparato
    nota_counters = annotation_counter.setdefault("_occurrences_nota", {})
    aparato_counters = annotation_counter.setdefault("_occurrences_aparato", {})
    
    # Función de reemplazo para procesar cada anotación encontrada
    def replace_annotation(match):
        symbol = match.group(1)  # '@', '%' o '@%'
        phrase = match.group(2)  # Palabra sin el símbolo
        key = normalize_word(phrase)
        
        if key not in all_keys:
            # Sin notas, simplemente quitar el símbolo y devolver la palabra
            return phrase
        
        # Construir elementos <note> para esta anotación
        note_str = ""
        
        # Procesar notas filológicas si el símbolo tiene @ (@ o @%)
        if '@' in symbol and key in nota_notes:
            nota_index = nota_counters.get(key, 0)
            # Validación defensiva: asegurar que nota_index nunca sea None
            if nota_index is None:
                nota_index = 0
            nota_list = nota_notes[key] if isinstance(nota_notes[key], list) else [nota_notes[key]]
            if nota_index < len(nota_list):
                content = nota_list[nota_index]
                xml_id_nota = f"n_{key}_{section}_{nota_index + 1}"
                xml_id_nota = re.sub(r'\s+', '_', xml_id_nota)
                xml_id_nota = re.sub(r'[^a-zA-Z0-9_]', '', xml_id_nota)
                xml_id_nota = xml_id_nota.lower()
                note_str += f'<note subtype="nota" xml:id="{xml_id_nota}">{content}</note>'
            nota_counters[key] = nota_index + 1
        
        # === APARATO CRÍTICO === solo si tiene % (% o @%)
        if '%' in symbol and key in aparato_notes:
            aparato_index = aparato_counters.get(key, 0)
            # Validación defensiva: asegurar que aparato_index nunca sea None
            if aparato_index is None:
                aparato_index = 0
            aparato_list = aparato_notes[key] if isinstance(aparato_notes[key], list) else [aparato_notes[key]]
            if aparato_index < len(aparato_list):
                content = aparato_list[aparato_index]
                xml_id_aparato = f"a_{key}_{section}_{aparato_index + 1}"
                xml_id_aparato = re.sub(r'\s+', '_', xml_id_aparato)
                xml_id_aparato = re.sub(r'[^a-zA-Z0-9_]', '', xml_id_aparato)
                xml_id_aparato = xml_id_aparato.lower()
                note_str += f'<note subtype="aparato" xml:id="{xml_id_aparato}">{content}</note>'
            aparato_counters[key] = aparato_index + 1
        
        # Devolver la palabra (ya escapada XML si fuera necesario) seguida de las notas
        return f'{escape_xml(phrase)}{note_str}'
    
    # Patrón regex: @palabra, %palabra o @%palabra
    # Captura: Grupo 1 = símbolo(s), Grupo 2 = palabra
    processed_text = re.sub(r'(@%?|%)(\w+)', replace_annotation, raw_text)
    
    return processed_text

def extract_notes_with_italics(docx_path: str) -> dict:
    """
    Extrae notas o aparato de un DOCX.
    Devuelve un dict donde las claves pueden ser:
    - int: versos normales (ej: 329)
    - str: palabras normalizadas (ej: "dedicatoria") o versos con sufijo alfabético (ej: "329a", "329b")
    
    Los valores son SIEMPRE LISTAS de strings (para manejar múltiples notas secuencialmente).
    
    SUFIJOS ALFABÉTICOS PARA VERSOS PARTIDOS:
    - Los versos partidos usan sufijos: 329a (primera parte), 329b (segunda parte), etc.
    - El sufijo se asigna secuencialmente según el orden de las partes
    - Ejemplo: verso con 3 partes tendría "329a:", "329b:", "329c:" en el archivo de notas
    - Funciona para cualquier número de partes (a-z soporta hasta 26 partes)
    
    FORMATO EN ARCHIVO DE NOTAS:
    - Verso normal: "329: contenido de la nota"
    - Verso partido parte 1: "329a: contenido de la nota"
    - Verso partido parte 2: "329b: contenido de la nota"
    - Palabra (nota filológica): "@dedicatoria: contenido de la nota"
    - Palabra (aparato crítico): "%dedicatoria: contenido de la nota"
    
    SÍMBOLOS PARA ANOTACIONES:
    - @ : para notas filológicas
    - % : para aparato crítico
    En el texto principal se usará:
    - @palabra : solo nota filológica
    - %palabra : solo aparato crítico
    - @%palabra : ambos tipos de notas
    """
    notes: dict = {}
    if not docx_path or not os.path.exists(docx_path):
        return notes

    def normalize_key(word):
        """Normaliza una palabra eliminando acentos y convirtiendo a minúsculas."""
        # Descomponer caracteres con acentos
        normalized = unicodedata.normalize('NFKD', word)
        # Eliminar marcas diacríticas y convertir a minúsculas
        normalized = normalized.encode('ASCII', 'ignore').decode('utf-8').lower().strip()
        return normalized

    doc = Document(docx_path)
    for para in doc.paragraphs:
        text = extract_text_with_italics(para).strip()

        # Notas tipo verso: "1: contenido" o "329a: contenido" (con sufijo alfabético)
        # El sufijo alfabético se usa para versos partidos: 329a, 329b, 329c, etc.
        match_verse = re.match(r'^(\d+[a-z]?):\s*(.*)', text)
        # Notas tipo @palabra o %palabra: "@palabra: contenido" o "%palabra: contenido"
        match_single = re.match(r'^[@%]([^@%]+?):\s*(.*)', text)

        if match_verse:
            verse_key = match_verse.group(1)  # Puede ser "329" o "329a"
            content = match_verse.group(2).strip()
            
            # Si tiene sufijo alfabético, usar como string; si no, convertir a int
            if re.match(r'^\d+[a-z]$', verse_key):
                # Tiene sufijo: usar string como clave (ej: "329a")
                key = verse_key
            else:
                # Sin sufijo: convertir a int para retrocompatibilidad
                key = int(verse_key)
            
            # Siempre usar listas para facilitar el acceso secuencial
            if key not in notes:
                notes[key] = []
            notes[key].append(content)
                
        elif match_single:
            key_original = match_single.group(1).strip()
            content = match_single.group(2).strip()
            
            # Normalizar la clave para insensibilidad a mayúsculas/acentos
            key = normalize_key(key_original)
            
            # Siempre usar listas para facilitar el acceso secuencial
            if key not in notes:
                notes[key] = []
            notes[key].append(content)

    return notes


# --- Procesamiento de notas y aparato

def process_annotations_with_ids(text, nota_notes, aparato_notes, annotation_counter, section):
    """
    Sustituye marcadores @palabra en el texto por notas TEI con xml:ids únicos.
    Usa sincronización secuencial: la 1ª ocurrencia de @palabra recibe la 1ª nota,
    la 2ª ocurrencia recibe la 2ª nota, etc.
    
    - nota_notes y aparato_notes son dicts con claves YA NORMALIZADAS y valores como LISTAS.
    - annotation_counter lleva el conteo de ocurrencias de cada palabra por sección.
    - section es el nombre de la sección (p.ej. 'p', 'speaker', etc.).
    """
    if not text:
        return ""

    # Aseguramos dicts válidos
    nota_notes = nota_notes or {}
    aparato_notes = aparato_notes or {}

    # Función de normalización para palabras del texto (debe coincidir con extract_notes_with_italics)
    def normalize_word(word):
        if isinstance(word, int):
            return word
        # Quitamos tags <hi> y descomponemos acentos
        plain = re.sub(r'<hi rend="italic">(.*?)</hi>', r'\1', word)
        normalized = unicodedata.normalize('NFKD', plain)
        normalized = normalized.encode('ASCII', 'ignore').decode('utf-8').lower().strip()
        return normalized

    # Las claves de las notas ya vienen normalizadas, no necesitamos procesarlas de nuevo
    # Solo necesitamos combinar los dicts
    all_keys = set(nota_notes.keys()) | set(aparato_notes.keys())

    new_text = text.strip()
    
    # Contadores globales de ocurrencias separados para notas y aparato
    nota_counters = annotation_counter.setdefault("_occurrences_nota", {})
    aparato_counters = annotation_counter.setdefault("_occurrences_aparato", {})
    
    # Función de reemplazo para usar con re.sub
    def replace_annotation(match):
        # Determinar qué patrón coincidió y extraer el símbolo y la palabra
        full_match = match.group(0)  # El match completo
        
        # Extraer el símbolo (@, % o @%) del full_match
        symbol = ''
        if full_match.startswith('@%') or full_match.startswith('<hi rend="italic">@%'):
            symbol = '@%'
        elif full_match.startswith('@') or full_match.startswith('<hi rend="italic">@'):
            symbol = '@'
        elif full_match.startswith('%') or full_match.startswith('<hi rend="italic">%'):
            symbol = '%'
        
        # El índice del grupo de captura puede variar según el patrón
        # Para los 3 patrones usados, siempre es el último grupo el que tiene la palabra
        phrase = match.group(match.lastindex)  # La palabra está en el último grupo
        
        # Determinar si la palabra debe estar en cursiva
        # Casos con cursiva: @<hi rend="italic">palabra</hi> o <hi rend="italic">@palabra</hi>
        in_italic = '<hi rend="italic">' in full_match
        
        key = normalize_word(phrase)
        
        if key not in all_keys:
            # Sin notas, solo devolver la palabra con cursiva si la tenía
            if in_italic:
                return f'<hi rend="italic">{phrase}</hi>'
            return phrase
        
        # Construir elementos <note> para cada anotación encontrada
        note_str = ""
        
        # Procesar notas filológicas si el símbolo tiene @ (@ o @%)
        if '@' in symbol and key in nota_notes:
            nota_index = nota_counters.get(key, 0)
            # Validación defensiva: asegurar que nota_index nunca sea None
            if nota_index is None:
                nota_index = 0
            nota_list = nota_notes[key] if isinstance(nota_notes[key], list) else [nota_notes[key]]
            if nota_index < len(nota_list):
                content = nota_list[nota_index]
                xml_id_nota = f'n_{key}_{section}_{nota_index + 1}'
                xml_id_nota = re.sub(r'\s+', '_', xml_id_nota)
                xml_id_nota = re.sub(r'[^a-zA-Z0-9_]', '', xml_id_nota)
                xml_id_nota = xml_id_nota.lower()
                note_str += f'<note subtype="nota" xml:id="{xml_id_nota}">{content}</note>'
            nota_counters[key] = nota_index + 1
        
        # Procesar aparato crítico si el símbolo tiene % (% o @%)
        if '%' in symbol and key in aparato_notes:
            aparato_index = aparato_counters.get(key, 0)
            # Validación defensiva: asegurar que aparato_index nunca sea None
            if aparato_index is None:
                aparato_index = 0
            aparato_list = aparato_notes[key] if isinstance(aparato_notes[key], list) else [aparato_notes[key]]
            if aparato_index < len(aparato_list):
                content = aparato_list[aparato_index]
                xml_id_aparato = f'a_{key}_{section}_{aparato_index + 1}'
                xml_id_aparato = re.sub(r'\s+', '_', xml_id_aparato)
                xml_id_aparato = re.sub(r'[^a-zA-Z0-9_]', '', xml_id_aparato)
                xml_id_aparato = xml_id_aparato.lower()
                note_str += f'<note subtype="aparato" xml:id="{xml_id_aparato}">{content}</note>'
            aparato_counters[key] = aparato_index + 1
        
        # Reconstruir el texto con la nota, manteniendo cursiva si la tenía
        if in_italic:
            return f'<hi rend="italic">{phrase}</hi>{note_str}'
        else:
            return f'{phrase}{note_str}'
    
    # Aplicar el reemplazo en múltiples pasadas para capturar todas las variantes:
    # Pasada 1: @<hi rend="italic">palabra</hi> o %<hi> o @%<hi> (símbolo fuera, palabra dentro)
    def replace_symbol_before_hi(match):
        return replace_annotation(match)
    
    new_text = re.sub(r'(@%?|%)<hi rend="italic">(\w+)</hi>', replace_symbol_before_hi, new_text)
    
    # Pasada 2: <hi rend="italic">@palabra</hi> o <hi>%palabra</hi> o <hi>@%palabra</hi> (ambos dentro)
    def replace_symbol_inside_hi(match):
        return replace_annotation(match)
    
    new_text = re.sub(r'<hi rend="italic">(@%?|%)(\w+)</hi>', replace_symbol_inside_hi, new_text)
    
    # Pasada 3: @palabra, %palabra o @%palabra (sin cursivas)
    def replace_plain_symbol(match):
        return replace_annotation(match)
    
    new_text = re.sub(r'(@%?|%)(\w+)', replace_plain_symbol, new_text)

    # Reagrupamos cursivas consecutivas
    new_text = merge_italic_text(new_text)
    return new_text


# --- Función principal de conversión DOCX → TEI
def convert_docx_to_tei(
    main_docx: str,
    notas_docx: Optional[str] = None,
    aparato_docx: Optional[str] = None,
    metadata_docx: Optional[str] = None,  
    tei_header: Optional[str] = None,
    output_file: Optional[str] = None,
    save: bool = True,
    header_mode: str = "prolope"
) -> Optional[str]:
    """
    Convierte uno o más DOCX a un XML-TEI completo.
    
    Args:
        main_docx: Ruta al archivo DOCX principal.
        notas_docx: Ruta al archivo DOCX con notas (opcional).
        aparato_docx: Ruta al archivo DOCX con aparato crítico (opcional).
        metadata_docx: Ruta al archivo DOCX con metadatos (opcional).
        tei_header: Header TEI personalizado (opcional).
        output_file: Ruta donde guardar el archivo TEI (opcional).
        save: Si se debe guardar el archivo (por defecto True).
        header_mode: "prolope" para header completo, "minimo" para header básico.
    """
    #Chequeo de existencia del principal
    if not main_docx.lower().endswith(".docx"):
        raise ValueError(f"Se esperaba un .docx, pero se obtuvo: {main_docx}")
    if not os.path.exists(main_docx):
        raise FileNotFoundError(f"No existe el archivo principal: {main_docx}")

    # Generación del header TEI a partir de metadata_docx si se proporciona
    if metadata_docx:
        if not os.path.exists(metadata_docx):
            raise FileNotFoundError(f"No existe el archivo de metadatos: {metadata_docx}")
        try:
            tei_header = parse_metadata_docx(metadata_docx, header_mode=header_mode)
        except Exception as e:
            raise RuntimeError(f"No se pudo parsear metadata DOCX '{metadata_docx}': {e}")
    elif not tei_header:
        # Cabecera mínima de reserva
        tei_header_respaldo = "<teiHeader>…</teiHeader>"

    header = tei_header if tei_header else tei_header_respaldo

    # Carga del DOCX principal
    try:
        doc = Document(main_docx)
    except Exception as e:
        raise RuntimeError(f"Error al abrir el archivo DOCX principal '{main_docx}': {e}")

    # --- SEPARACIÓN FRONT/BODY BASADA EN 'Titulo_comedia' ---

    # 1) Buscar todos los párrafos no vacíos con estilo 'Titulo_comedia' (máximo 2: título y subtítulo)
    # Ignora líneas vacías (incluidas las que tengan estilo) y detiene en el primer
    # párrafo no vacío fuera del bloque de título.
    title_paragraphs = []
    found_first_title = False
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""
        is_empty_for_parse = is_parse_empty_paragraph(p)

        if not found_first_title:
            if style_name == "Titulo_comedia":
                if is_empty_for_parse:
                    continue
                title_paragraphs.append(i)
                found_first_title = True
                if len(title_paragraphs) == 2:
                    break
            continue

        # Tras encontrar el primer título, ignoramos líneas vacías entre título/subtítulo
        if is_empty_for_parse:
            continue

        if style_name == "Titulo_comedia":
            if looks_like_pre_act_sequence(doc.paragraphs, i, require_dramatis=True):
                break
            title_paragraphs.append(i)
            if len(title_paragraphs) == 2:
                break
            continue

        # Primer párrafo no vacío que no es Titulo_comedia: fin del bloque de títulos
        break
    
    if not title_paragraphs:
        raise RuntimeError("No se encontró ningún párrafo con estilo 'Titulo_comedia' en el documento")

    title_idx = title_paragraphs[0]
    subtitle_idx = title_paragraphs[1] if len(title_paragraphs) > 1 else None

    # 2) Divide el documento en front y body.
    # El front mantiene el orden real de párrafos y tablas.
    # El body comienza después del último título válido (título o subtítulo).
    last_title_idx = title_paragraphs[-1]
    body_start_idx = last_title_idx + 1
    front_blocks = get_front_blocks(doc, doc.paragraphs[title_idx])
    body_paragraphs  = list(doc.paragraphs[body_start_idx:])

    # --- Extracción del título ---
    raw_title = doc.paragraphs[title_idx].text.strip()
    # Generar la clave/slug a partir del título (sin marcadores @)
    clean_title_for_filename = re.sub(r'@', '', raw_title)
    title_key = generate_filename(clean_title_for_filename)


    # --- Determinación y validación de rutas de notas y aparato ---
    nota_notes = {}
    if notas_docx:
        if not notas_docx.lower().endswith(".docx"):
            raise ValueError(f"El archivo de notas debe ser .docx: {notas_docx}")
        if not os.path.exists(notas_docx):
            raise FileNotFoundError(f"No existe el archivo de notas: {notas_docx}")
        nota_notes = extract_notes_with_italics(notas_docx)

    aparato_notes = {}
    if aparato_docx:
        if not aparato_docx.lower().endswith(".docx"):
            raise ValueError(f"El archivo de aparato debe ser .docx: {aparato_docx}")
        if not os.path.exists(aparato_docx):
            raise FileNotFoundError(f"No existe el archivo de aparato: {aparato_docx}")
        aparato_notes = extract_notes_with_italics(aparato_docx)

    
    # Contadores y estado
    annotation_counter = {}
    state: dict[str, Any] = {
        "in_sp": False,
        "in_cast_list": False,
        "in_dedicatoria": False,
        "in_act": False,
        "cast_list_div_indent": None,
        "cast_list_indent": None,
        "cast_item_indent": None,
        "cast_list_scope": None,
        "pending_split_verse": None,
        "current_split_verse": None,
        "split_verse_part_index": None,
    }
    
    # Dramatis global y dramatis específico del acto actual
    global_characters = {}
    current_act_characters = {}
    
    act_counter = 0
    verse_counter = 1

    # Título procesado con el mismo contador de anotaciones
    title_para = doc.paragraphs[title_idx]
    processed_title = extract_text_with_italics_and_annotations(
        title_para,
        nota_notes,
        aparato_notes,
        annotation_counter,
        "head"
    )
    # Convertir título a mayúsculas preservando etiquetas XML
    processed_title = uppercase_preserve_tags_and_note_content(processed_title)

    # Subtítulo procesado (si existe)
    processed_subtitle = None
    if subtitle_idx is not None:
        subtitle_para = doc.paragraphs[subtitle_idx]
        processed_subtitle = extract_text_with_italics_and_annotations(
            subtitle_para,
            nota_notes,
            aparato_notes,
            annotation_counter,
            "head"
        )
        # Convertir subtítulo a mayúsculas preservando etiquetas XML
        processed_subtitle = uppercase_preserve_tags_and_note_content(processed_subtitle)

    # Notas introductorias
    footnotes_intro = extract_intro_footnotes(main_docx)


    # --- Construcción de <front> y apertura de <body> ---
    tei = [
        '<?xml version="1.0" encoding="UTF-8"?>',  # línea XML
        '<?xml-model href="http://www.tei-c.org/release/xml/tei/custom/schema/relaxng/tei_all.rng" '
        'schematypens="http://relaxng.org/ns/structure/1.0"?>',
        '<TEI xmlns="http://www.tei-c.org/ns/1.0">',
        header,                       # ya generado arriba
        '  <text>',
        '    <front xml:id="front">',
        '      <div type="Introducción" xml:id="prologo">'
    ]

    # Inserta el contenido de <front>, incluyendo notas introductorias y tablas
    tei.append(process_front_paragraphs_with_tables(front_blocks, footnotes_intro))

    # Cerramos el front y abrimos el body con el título principal (y subtítulo si existe)
    tei.extend([
        '      </div>',    # cierra <div type="Introducción">
        '    </front>',
        '    <body xml:id="body">',
        '      <div type="Texto" subtype="TEXTO" xml:id="comedia">',
        f'        <head type="mainTitle" xml:id="titulo">{processed_title}</head>',
    ])
    
    # Añadir subtítulo si existe
    if processed_subtitle:
        tei.append(f'        <head type="subTitle">{processed_subtitle}</head>')


    # Recorre los párrafos significativos del cuerpo con lookahead para detectar
    # títulos repetidos pegados al encabezado de acto.
    significant_body_paragraphs = [para for para in body_paragraphs if not is_parse_empty_paragraph(para)]
    i = 0
    while i < len(significant_body_paragraphs):
        para = significant_body_paragraphs[i]
        style = get_paragraph_style_name(para)

        # Para detección de milestones, usamos texto simple
        text_simple = para.text.strip()

        if style == "Titulo_comedia":
            repeated_titles, next_idx = collect_consecutive_title_paragraphs(significant_body_paragraphs, i)
            if next_idx < len(significant_body_paragraphs):
                dramatis_head, dramatis_entries, after_dramatis_idx = collect_dramatis_block(
                    significant_body_paragraphs, next_idx
                )
                if (
                    dramatis_head is not None
                    and after_dramatis_idx < len(significant_body_paragraphs)
                    and get_paragraph_style_name(significant_body_paragraphs[after_dramatis_idx]) == "Acto"
                ):
                    act_para = significant_body_paragraphs[after_dramatis_idx]
                    close_current_blocks(tei, state, current_act_characters)
                    act_counter += 1
                    open_act_block(tei, state, act_counter)
                    append_repeated_title_heads(
                        tei, repeated_titles, nota_notes, aparato_notes, annotation_counter
                    )

                    processed_dramatis_head = extract_text_with_italics_and_annotations(
                        dramatis_head, nota_notes, aparato_notes, annotation_counter, "head"
                    )
                    open_cast_list_block(
                        tei,
                        state,
                        processed_dramatis_head,
                        f'personajes_acto{act_counter}',
                        inside_act=True
                    )
                    append_dramatis_entries(
                        tei,
                        dramatis_entries,
                        state,
                        nota_notes,
                        aparato_notes,
                        annotation_counter,
                        global_characters,
                        current_act_characters,
                        act_counter,
                    )
                    close_cast_list(tei, state)
                    append_act_head(
                        tei, act_para, nota_notes, aparato_notes, annotation_counter
                    )
                    i = after_dramatis_idx + 1
                    continue

            if next_idx < len(significant_body_paragraphs) and get_paragraph_style_name(significant_body_paragraphs[next_idx]) == "Acto":
                act_para = significant_body_paragraphs[next_idx]
                close_current_blocks(tei, state, current_act_characters)
                act_counter += 1
                open_act_block(tei, state, act_counter)
                append_repeated_title_heads(
                    tei, repeated_titles, nota_notes, aparato_notes, annotation_counter
                )
                append_act_head(
                    tei, act_para, nota_notes, aparato_notes, annotation_counter
                )
                i = next_idx + 1
                continue

            if state.get("in_cast_list"):
                close_cast_list(tei, state)
            i = next_idx
            continue

        if style == "Acto":
            close_current_blocks(tei, state, current_act_characters)
            act_counter += 1
            open_act_block(tei, state, act_counter)
            append_act_head(
                tei, para, nota_notes, aparato_notes, annotation_counter
            )
            i += 1

            if i < len(significant_body_paragraphs) and get_paragraph_style_name(significant_body_paragraphs[i]) == "Titulo_comedia":
                repeated_titles, i = collect_consecutive_title_paragraphs(significant_body_paragraphs, i)
                append_repeated_title_heads(
                    tei, repeated_titles, nota_notes, aparato_notes, annotation_counter
                )
            continue

        if state.get("in_cast_list") and style not in ["Dramatis_lista", "Epigr_Dramatis"]:
            close_cast_list(tei, state)

        # 1) Detección de estrofas marcadas con $tipo de estrofa
        if text_simple.startswith("$"):
            milestone_match = re.match(r'^\$\s*(.+?)\s*$', text_simple)
            if not milestone_match:
                raise ValueError(
                    f"Marcador estrófico inválido: '{text_simple}'. "
                    "Debe contener texto tras '$'."
                )

            raw_milestone_type = milestone_match.group(1)
            milestone_type = normalize_milestone_type(raw_milestone_type)
            if not milestone_type:
                raise ValueError(
                    f"Marcador estrófico inválido: '{text_simple}'. "
                    "Debe contener al menos un carácter alfanumérico tras '$'."
                )
            # Insertar el milestone inmediatamente en la posición actual
            if state["in_sp"]:
                tei.append(f'            <milestone unit="stanza" type="{milestone_type}"/>')
            elif state["in_dedicatoria"]:
                tei.append(f'          <milestone unit="stanza" type="{milestone_type}"/>')
            i += 1
            continue  # saltar el resto del procesamiento para este párrafo

        # 2) Antes de abrir un nuevo bloque estilístico, cerramos los que estén abiertos
        # Nota: Epigr_Dedic no cierra dedicatoria si ya está abierta (para permitir dos head consecutivos)
        if style == "Epigr_Dramatis":
            if state["in_sp"]:
                tei.append('        </sp>')
                state["in_sp"] = False
            close_cast_list(tei, state)
            if state["in_dedicatoria"]:
                tei.append('        </div>')
                state["in_dedicatoria"] = False
        elif style == "Epigr_Dedic" and not state["in_dedicatoria"]:
            # Solo cerrar bloques si no estamos ya en una dedicatoria
            close_current_blocks(tei, state, current_act_characters)


        if style == "Epigr_Dedic":
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "head")
            if not state["in_dedicatoria"]:
                # Primer head de la dedicatoria: abrir div y usar mainTitle
                tei.append('        <div type="dedicatoria" xml:id="dedicatoria">')
                tei.append(f'          <head type="mainTitle">{processed_text}</head>')
                state["in_dedicatoria"] = True
            else:
                # Segundo head consecutivo: usar subTitle (no abrir nuevo div)
                tei.append(f'          <head type="subTitle">{processed_text}</head>')

        elif style == "Epigr_Dramatis":
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "head")
            cast_list_id = f'personajes_acto{act_counter}' if state["in_act"] else "personajes"
            open_cast_list_block(
                tei,
                state,
                processed_text,
                cast_list_id,
                inside_act=state["in_act"]
            )


        elif style == "Dramatis_lista":
            processed_role_name = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "role")
            role_name = para.text.strip()
            if role_name:
                role_name_clean = re.sub(r'@', '', role_name)
                role_slug = normalize_id(role_name_clean)
                if state.get("cast_list_scope") == "act" and state["in_act"]:
                    role_id = f"acto{act_counter}_{role_slug}"
                    current_act_characters[role_name_clean] = role_id
                else:
                    role_id = role_slug
                    global_characters[role_name_clean] = role_id
                item_indent = state.get("cast_item_indent", "            ")
                tei.append(f'{item_indent}<castItem><role xml:id="{role_id}">{processed_role_name}</role></castItem>')

        elif style == "Verso":
            if state["in_dedicatoria"]:
                processed_verse = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "l")
                tei.append(f'          <l>{processed_verse}</l>')
            elif state["in_sp"]:
                verse_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "l")
                
                # Procesar notas
                if verse_counter in nota_notes:
                    note_list = nota_notes[verse_counter]
                    # note_list siempre es una lista
                    for note_idx, content in enumerate(note_list, 1):
                        verse_text += f'<note subtype="nota" n="{verse_counter}" xml:id="nota_{verse_counter}_{note_idx}">{content}</note>'
                
                # Mismo tratamiento para aparato
                if verse_counter in aparato_notes:
                    aparato_list = aparato_notes[verse_counter]
                    # aparato_list siempre es una lista
                    for note_idx, content in enumerate(aparato_list, 1):
                        verse_text += f'<note subtype="aparato" n="{verse_counter}" xml:id="aparato_{verse_counter}_{note_idx}">{content}</note>'
                
                tei.append(f'            <l n="{verse_counter}">{verse_text}</l>')
                verse_counter += 1

        elif style == "Laguna":
            # Laguna de extensión incierta - no incrementa el contador de versos
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "gap")
            if state["in_sp"]:
                tei.append(f'            <gap>{processed_text}</gap>')
            elif state["in_dedicatoria"]:
                tei.append(f'          <gap>{processed_text}</gap>')

        elif style == "Partido_inicial":
            # Iniciar verso partido con sistema de sufijos alfabéticos
            # El sufijo 'a' se asigna a la primera parte, 'b' a la segunda, etc.
            verse_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "l")
            text_simple = para.text.strip()
            
            # Inicializar estado del verso partido
            state["current_split_verse"] = verse_counter
            state["split_verse_part_index"] = 0
            
            # Calcular sufijo alfabético para esta parte (primera parte = 'a')
            letra = chr(97 + state["split_verse_part_index"])  # 97 = 'a' en ASCII
            verse_key_with_suffix = f"{verse_counter}{letra}"
            
            pending_split_verse: PendingSplitVerse = {
                "verse_number": verse_counter,
                "initial_text": text_simple,
                "parts": [text_simple],
                "has_notes": verse_counter in nota_notes or verse_counter in aparato_notes or verse_key_with_suffix in nota_notes or verse_key_with_suffix in aparato_notes
            }
            state["pending_split_verse"] = pending_split_verse
            
            # Procesar notas con clave que incluye sufijo (ej: "329a")
            # Buscar primero con sufijo, luego sin sufijo para retrocompatibilidad
            if verse_key_with_suffix in nota_notes:
                note_list = nota_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(note_list, 1):
                    verse_text += f'<note subtype="nota" n="{verse_key_with_suffix}" xml:id="nota_{verse_counter}{letra}_{note_idx}">{content}</note>'
            elif verse_counter in nota_notes:
                # Retrocompatibilidad: buscar sin sufijo
                note_list = nota_notes[verse_counter]
                for note_idx, content in enumerate(note_list, 1):
                    verse_text += f'<note subtype="nota" n="{verse_key_with_suffix}" xml:id="nota_{verse_counter}{letra}_{note_idx}">{content}</note>'
            
            # Mismo tratamiento para aparato
            if verse_key_with_suffix in aparato_notes:
                aparato_list = aparato_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(aparato_list, 1):
                    verse_text += f'<note subtype="aparato" n="{verse_key_with_suffix}" xml:id="aparato_{verse_counter}{letra}_{note_idx}">{content}</note>'
            elif verse_counter in aparato_notes:
                # Retrocompatibilidad: buscar sin sufijo
                aparato_list = aparato_notes[verse_counter]
                for note_idx, content in enumerate(aparato_list, 1):
                    verse_text += f'<note subtype="aparato" n="{verse_key_with_suffix}" xml:id="aparato_{verse_counter}{letra}_{note_idx}">{content}</note>'
            
            # Incrementar índice de parte para la siguiente parte del verso
            state["split_verse_part_index"] += 1
            
            tei.append(f'            <l part="I" n="{verse_key_with_suffix}">{verse_text}</l>')
            verse_counter += 1

        elif style == "Partido_medio":
            # Procesar parte media del verso partido con sufijo alfabético
            text_simple = para.text.strip()
            verse_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "l")
            
            # Recuperar número base del verso partido
            if state.get("current_split_verse") is not None:
                base_verse = state["current_split_verse"]
                part_index = state.get("split_verse_part_index")
                # Validación defensiva: asegurar que part_index nunca sea None
                if part_index is None:
                    part_index = 1
            else:
                # Fallback: si no hay estado, usar verso anterior (puede ocurrir en secuencias válidas)
                base_verse = verse_counter - 1
                part_index = 1
            
            # Calcular sufijo alfabético para esta parte (segunda parte = 'b', tercera = 'c', etc.)
            letra = chr(97 + part_index)  # 97 = 'a' en ASCII
            verse_key_with_suffix = f"{base_verse}{letra}"
            
            pending = get_pending_split_verse(state)
            if pending is not None:
                pending["parts"].append(text_simple)
            
            # Procesar notas con clave que incluye sufijo (ej: "329b", "329c")
            if verse_key_with_suffix in nota_notes:
                note_list = nota_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(note_list, 1):
                    verse_text += f'<note subtype="nota" n="{verse_key_with_suffix}" xml:id="nota_{base_verse}{letra}_{note_idx}">{content}</note>'
            
            if verse_key_with_suffix in aparato_notes:
                aparato_list = aparato_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(aparato_list, 1):
                    verse_text += f'<note subtype="aparato" n="{verse_key_with_suffix}" xml:id="aparato_{base_verse}{letra}_{note_idx}">{content}</note>'
            
            # Incrementar índice de parte para la siguiente parte
            if "split_verse_part_index" in state and state["split_verse_part_index"] is not None:
                state["split_verse_part_index"] += 1
            else:
                # Si no existe o es None, inicializar a 2 (ya procesamos la primera parte)
                state["split_verse_part_index"] = 2
            
            tei.append(f'            <l part="M" n="{verse_key_with_suffix}">{verse_text}</l>')

        elif style == "Partido_final":
            # Completar el verso partido con sufijo alfabético y limpiar estado
            text_simple = para.text.strip()
            verse_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "l")
            
            # Recuperar número base del verso partido
            if state.get("current_split_verse") is not None:
                base_verse = state["current_split_verse"]
                part_index = state.get("split_verse_part_index")
                # Validación defensiva: asegurar que part_index nunca sea None
                if part_index is None:
                    part_index = 1
            else:
                # Fallback: si no hay estado, usar verso anterior (puede ocurrir en secuencias válidas)
                base_verse = verse_counter - 1
                part_index = 1
            
            # Calcular sufijo alfabético para esta parte final
            letra = chr(97 + part_index)  # 97 = 'a' en ASCII
            verse_key_with_suffix = f"{base_verse}{letra}"
            
            pending = get_pending_split_verse(state)
            if pending is not None:
                pending["parts"].append(text_simple)
                # El verso partido está completo, limpiar estado
                state["pending_split_verse"] = None
            
            # Procesar notas con clave que incluye sufijo (ej: "329c", "329d")
            if verse_key_with_suffix in nota_notes:
                note_list = nota_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(note_list, 1):
                    verse_text += f'<note subtype="nota" n="{verse_key_with_suffix}" xml:id="nota_{base_verse}{letra}_{note_idx}">{content}</note>'
            
            if verse_key_with_suffix in aparato_notes:
                aparato_list = aparato_notes[verse_key_with_suffix]
                for note_idx, content in enumerate(aparato_list, 1):
                    verse_text += f'<note subtype="aparato" n="{verse_key_with_suffix}" xml:id="aparato_{base_verse}{letra}_{note_idx}">{content}</note>'
            
            # Limpiar estado del verso partido
            state["current_split_verse"] = None
            state["split_verse_part_index"] = None
            
            tei.append(f'            <l part="F" n="{verse_key_with_suffix}">{verse_text}</l>')

        elif style == "Acot":
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "stage")
            if state["in_sp"]:
                # Si estamos dentro de un <sp>, insertar el <stage> dentro con la misma indentación que <l>
                tei.append(f'            <stage>{processed_text}</stage>')
            else:
                # Si no hay <sp> abierto, insertar <stage> standalone
                tei.append(f'        <stage>{processed_text}</stage>')

        elif style == "Personaje":
            text_simple = para.text.strip()
            who_id = find_who_id_with_fallback(text_simple, current_act_characters, global_characters)
            processed = extract_text_with_italics_and_annotations(
                para, nota_notes, aparato_notes, annotation_counter, "speaker"
            )

            # Cierra <sp> anterior si es necesario
            if state["in_sp"]:
                tei.append('        </sp>')
                state["in_sp"] = False

            # Abrir <sp> con who si está disponible
            if who_id:
                tei.append(f'        <sp who="#{who_id}">')
            else:
                # No hay who_id (personaje no encontrado en dramatis personae)
                tei.append('        <sp>')
            
            # SIEMPRE insertar <speaker> en cada nuevo <sp>
            # Cada intervención es un <sp> separado y debe tener su propio <speaker>
            # Convertir speaker a mayúsculas, preservando etiquetas XML
            processed_upper = uppercase_preserve_tags_and_note_content(processed)
            tei.append(f'          <speaker>{processed_upper}</speaker>')
            
            state["in_sp"] = True

        elif style == "Prosa":
            # Párrafos en prosa, pueden estar en dedicatoria o en otras secciones
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "p")
            if state["in_dedicatoria"]:
                tei.append(f'          <p>{processed_text}</p>')
            elif state["in_cast_list"]:
                item_indent = state.get("cast_item_indent", "            ")
                tei.append(f'{item_indent}<p>{processed_text}</p>')
            elif state["in_sp"]:
                tei.append(f'            <p>{processed_text}</p>')
            else:
                # Prosa en contexto general
                tei.append(f'        <p>{processed_text}</p>')

        elif style == "Epigr_final":
            processed_text = extract_text_with_italics_and_annotations(para, nota_notes, aparato_notes, annotation_counter, "trailer")
            if processed_text.strip():
                tei.append(f'          <trailer>{processed_text}</trailer>')

        i += 1



    # Cierre final de todos los bloques aún abiertos
    close_current_blocks(tei, state, current_act_characters)

    # Verificar si hay versos partidos incompletos al final del procesamiento
    pending = get_pending_split_verse(state)
    if pending is not None:
        print(f"⚠️ Advertencia: Verso partido incompleto detectado durante procesamiento:")
        print(f"   Verso {pending['verse_number']}: '{pending['initial_text'][:50]}...'")
        print(f"   - Falta Partido_final para completar el verso")

    # Cierre de secciones TEI
    tei.append('      </div>')  # cierra Texto
    tei.append('    </body>')
    tei.append('  </text>')
    tei.append('</TEI>')



    # Serializamos el TEI a string
    tei = [fragment for fragment in tei if isinstance(fragment, str)]
    tei_str = "\n".join(tei)

    # Si no queremos guardar en disco, devolvemos el string
    if not save:
        return tei_str

    # Si llegamos aquí, save == True: escribimos el fichero
    if not output_file:
        # Generamos nombre por defecto si hace falta
        output_file = f"{title_key}.xml"


    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(tei_str)
    # Devolvemos None para indicar que se escribió en disco
    return None


# --- Validación y análisis de los documentos
def count_verses_in_document(main_docx, include_dedication=False):
    """
    Cuenta los versos en un documento DOCX usando la misma lógica que el procesamiento principal.
    
    Args:
        main_docx: Ruta al archivo DOCX
        include_dedication: Si True, cuenta versos desde Titulo_comedia; si False, desde primer Acto
    
    Returns:
        Lista de tuplas (paragraph_index, verse_number, style, text) para cada verso encontrado
    """
    doc = Document(main_docx)
    verses = []
    found_start = False
    verse_counter = 1
    
    for para_idx, para in enumerate(doc.paragraphs):
        style: str = para.style.name if para.style else ""
        text = para.text.strip() if para.text else ""
        
        # Determinar punto de inicio según parámetro
        if not found_start:
            start_style = "Titulo_comedia" if include_dedication else "Acto"
            if style == start_style and not is_parse_empty_paragraph(para):
                found_start = True
                if not include_dedication:  # Si empezamos en Acto, reiniciar contador
                    verse_counter = 1
            continue
        
        # Aplicar los mismos filtros que en el procesamiento principal
        if is_parse_empty_paragraph(para):  # Párrafos vacíos para parseo
            continue
        if re.match(r'^\$\w+', text):  # Milestones
            continue
        if style in [
            "Personaje", "Acot", "Prosa", 
            "Epigr_Dedic", "Epigr_Dramatis", "Dramatis_lista", "Epigr_final",
            "Acto", "Cita", "Heading 1", "Heading 2", "Heading 3", "Normal"
        ]:
            continue
        
        # Contar versos reales (excluyendo Laguna que no incrementa numeración)
        if style == "Verso":
            verses.append((para_idx, verse_counter, style, text))
            verse_counter += 1
        elif style == "Partido_inicial":
            verses.append((para_idx, verse_counter, style, text))
            verse_counter += 1
        elif style in ["Partido_medio", "Partido_final"]:
            # Medio y final usan el número del inicial (counter - 1)
            verses.append((para_idx, verse_counter - 1, style, text))
        elif style == "Laguna":
            # Registrar laguna pero sin incrementar contador
            verses.append((para_idx, verse_counter, style, text))
    
    return verses

def get_verse_number_at_position(main_docx, target_para_index, include_dedication=False):
    """
    Obtiene el número del último verso antes de una posición específica en el documento.
    
    Args:
        main_docx: Ruta al archivo DOCX
        target_para_index: Índice del párrafo objetivo
        include_dedication: Si True, cuenta versos desde Titulo_comedia; si False, desde primer Acto
    
    Returns:
        int: Número del último verso antes de la posición, o 0 si no hay versos previos
    """
    verses = count_verses_in_document(main_docx, include_dedication)
    
    # Buscar el último verso antes de la posición objetivo
    last_verse_number = 0
    for para_idx, verse_number, style, text in verses:
        if para_idx < target_para_index:
            # Solo contar versos que incrementan el contador (no medio/final)
            if style in ["Verso", "Partido_inicial"]:
                last_verse_number = verse_number
        else:
            break
    
    return last_verse_number

def is_parse_empty_paragraph(para) -> bool:
    """
    Determina si un párrafo debe considerarse vacío durante el parseo.
    Solo considera vacíos los párrafos con blancos/caracteres invisibles.
    La puntuación o cualquier otro carácter visible cuenta como contenido.
    """
    if para is None:
        return True

    # Para parseo, priorizamos runs concatenados; fallback a para.text
    if para.runs:
        raw_text = "".join((run.text or "") for run in para.runs)
    else:
        raw_text = para.text or ""

    if not raw_text:
        return True

    # Espacios visibles, tabs/saltos y rango de espacios invisibles Unicode
    invisible_pattern = r'^[\s\u00A0\u2000-\u200D\u2028-\u202F\u205F\u3000\uFEFF]*$'
    return re.match(invisible_pattern, raw_text) is not None

def is_empty_paragraph(para) -> bool:
    """
    Determina si un párrafo está vacío o solo contiene espacios/caracteres invisibles.
    """
    if not para.text:
        return True
    
    text = para.text.strip()
    if not text:
        return True
    
    # Verificar espacios invisibles, tabulaciones y caracteres especiales
    clean_text = text.replace('\u00A0', '').replace('\t', '').strip()
    if not clean_text:
        return True
    
    # Verificar si todos los runs están vacíos
    if not para.runs or all(not run.text.strip() for run in para.runs):
        return True
    
    # Verificar líneas con solo puntuación, espacios o caracteres de control
    if re.match(r'^[\s\.:!?\(\)"\']+$', text):
        return True
    
    # Verificar líneas con solo caracteres de control o espacios no separables
    if re.match(r'^[\s\u00A0\u2000-\u200F\u2028-\u202F]+$', text):
        return True
    
    return False

def should_skip_paragraph(para: Paragraph, text: str, style: str) -> bool:
    """
    Determina si un párrafo debe ser omitido durante la validación.
    """
    # Párrafos vacíos
    if is_empty_paragraph(para):
        return True

    # Milestones que empiezan con '$'
    if re.match(r'^\$\s*\S+', text):
        return True

    # Front-matter con almohadilla (títulos de sección del prólogo)
    if text.startswith("#"):
        return True

    # Párrafos dentro de tablas (sinopsis, metadatos, etc.)
    if para._element.xpath("ancestor::w:tbl"):
        return True

    return False

def analyze_main_text(main_docx) -> list[str]:
    """
    Analiza el archivo principal y devuelve avisos de párrafos sin estilo
    solo en el cuerpo de la obra (tras Titulo_comedia), ignorando front matter
    y milestones ($.). Incluye dramatis personae pero no cuenta versos.
    """
    warnings: list[str] = []
    unstyled_paragraphs: list[tuple[str, str]] = []  # (text, location_info)

    doc = Document(main_docx)
    found_body = False
    last_act_name = None

    for para_idx, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip() if para.text else ""

        # 1) Buscamos el inicio del body (incluyendo dramatis personae)
        if not found_body:
            if style == "Titulo_comedia":
                found_body = True
            continue

        # Registrar actos para determinar ubicación
        if style == "Acto":
            last_act_name = text

        # 2) Aplicamos los filtros comunes para detectar párrafos problemáticos
        if should_skip_paragraph(para, text, style):
            continue

        # 3) Solo revisamos estilos 'Normal' o None para párrafos sin estilo
        if style in ["Normal", ""]:
            # Obtener número del último verso antes de esta posición
            last_verse = get_verse_number_at_position(main_docx, para_idx, include_dedication=False)
            
            # Determinar el contexto de localización
            if last_verse > 0:
                # Hay versos antes, está dentro de un acto
                location_info = f" (después del verso {last_verse})"
            elif last_act_name:
                # Hay un acto definido pero aún no hay versos
                location_info = f" (en {last_act_name}, antes del primer verso marcado)"
            else:
                # Está antes del primer acto (en dramatis personae o dedicatoria)
                location_info = " (en dramatis personae o dedicatoria)"
            
            unstyled_paragraphs.append((text, location_info))

    # 4) Tras recorrer todas las líneas, añadimos el warning si hay alguno
    if unstyled_paragraphs:
        count = len(unstyled_paragraphs)
        warnings.append(
            f"❌ LÍNEAS SIN ESTILO DETECTADAS ({count})\n"
            f"   Archivo: {os.path.basename(main_docx)}\n"
            f"   Revisa que todas las líneas tengan el estilo correcto aplicado."
        )
        
        # Añadir detalles de cada párrafo sin estilo (mostrar máximo 5)
        for i, (text, location) in enumerate(unstyled_paragraphs[:5], 1):
            snippet = text[:60] + "..." if len(text) > 60 else text
            warnings.append(f"   {i}. «{snippet}»{location}")
        
        if len(unstyled_paragraphs) > 5:
            remaining = len(unstyled_paragraphs) - 5
            warnings.append(f"   ... y {remaining} {'línea más' if remaining == 1 else 'líneas más'}")

    return warnings




def analyze_notes(
    notes: dict, 
    note_type: str
) -> list[str]:
    """
    Analiza el dict de notas (aparato o nota) y devuelve
    una lista de strings con posibles notas mal formateadas o múltiples.
    """
    warnings: list[str] = []

    for key, content in notes.items():
        # key puede ser int (verso) o str (palabra normalizada)
        
        # Verificar si hay múltiples notas para la misma clave
        if isinstance(content, list) and len(content) > 1:
            if isinstance(key, str):
                # Determinar el símbolo correcto según el tipo de nota
                symbol = '@' if note_type.lower() == 'nota' else '%'
                # Notas @palabra o %palabra con múltiples entradas
                warnings.append(
                    f"⚠️ MÚLTIPLES {note_type.upper()}S PARA '{symbol}{key}' ({len(content)})\n"
                    f"   Verifica que la asignación secuencial en el texto sea correcta."
                )
            elif isinstance(key, int):
                # Notas de verso con múltiples entradas
                warnings.append(
                    f"⚠️ MÚLTIPLES {note_type.upper()}S PARA VERSO {key} ({len(content)})\n"
                    f"   Verifica que todas las {note_type}s sean correctas."
                )
        
        # Validar que el contenido no esté vacío
        if isinstance(content, list):
            for i, text in enumerate(content, 1):
                if not text.strip():
                    if isinstance(key, str):
                        warnings.append(f"❌ {note_type.capitalize()} vacía: '@{key}' (entrada #{i})")
                    elif isinstance(key, int):
                        warnings.append(f"❌ {note_type.capitalize()} vacía: verso {key} (entrada #{i})")
        elif not str(content).strip():
            # Caso de contenido no-lista (por compatibilidad)
            if isinstance(key, str):
                warnings.append(f"❌ {note_type.capitalize()} vacía: '@{key}'")
            elif isinstance(key, int):
                warnings.append(f"❌ {note_type.capitalize()} vacía: verso {key}")

    return warnings



def validate_split_verses_impact_on_numbering(main_docx) -> list[str]:
    """
    Valida que los versos partidos incompletos no afecten la numeración total.
    Compara el número esperado de versos completos vs el número real.
    """
    warnings: list[str] = []
    
    # Obtener todos los versos
    verses = count_verses_in_document(main_docx, include_dedication=False)
    
    total_verses = 0  # Versos normales
    split_verse_initials = 0  # Partido_inicial (cada uno debería ser un verso)
    split_verse_groups = 0  # Grupos completos de versos partidos (I + [M...] + F)
    incomplete_splits = 0  # Partido_inicial sin Partido_final
    
    # Separar por tipo de verso
    verse_sequence = [(style, text) for _, _, style, text in verses]
    
    # 2. Analizar secuencias de versos partidos
    i = 0
    while i < len(verse_sequence):
        style, text = verse_sequence[i]
        
        if style == "Verso":
            total_verses += 1
            i += 1
        elif style == "Partido_inicial":
            split_verse_initials += 1
            j = i + 1
            found_final = False
            
            # Buscar el final correspondiente
            while j < len(verse_sequence):
                next_style, _ = verse_sequence[j]
                if next_style == "Partido_final":
                    found_final = True
                    split_verse_groups += 1
                    break
                elif next_style in ["Partido_inicial", "Verso"]:
                    break
                j += 1
            
            if not found_final:
                incomplete_splits += 1
            
            i = j + 1 if found_final else i + 1
        else:  # Partido_medio, Partido_final
            i += 1
    
    # 3. Validaciones - Solo mostrar desajuste si existe
    expected_total_verses = total_verses + split_verse_groups
    actual_verse_increments = total_verses + split_verse_initials
    
    if actual_verse_increments != expected_total_verses:
        diff = actual_verse_increments - expected_total_verses
        plural_s = 's' if abs(diff) > 1 else ''
        warnings.append(
            f"\n⚠️ DESAJUSTE EN LA NUMERACIÓN DE VERSOS\n"
            f"Se esperan {expected_total_verses} pero se numerarán {actual_verse_increments}. "
            f"Hay una diferencia de {diff:+d} verso{plural_s} debido a versos partidos incompletos."
        )
    
    return warnings

def validate_split_verses(main_docx) -> list[str]:
    """
    Valida que los versos partidos sigan la secuencia lógica correcta:
    - Partido_inicial debe tener al menos un Partido_final después
    - Entre Partido_inicial y Partido_final puede haber 0 o más Partido_medio
    - No puede haber Partido_medio o Partido_final sin Partido_inicial previo
    """
    warnings: list[str] = []
    verse_problems: list[tuple[int, str, str]] = []  # (verse_num, text, problem_description)
    
    # Obtener todos los versos con su numeración correcta
    verses = count_verses_in_document(main_docx, include_dedication=False)
    
    # 2. Validar secuencia de versos partidos y recopilar problemas
    i = 0
    while i < len(verses):
        para_idx, verse_num, style, text = verses[i]
        
        if style == "Partido_inicial":
            # Buscar el final correspondiente
            j = i + 1
            found_final = False
            middle_count = 0
            
            while j < len(verses):
                next_para_idx, next_verse_num, next_style, next_text = verses[j]
                
                if next_style == "Partido_medio" and next_verse_num == verse_num:
                    middle_count += 1
                elif next_style == "Partido_final" and next_verse_num == verse_num:
                    found_final = True
                    break
                elif next_style in ["Verso", "Partido_inicial"]:
                    # Nuevo verso antes de encontrar final
                    break
                j += 1
            
            if not found_final:
                snippet = text[:50] + "..." if len(text) > 50 else text
                verse_problems.append((
                    verse_num,
                    snippet,
                    "Falta 'Partido_final' después de 'Partido_inicial'."
                ))
            
            # Avanzar hasta después del grupo procesado
            i = j + 1 if found_final else i + 1
            
        elif style == "Partido_medio":
            # Partido_medio sin Partido_inicial previo
            has_initial = False
            for k in range(i - 1, -1, -1):
                prev_para_idx, prev_verse_num, prev_style, prev_text = verses[k]
                if prev_style == "Partido_inicial" and prev_verse_num == verse_num:
                    has_initial = True
                    break
                elif prev_style in ["Verso"] or (prev_style == "Partido_inicial" and prev_verse_num != verse_num):
                    break
            
            if not has_initial:
                snippet = text[:50] + "..." if len(text) > 50 else text
                verse_problems.append((
                    verse_num,
                    snippet,
                    "Falta 'Partido_inicial' antes de este 'Partido_medio'."
                ))
            i += 1
            
        elif style == "Partido_final":
            # Partido_final sin Partido_inicial previo
            has_initial = False
            for k in range(i - 1, -1, -1):
                prev_para_idx, prev_verse_num, prev_style, prev_text = verses[k]
                if prev_style in ["Partido_inicial", "Partido_medio"] and prev_verse_num == verse_num:
                    has_initial = True
                    break
                elif prev_style in ["Verso"] or (prev_style in ["Partido_inicial", "Partido_final"] and prev_verse_num != verse_num):
                    break
            
            if not has_initial:
                snippet = text[:50] + "..." if len(text) > 50 else text
                verse_problems.append((
                    verse_num,
                    snippet,
                    "Falta 'Partido_inicial' o 'Partido_medio' previo."
                ))
            i += 1
            
        else:  # style == "Verso"
            i += 1
    
    # 3. Construir mensaje consolidado si hay problemas
    if verse_problems:
        count = len(verse_problems)
        plural_s = 's' if count > 1 else ''
        
        # Encabezado
        message = f"❌ ({count}) VERSO{plural_s.upper()} PARTIDO{plural_s.upper()} INCOMPLETO{plural_s.upper()}\n"
        message += "Revisa los siguientes versos:\n\n"
        
        # Listar cada verso con su problema
        for verse_num, text, problem in verse_problems:
            message += f"Verso {verse_num}. Texto: {text}\n"
            message += f"      Problema: {problem}\n\n"
        
        warnings.append(message.rstrip())
    
    return warnings


def validate_Laguna(main_docx) -> list[str]:
    """
    Valida que las lagunas marcadas como Laguna no sean versos específicos perdidos
    que deberían marcarse como Verso normal para mantener la numeración.
    """
    warnings: list[str] = []
    
    doc = Document(main_docx)
    found_body = False
    
    for para_idx, para in enumerate(doc.paragraphs):
        style: str = para.style.name if para.style else ""
        text = para.text.strip() if para.text else ""
        
        # Esperar hasta el inicio del cuerpo principal
        if not found_body:
            if style == "Acto":
                found_body = True
            continue
        
        if style == "Laguna":
            # Obtener el número de verso en la posición actual
            verse_num = get_verse_number_at_position(main_docx, para_idx, include_dedication=False)
            
            # Contar total de versos para contexto
            total_verses = len([v for v in count_verses_in_document(main_docx, include_dedication=False) 
                              if v[2] in ["Verso", "Partido_inicial"]])
            
            snippet = text[:50] + "..." if len(text) > 50 else text
            warnings.append(
                f"⚠️ LAGUNA DETECTADA (después del verso {verse_num})\n"
                f"   Texto: {snippet}\n"
                f"   Revisa: ¿Es una laguna de extensión incierta o un verso específico faltante?\n"
                f"   Si es un verso específico faltante, márcalo como 'Verso' con corchetes.\n"
                f"   Total de versos actual: {total_verses}"
            )
    
    return warnings


def validate_verso_con_corchetes(main_docx) -> list[str]:
    """
    Valida que los versos marcados como 'Verso' que contienen solo corchetes
    no sean lagunas que deberían marcarse como 'Laguna' para no contar en la numeración.
    """
    warnings: list[str] = []
    
    doc = Document(main_docx)
    found_body = False
    
    # Patrón para detectar texto que consiste principalmente en corchetes con puntos o puntos suspensivos
    import re
    # Incluye tanto puntos normales (.) como puntos suspensivos (…)
    corchetes_pattern = re.compile(r'^\s*\[[\.…]{1,}\]\s*$|^\s*\[\s*[\.…\s]+\s*\]\s*$')
    
    for para_idx, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip() if para.text else ""
        
        # Esperar hasta el inicio del cuerpo principal
        if not found_body:
            if style in ["Titulo_comedia", "Acto"]:
                found_body = True
            continue
        
        if style == "Verso" and corchetes_pattern.match(text):
            # Obtener el número de verso en la posición actual
            verse_num = get_verse_number_at_position(main_docx, para_idx, include_dedication=False)
            
            # Contar total de versos para contexto
            total_verses = len([v for v in count_verses_in_document(main_docx, include_dedication=False) 
                              if v[2] in ["Verso", "Partido_inicial"]])
            
            warnings.append(
                f"⚠️ VERSO CON CORCHETES DETECTADO (verso {verse_num})\n"
                f"   Texto: {text}\n"
                f"   Revisa: ¿Es una laguna de extensión incierta?\n"
                f"   Si no sabes cuántos versos faltan, márcalo como 'Laguna' para no contarlo.\n"
                f"   Total de versos actual: {total_verses}"
            )
    
    return warnings


def validate_note_format(docx_path: str, note_type: str) -> list[str]:
    """
    Valida que todas las entradas en el archivo de notas o aparato crítico
    sigan el formato correcto:
    - NÚMERO: contenido (ej: 6: 6 cursiva y van bien centradas...)
    - @PALABRA: contenido (para notas filológicas, ej: @dedicatoria: Esta es una nota...)
    - %PALABRA: contenido (para aparato crítico, ej: %dedicatoria: Variante...)
    
    Devuelve una lista de warnings con las entradas que no cumplan el formato.
    """
    warnings: list[str] = []
    
    if not docx_path or not os.path.exists(docx_path):
        return warnings
    
    # Obtener el nombre del archivo para mostrarlo en el mensaje
    filename = os.path.basename(docx_path)
    
    doc = Document(docx_path)
    
    # Patrón para validar el formato correcto
    # Debe comenzar con número:, @palabra: o %palabra:
    pattern_verse = re.compile(r'^\d+[a-z]?:\s*')  # Números seguidos opcionalmente de letra y :
    pattern_nota = re.compile(r'^@[^@%\s]+:\s*')  # @palabra seguido de :
    pattern_aparato = re.compile(r'^%[^@%\s]+:\s*')  # %palabra seguido de :
    
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        
        # Ignorar párrafos vacíos o con solo espacios en blanco
        if not text:
            continue
        
        # Verificar si el párrafo cumple alguno de los formatos válidos
        is_verse_format = pattern_verse.match(text)
        is_nota_format = pattern_nota.match(text)
        is_aparato_format = pattern_aparato.match(text)
        
        if not is_verse_format and not is_nota_format and not is_aparato_format:
            # El párrafo no cumple ninguno de los formatos válidos
            snippet = text[:80] + "..." if len(text) > 80 else text
            warnings.append(
                f"❌ Formato incorrecto en archivo '{filename}' ({note_type}, párrafo {i}): "
                f"Debe comenzar con 'NÚMERO:', '@PALABRA:' o '%PALABRA:' → Texto: {snippet}"
            )
    
    return warnings


def validate_documents(main_docx, aparato_docx=None, notas_docx=None) -> list[str]:
    """
    Ejecuta las comprobaciones sobre los DOCX y devuelve una lista
    de strings con los avisos encontrados (vacía si no hay warnings).
    """
    warnings: list[str] = []

    # 1) Comprueba existencia del principal
    if not os.path.exists(main_docx):
        warnings.append(f"❌ No existe el archivo principal: {main_docx}")
        return warnings

    # 2) Validación de estilos en el body
    ESTILOS_VALIDOS = {
        "Titulo_comedia", "Acto", "Prosa", "Verso", "Partido_inicial",
        "Partido_medio", "Partido_final", "Personaje", "Acot",
        "Epigr_Dedic", "Epigr_Dramatis", "Dramatis_lista", "Epigr_final",
        "Laguna"  # Nuevo estilo para lagunas de extensión incierta
    }
    # Estilos que se omiten en esta validación básica porque tienen validación específica
    SKIP_STYLES = {"Cita", "Heading 1", "Heading 2", "Heading 3", "Normal"}
    doc = Document(main_docx)
    found_body = False

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip() if para.text else ""

        # 2.1) Esperar hasta el inicio de cuerpo
        if not found_body:
            if style in ("Titulo_comedia", "Acto"):
                found_body = True
            continue

        # 2.2) Aplicar filtros comunes para omitir párrafos
        if should_skip_paragraph(para, text, style):
            continue
        
        # 2.3) Omitir estilos específicos que no necesitan validación
        if style in SKIP_STYLES:
            continue

        # 2.4) Validar estilo permitido (solo si no es un párrafo a omitir)
        if style not in ESTILOS_VALIDOS:
            snippet = text[:60]
            warnings.append(f"❌ Estilo no válido: {style or 'None'} — Texto: {snippet}")

    # 3) Análisis avanzado del texto principal (detección de párrafos sin estilo)
    warnings.extend(analyze_main_text(main_docx))

    # 4) Notas de aparato
    if aparato_docx:
        if not os.path.exists(aparato_docx):
            warnings.append(f"❌ El archivo de notas de aparato: {aparato_docx}")
        else:
            # Validar formato de entrada (NÚMERO: o @PALABRA:)
            warnings.extend(validate_note_format(aparato_docx, "aparato crítico"))
            # Validar contenido de las notas
            aparato_notes = extract_notes_with_italics(aparato_docx)
            warnings.extend(analyze_notes(aparato_notes, "aparato"))

    # 5) Notas
    if notas_docx:
        if not os.path.exists(notas_docx):
            warnings.append(f"❌ El archivo de notas no existe: {notas_docx}")
        else:
            # Validar formato de entrada (NÚMERO: o @PALABRA:)
            warnings.extend(validate_note_format(notas_docx, "notas"))
            # Validar contenido de las notas
            nota_notes = extract_notes_with_italics(notas_docx)
            warnings.extend(analyze_notes(nota_notes, "nota"))

    # 6) Validación de versos partidos
    warnings.extend(validate_split_verses(main_docx))
    warnings.extend(validate_split_verses_impact_on_numbering(main_docx))

    # 7) Validación de lagunas marcadas como Laguna
    warnings.extend(validate_Laguna(main_docx))

    # 8) Validación de versos con corchetes que podrían ser lagunas
    warnings.extend(validate_verso_con_corchetes(main_docx))

    return warnings


